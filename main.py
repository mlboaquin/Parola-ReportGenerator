# =========================================================
# Patent Report Generator - Main Application
# =========================================================
# This application generates patent reports (Invalidity, FTO, etc.) by processing
# Excel data files and populating Word document templates with the extracted information.
# =========================================================

# Standard library imports
import sys
import os
import re
from datetime import datetime
from copy import deepcopy
from queue import Queue
import io

# Third-party imports for data processing and document manipulation
import pandas as pd
from io import BytesIO
import requests
from bs4 import BeautifulSoup
import difflib

# Feb10: for Excel date handling via displayed formats
try:
    from openpyxl import load_workbook
    from openpyxl.utils.datetime import from_excel
except Exception:
    load_workbook = None
    from_excel = None

# Python-docx imports for Word document processing
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.text.paragraph import Paragraph
from docx.table import Table
try:
    import msoffcrypto
except Exception:
    msoffcrypto = None

# PyQt6 imports for GUI application
from PyQt6.QtCore import QTimer, QThread, pyqtSignal as Signal, Qt
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QTextEdit, QFileDialog, QProgressBar,
    QMessageBox, QComboBox
)

def unlock_password_protected_docx(file_bytes, password):
    """Decrypt password-protected Word file"""
    try:
        if msoffcrypto is None:
            print("⚠ Note: msoffcrypto not available, assuming file is not password-protected")
            return io.BytesIO(file_bytes)
            
        encrypted_file = io.BytesIO(file_bytes)
        decrypted_file = io.BytesIO()

        office_file = msoffcrypto.OfficeFile(encrypted_file)
        office_file.load_key(password=password)
        office_file.decrypt(decrypted_file)

        decrypted_file.seek(0)
        print("✓ Template unlocked successfully!")
        return decrypted_file
    except Exception as e:
        print(f"⚠ Note: {e}")
        print("If file is not password-protected, this is normal. Proceeding...")
        # Return original bytes if decryption fails (file might not be protected)
        return io.BytesIO(file_bytes)

def extract_mapping_section(edited_doc):
    """Extract the Mapping section from edited document"""
    if not edited_doc:
        return []
    body = edited_doc.element.body
    elements = []
    start_found = False
    
    # Debug: Print all paragraph text to help identify the correct patterns
    print("DEBUG: Searching for mapping section in edited document...")
    for i, child in enumerate(list(body)):
        if child.tag == qn('w:p'):
            p = Paragraph(child, edited_doc)
            text = p.text.strip()
            if text and len(text) < 100:  # Only print short text (likely headings)
                print(f"  Paragraph {i}: '{text}'")
    
    for child in list(body):
        if child.tag == qn('w:p'):
            p = Paragraph(child, edited_doc)
            text_lower = p.text.lower().strip()
            # Match exactly on the section header text
            if not start_found and "mappings based on selected references" in text_lower:
                print(f"DEBUG: Found mapping section start with text: '{p.text}'")
                start_found = True
                # Include the header itself
                elements.append(child)
                continue
        if start_found:
            if child.tag == qn('w:p'):
                p = Paragraph(child, edited_doc)
                text_lower = p.text.lower().strip()
                # Stop at Disclaimer section - everything after should be regenerated
                # Note: Appendices and Search Strategies are regenerated from Excel
                if "disclaimer" in text_lower:
                    print(f"DEBUG: Found mapping section end with text: '{p.text}'")
                    break
            # Include all elements (paragraphs, tables, drawings, etc.)
            if child.tag == qn('w:p'):
                p = Paragraph(child, edited_doc)
                # Check if this paragraph contains page breaks by looking at XML
                has_page_break = False
                try:
                    for run_element in child.xpath('.//w:br'):
                        break_type = run_element.get(qn('w:type'))
                        if break_type == 'page':
                            has_page_break = True
                            break
                except Exception:
                    pass
                
                text_content = p.text.strip()
                if text_content or has_page_break:
                    # Include paragraph if it has text OR if it contains page breaks
                    elements.append(child)
            else:
                # Include all other elements (tables, drawings, etc.)
                elements.append(child)
    
    print(f"DEBUG: Extracted {len(elements)} elements from mapping section")
    
    # If no elements found, try a more general search
    if len(elements) == 0:
        print("DEBUG: No mapping section found with specific patterns, trying general search...")
        for i, child in enumerate(list(body)):
            if child.tag == qn('w:p'):
                p = Paragraph(child, edited_doc)
                text_lower = p.text.lower().strip()
                if "mappings based" in text_lower:
                    print(f"  Found potential mapping header at paragraph {i}: '{p.text[:100]}...'")
    
    return [deepcopy(el) for el in elements]

def extract_criteria_section(edited_doc):
    """Extract the Criteria for Publication Search section from edited document"""
    if not edited_doc:
        return []
    body = edited_doc.element.body
    elements = []
    start_found = False
    
    # Debug: Print all paragraph text to help identify the correct patterns
    print("DEBUG: Searching for criteria section in edited document...")
    for i, child in enumerate(list(body)):
        if child.tag == qn('w:p'):
            p = Paragraph(child, edited_doc)
            text = p.text.strip()
            if text and len(text) < 100:  # Only print short text (likely headings)
                print(f"  Paragraph {i}: '{text}'")
    
    for child in list(body):
        if child.tag == qn('w:p'):
            p = Paragraph(child, edited_doc)
            text_lower = p.text.lower().strip()
            # Match exactly on the section header text
            if not start_found and "criteria for the publication search" in text_lower:
                print(f"DEBUG: Found criteria section start with text: '{p.text}'")
                start_found = True
                # Skip the header itself, start collecting from next element
                continue
        if start_found:
            if child.tag == qn('w:p'):
                p = Paragraph(child, edited_doc)
                text_lower = p.text.lower().strip()
                # Look for mapping section header as end marker
                if "mappings based on selected references" in text_lower:
                    print(f"DEBUG: Found criteria section end with text: '{p.text}'")
                    break
            # Only append paragraphs that are not empty OR contain page breaks
            if child.tag == qn('w:p'):
                p = Paragraph(child, edited_doc)
                
                # Check if this paragraph contains page breaks by looking at XML
                has_page_break = False
                try:
                    for run_element in child.xpath('.//w:br'):
                        break_type = run_element.get(qn('w:type'))
                        if break_type == 'page':
                            has_page_break = True
                            break
                except Exception:
                    pass
                
                if p.text.strip() or has_page_break:
                    elements.append(child)
            # Include all tables
            elif child.tag == qn('w:tbl'):
                elements.append(child)
    
    print(f"DEBUG: Extracted {len(elements)} elements from criteria section")
    
    # If no elements found, try a more general search
    if len(elements) == 0:
        print("DEBUG: No criteria section found with specific patterns, trying general search...")
        for i, child in enumerate(list(body)):
            if child.tag == qn('w:p'):
                p = Paragraph(child, edited_doc)
                text_lower = p.text.lower().strip()
                if "criteria for" in text_lower:
                    print(f"  Found potential criteria header at paragraph {i}: '{p.text[:100]}...'")
    
    return [deepcopy(el) for el in elements]

def remove_section(doc, start_key, end_key):
    """Remove a section from document between two markers"""
    body = doc.element.body
    to_remove = []
    start_found = False
    for child in list(body):
        if child.tag == qn('w:p'):
            p = Paragraph(child, doc)
            text_lower = p.text.lower().strip()
            if not start_found and start_key in text_lower:
                start_found = True
                to_remove.append(child)
                continue
        if start_found:
            if child.tag == qn('w:p'):
                p = Paragraph(child, doc)
                text_lower = p.text.lower().strip()
                if end_key in text_lower:
                    break
            to_remove.append(child)
    for el in to_remove:
        body.remove(el)

def insert_element_after(anchor, element):
    """Insert element after anchor, sanitizing relationships"""
    try:
        new_elem = deepcopy(element)
        # Sanitize relationships/hyperlinks copied from another document
        def _strip_relationship_ids(node):
            # Remove any r:id attributes to avoid broken relationships
            # BUT preserve image relationships by skipping certain element types
            for el in node.iter():
                # Don't strip relationship IDs from inline drawings/images
                if 'pict' in el.tag.lower() or 'drawing' in el.tag.lower() or 'inline' in el.tag.lower():
                    continue
                # Remove relationship id attributes if present
                for attr in list(el.attrib.keys()):
                    if attr.endswith('}id'):
                        del el.attrib[attr]
        def _flatten_hyperlinks(node):
            # Replace w:hyperlink elements with their children runs to keep text
            to_replace = []
            for el in list(node.iter()):
                if el.tag.endswith('}hyperlink'):
                    to_replace.append(el)
            for hl in to_replace:
                parent = hl.getparent()
                if parent is None:
                    continue
                idx = list(parent).index(hl)
                for child in list(hl):
                    parent.insert(idx, child)
                    idx += 1
                parent.remove(hl)
        _flatten_hyperlinks(new_elem)
        _strip_relationship_ids(new_elem)
        anchor._element.addnext(new_elem)
        if new_elem.tag == qn('w:p'):
            return Paragraph(new_elem, anchor._parent)
        elif new_elem.tag == qn('w:tbl'):
            from docx.table import Table
            return Table(new_elem, anchor._parent)
        else:
            # Return the new element even if we can't wrap it
            return new_elem
    except Exception as e:
        print(f"Error in insert_element_after: {e}")
        return None

class PatentReportGenerator:
    """
    Main class for generating patent reports from Excel data and Word templates.
    
    This class handles the complete process of:
    1. Loading Excel data files containing patent information
    2. Loading Word document templates with placeholders
    3. Extracting and processing patent data, claims, and references
    4. Populating Word documents with formatted content
    5. Generating final reports (Invalidity, FTO, etc.)
    
    Attributes:
        log_callback: Function to handle logging messages
        progress_callback: Function to handle progress updates
        report_type: Type of report being generated (Invalidity, FTO, etc.)
        df: Pandas DataFrame containing Excel data
        doc: Word document object for template processing
        excel_filename: Name of the loaded Excel file
        template_filename: Name of the loaded Word template
        global_color_index: Counter for consistent color cycling in mappings
    """
    
    def __init__(self, log_callback, progress_callback, report_type, update_mode=False, edited_report_path=None, template_password="parolatools"):
        """
        Initialize the PatentReportGenerator.
        
        Args:
            log_callback: Function to handle logging messages
            progress_callback: Function to handle progress updates  
            report_type: Type of report being generated (Invalidity, FTO, etc.)
            update_mode: Whether to use update mode (merge with edited report)
            edited_report_path: Path to edited report for update mode
            template_password: Password for password-protected templates
        """
        self.log_callback = log_callback
        self.progress_callback = progress_callback
        self.report_type = report_type
        self.update_mode = update_mode
        self.edited_report_path = edited_report_path
        self.template_password = template_password
        self.df = None
        self.doc = None
        self.edited_doc = None
        self.gen_doc = None  # For generating fresh sections in update mode
        self.excel_filename = None
        self.template_filename = None
        self.global_color_index = 0  # For consistent color cycling across claims
        # Feb10: openpyxl worksheet for precise date formatting via Excel number_format
        self.ws = None

    class Reference:
        """
        Data class representing a patent or non-patent literature reference.
        
        This class stores all relevant information about a reference found in the
        patent search, including publication details, assignee information, and
        classification as patent or non-patent literature (NPL).
        """
        def __init__(self):
            self.PublicationNumber = ""      # Cleaned publication number
            self.PriorityDate = ""           # Priority filing date
            self.FilingDate = ""             # Filing date
            self.PublicationDate = ""        # Publication date
            self.OriginalAssignee = ""       # Original assignee
            self.CurrentAssignee = ""         # Current assignee
            self.Title = ""                  # Patent/Publication title
            self.URL = ""                    # URL to the reference
            self.Rank = ""                   # Ranking (A, B, C, etc.)
            self.isNPL = None                # Boolean: True if Non-Patent Literature
            self.PublicationName = None      # Formatted publication name
            self.RawPublicationNumber = ""   # Original publication number from Excel

    def log(self, message):
        """Log a message using the configured callback function."""
        self.log_callback(message)

    def progress(self, value, message=""):
        """Update progress using the configured callback function."""
        self.progress_callback(value, message)

    def load_excel(self, file_path):
        """
        Load Excel file containing patent data.
        
        Args:
            file_path: Path to the Excel file to load
        """
        self.log("Loading Excel file...")
        try:
            self.df = pd.read_excel(file_path, header=None)
            self.excel_filename = os.path.basename(file_path)
            self.log("Excel file loaded successfully.")

            # Feb10: also load workbook via openpyxl so we can honor Excel's displayed date formats
            self.ws = None
            if load_workbook is not None:
                try:
                    with open(file_path, "rb") as f:
                        excel_bytes = f.read()
                    wb = load_workbook(BytesIO(excel_bytes), data_only=True)
                    self.ws = wb.active
                except Exception as e:
                    self.log(f"Warning: Could not load workbook with openpyxl for precise date formatting: {str(e)}")
        except Exception as e:
            self.log(f"Error loading Excel file: {str(e)}")
            raise

    def load_template(self, file_path):
        """
        Load Word document template.
        
        Args:
            file_path: Path to the Word template file to load
        """
        self.log("Loading Word template...")
        try:
            try:
                # Store decrypted bytes for gen_doc creation in update mode
                with open(file_path, 'rb') as f:
                    template_bytes = f.read()
                self.doc = Document(io.BytesIO(template_bytes))
                self.template_filename = os.path.basename(file_path)
                self.log("Word template loaded successfully.")
                # Store bytes for gen_doc creation
                self.template_bytes = template_bytes
            except Exception as e_normal:
                if msoffcrypto is None:
                    self.log(f"Note: msoffcrypto not available. Error opening template: {str(e_normal)}")
                    raise
                try:
                    with open(file_path, 'rb') as f:
                        encrypted_bytes = f.read()
                    encrypted_file = io.BytesIO(encrypted_bytes)
                    decrypted_file = io.BytesIO()
                    office_file = msoffcrypto.OfficeFile(encrypted_file)
                    office_file.load_key(password=self.template_password)
                    office_file.decrypt(decrypted_file)
                    decrypted_file.seek(0)
                    self.doc = Document(decrypted_file)
                    self.template_filename = os.path.basename(file_path)
                    self.log("Template unlocked and loaded successfully.")
                    # Store decrypted bytes for gen_doc creation
                    self.template_bytes = decrypted_file.getvalue()
                except Exception as e_unlock:
                    self.log(f"Error unlocking template: {str(e_unlock)}")
                    raise e_normal
        except Exception as e:
            self.log(f"Error loading Word template: {str(e)}")
            raise
    
    def setup_update_mode_documents(self):
        """
        Set up document structure for update mode.
        Matches the colab implementation: use edited_doc as base, gen_doc for fresh sections.
        """
        if self.update_mode and self.edited_doc is not None:
            # Use edited as base to preserve images in preserved sections
            old_doc = self.doc
            self.doc = self.edited_doc
            self.log("✓ Using edited report as base document (update mode)")
            
            # Prepare a separate generated document from blank template
            if hasattr(self, 'template_bytes') and self.template_bytes:
                decrypted_blank = unlock_password_protected_docx(self.template_bytes, self.template_password)
                self.gen_doc = Document(decrypted_blank)
                self.log("✓ Prepared fresh document for regenerated sections")
            else:
                self.gen_doc = old_doc
                self.log("✓ Using existing document as gen_doc")
        else:
            # Non-update mode: just use the loaded template
            self.log("✓ Blank template loaded and ready!")
    
    def get_target_doc(self, section_name="general"):
        """
        Get the target document for section generation.
        In update mode, regenerated sections go to gen_doc, preserved sections to doc.
        Matches colab lines 970-976.
        """
        if self.update_mode and self.gen_doc is not None:
            # Regenerated sections go to gen_doc (Title, Objectives, References, Patent-at-Issue, Criteria, Search Strings)
            if section_name in ["title", "objectives", "references", "patent", "criteria", "search"]:
                return self.gen_doc
            else:
                # Preserved sections (Mappings) stay in doc
                return self.doc
        else:
            # Non-update mode: use doc for everything
            return self.doc

    def load_edited_report(self):
        """Load the edited report for update mode"""
        if not self.update_mode or not self.edited_report_path:
            return
            
        try:
            with open(self.edited_report_path, 'rb') as f:
                file_bytes = f.read()
            
            # Try to unlock password-protected document
            decrypted_bytes = unlock_password_protected_docx(file_bytes, self.template_password)
            self.edited_doc = Document(decrypted_bytes)
            self.log("Edited report loaded successfully.")
            
            # Debug: Print some basic info about the edited document
            self.log(f"DEBUG: Edited document has {len(self.edited_doc.paragraphs)} paragraphs")
            self.log(f"DEBUG: Edited document has {len(self.edited_doc.tables)} tables")
            
            # Print first few paragraph texts to help debug
            self.log("DEBUG: First 10 paragraphs in edited document:")
            for i, p in enumerate(self.edited_doc.paragraphs[:10]):
                text = p.text.strip()
                if text:
                    self.log(f"  {i}: '{text[:100]}{'...' if len(text) > 100 else ''}'")
                    
        except Exception as e:
            self.log(f"Error loading edited report: {str(e)}")
            self.edited_doc = None

    def get_short_patent_name_with_suffix(self, patent_number):
        """Updated to handle non-US patents"""
        cleaned = patent_number.replace(',', '').replace(' ', '')

        # Check if US patent and remove US prefix
        is_us_patent = cleaned.upper().startswith('US')
        if is_us_patent:
            cleaned = cleaned[2:]

        match = re.match(r'([^\d]*)(\d+)([A-Z]\d{1,2})?$', cleaned)
        if not match:
            digits_only = re.sub(r'\D', '', cleaned)
            short_num = digits_only[-3:] if len(digits_only) >= 3 else digits_only
            return f"'{short_num} Patent"

        prefix, digits, suffix = match.groups()
        short_digits = digits[-3:] if len(digits) >= 3 else digits

        return f"'{short_digits} Patent"

    def get_short_patent_name_v2(self, patent_number):
        """Updated to handle non-US patents"""
        cleaned = patent_number.replace(',', '').replace(' ', '')

        # Check if US patent and remove US prefix
        is_us_patent = cleaned.upper().startswith('US')
        if is_us_patent:
            cleaned = cleaned[2:]

        match = re.match(r'([^\d]*)(\d+)([A-Z]\d{1,2})?$', cleaned)

        if not match:
            digits_only = re.sub(r'\D', '', cleaned)
            short_num = digits_only[-3:] if len(digits_only) >= 3 else digits_only
            return f"'{short_num}"

        prefix, digits, suffix = match.groups()
        short_digits = digits[-3:] if len(digits) >= 3 else digits

        return f"'{short_digits}"

    def format_patent_display(self, patent_number, include_prefix=True):
        """
        Format patent number for display based on country code
        - US patents: "U.S. Patent No. 10,123,456"
        - Non-US patents: "CN12345678A" (NO prefix, just raw number)
        """
        if pd.isna(patent_number) or not patent_number:
            return ""

        patent_str = str(patent_number).strip()

        # Check if it's a US patent
        if patent_str.upper().startswith("US"):
            core = patent_str[2:]  # Remove "US" prefix
            # Remove any suffix (A1, B2, etc.) for display
            match = re.match(r'(\d+)', core)
            if match:
                digits = match.group(1)
                formatted = self.format_number_with_commas(digits)
                if include_prefix:
                    return f"U.S. Patent No. {formatted}"
                else:
                    return formatted
            return patent_str
        else:
            # Non-US patent - return JUST the publication number, NO prefix
            return patent_str

    def extract_patent_number(self, cell_value):
        """
        Extract patent number from cell value that may contain extra text.
        Handles various patent number formats including US patents with optional suffixes.
        """
        if pd.isna(cell_value):
            return ""

        cell_str = str(cell_value).strip()

        # Pattern to match US patent numbers with optional suffix
        pattern = r'(US\d{7,11}(?:[A-Z]\d{1,2})?)'
        match = re.search(pattern, cell_str, re.IGNORECASE)
        if match:
            return match.group(1).upper()

        # Fallback for other country codes
        pattern_generic = r'([A-Z]{2}\d{7,}(?:[A-Z]\d{1,2})?)'
        match_generic = re.search(pattern_generic, cell_str, re.IGNORECASE)
        if match_generic:
            return match_generic.group(1).upper()

        return cell_str

    def parse_claim_numbers(self, claim_input, patent_number):
        """
        Parse claim numbers from various input formats:
        - "1-30" -> [1, 2, 3, ..., 30]
        - "All" or "all" -> fetch all claims from Google Patents
        - "1, 5, 10" -> [1, 5, 10]
        - "1-5, 10, 15-20" -> [1, 2, 3, 4, 5, 10, 15, 16, 17, 18, 19, 20]
        """
        if pd.isna(claim_input):
            return []

        claim_str = str(claim_input).strip().upper()

        # Handle "All" case
        if claim_str == "ALL":
            return self.get_all_claim_numbers_from_google(patent_number)

        claim_numbers = []
        parts = [part.strip() for part in claim_str.split(',')]

        for part in parts:
            if '-' in part:
                range_parts = part.split('-')
                if len(range_parts) == 2:
                    try:
                        start = int(range_parts[0].strip())
                        end = int(range_parts[1].strip())
                        claim_numbers.extend(range(start, end + 1))
                    except ValueError:
                        continue
            else:
                try:
                    num_match = re.search(r'\d+', part)
                    if num_match:
                        claim_numbers.append(int(num_match.group()))
                except ValueError:
                    continue

        claim_numbers = sorted(list(set(claim_numbers)))
        return [str(num) for num in claim_numbers]

    def get_all_claim_numbers_from_google(self, patent_number):
        """Fetch all claim numbers from Google Patents"""
        try:
            url = f"https://patents.google.com/patent/{patent_number}/en"
            response = requests.get(url, timeout=10)
            if response.status_code != 200:
                return []

            soup = BeautifulSoup(response.content, "html.parser")
            claims_div = soup.find("section", itemprop="claims")
            if not claims_div:
                return []

            claims_text = claims_div.get_text(separator="\n")
            claim_lines = claims_text.strip().split("\n")

            claim_numbers = []
            for line in claim_lines:
                match = re.match(r'^(\d+)\.', line.strip())
                if match:
                    claim_numbers.append(match.group(1))

            return claim_numbers
        except Exception:
            return []

    def format_claims_as_ranges(self, claim_numbers_list):
        """
        Convert list of claim numbers to range format.
        Example: [1,2,3,5,6,10] -> "1-3, 5, 6, and 10"
        Example: [1,2,3,4,5] -> "1-5"
        """
        if not claim_numbers_list:
            return ""

        # Convert to integers and sort
        nums = sorted([int(c) for c in claim_numbers_list])

        if len(nums) == 1:
            return str(nums[0])

        ranges = []
        start = nums[0]
        end = nums[0]

        for i in range(1, len(nums)):
            if nums[i] == end + 1:
                end = nums[i]
            else:
                if start == end:
                    ranges.append(str(start))
                elif end == start + 1:
                    ranges.append(str(start))
                    ranges.append(str(end))
                else:
                    ranges.append(f"{start}-{end}")
                start = nums[i]
                end = nums[i]

        # Handle last range
        if start == end:
            ranges.append(str(start))
        elif end == start + 1:
            ranges.append(str(start))
            ranges.append(str(end))
        else:
            ranges.append(f"{start}-{end}")

        # Format with commas and "and"
        if len(ranges) == 1:
            return ranges[0]
        elif len(ranges) == 2:
            return f"{ranges[0]} and {ranges[1]}"
        else:
            return ", ".join(ranges[:-1]) + f", and {ranges[-1]}"

    def clean_publication_number(self, pub_num):
        pub_num = str(pub_num)
        match = re.match(r'(US)(\d+)', pub_num)
        if match:
            return match.group(1) + match.group(2)
        else:
            return pub_num

    def format_date(self, row, col):
        """
        Feb10: Return the date as Excel displays it, using the cell's number_format
        when possible. row/col are 0-based indices into self.df (same as df.iloc).
        Falls back to legacy string-based handling when openpyxl context is absent.
        """
        # Primary path: use openpyxl worksheet if available
        if getattr(self, "ws", None) is not None and from_excel is not None:
            try:
                cell = self.ws.cell(row=row + 1, column=col + 1)  # openpyxl is 1-based
                v = cell.value
                fmt = (cell.number_format or "").lower()

                if v in (None, ""):
                    return ""

                # Excel serial number -> datetime
                if isinstance(v, (int, float)):
                    try:
                        v = from_excel(v)
                    except Exception:
                        return str(v)

                from datetime import datetime as _dt
                if isinstance(v, _dt):
                    # Month + Year (no day), e.g. "January 2014"
                    if ("mmmm" in fmt or "mmm" in fmt) and "yyyy" in fmt and "d" not in fmt:
                        return v.strftime("%B %Y") if "mmmm" in fmt else v.strftime("%b %Y")
                    # Year-only
                    if "yyyy" in fmt and "m" not in fmt and "d" not in fmt:
                        return v.strftime("%Y")
                    # Default: Day Month Year without leading zero
                    out = v.strftime("%d %B %Y")
                    return out[1:] if out.startswith("0") else out

                # If Excel stored it as text already, return text (strip leading day zero if present)
                s = str(v).strip()
                if s.lower() == "nan":
                    return ""
                if re.match(r"^0\d\s", s):
                    return s[1:]
                return s
            except Exception:
                # Fall through to legacy behavior below
                pass

        # Fallback: legacy behavior using raw value from dataframe (pre-Feb10 appV6 logic)
        try:
            date_val = self.df.iloc[row, col]
        except Exception:
            return ""

        if pd.isna(date_val):
            return ""

        date_str = str(date_val).strip()
        if date_str.lower() == "nan" or date_str == "":
            return ""

        # If already in textual full-date form like "13 June 2008", normalize leading zero
        if re.match(r"^\d{1,2}\s+[A-Za-z]+\s+\d{4}$", date_str):
            return date_str.lstrip("0")

        # Preserve partial dates exactly (Month YYYY or YYYY)
        if re.match(r"^[A-Za-z]+\s+\d{4}$", date_str) or re.match(r"^\d{4}$", date_str):
            return date_str

        # If not a pandas Timestamp, attempt to parse string dates (e.g., "2008-06-13 00:00:00")
        if not isinstance(date_val, pd.Timestamp):
            parsed = pd.to_datetime(date_str, errors="coerce")
            if isinstance(parsed, pd.Timestamp) and not pd.isna(parsed):
                try:
                    return parsed.strftime("%-d %B %Y").strip()
                except Exception:
                    out = parsed.strftime("%d %B %Y").strip()
                    return out[1:] if out and out[0] == "0" else out
            # Fallback for unparsed text: strip leading zero day if present
            if re.match(r"^0\d\s", date_str):
                return date_str[1:]
            return date_str

        # Timestamp case: format as "Day Month Year" without leading zero
        try:
            formatted = date_val.strftime("%-d %B %Y").strip()
        except Exception:
            formatted = date_val.strftime("%d %B %Y").strip()
            if formatted and formatted[0] == "0":
                formatted = formatted[1:]
        return formatted

    def apply_font_style(self, paragraph, size=10, bold=False):
        for run in paragraph.runs:
            run.font.name = 'Inter'
            run.font.size = Pt(size)
            run.bold = bold

    def isUSPatent(self, ref):
        """Updated to properly format both US and non-US patents"""
        if ref.PublicationNumber.startswith("US"):
            # US Patent logic (existing)
            name = ref.PublicationNumber[2:]
            name = re.sub(r'[^\d]', '', name)
            if len(name) >= 11 and name[:4].isdigit():
                year = name[:4]
                pub_seq = name[4:11]
                ref.PublicationName = f"{year}/{pub_seq}"
            else:
                try:
                    number = int(name.lstrip("0"))
                    ref.PublicationName = "{:,}".format(number)
                except:
                    ref.PublicationName = name
        else:
            # Non-US Patent - keep as-is
            ref.PublicationName = ref.PublicationNumber

    def fetch_abstract(self, publication_number):
        url = f"https://patents.google.com/patent/{publication_number}/en"
        try:
            response = requests.get(url, timeout=10)
            if response.status_code != 200:
                return "Abstract not found."
            soup = BeautifulSoup(response.content, "html.parser")
            abstract_tag = soup.find("meta", {"name": "DC.description"})
            if abstract_tag:
                return abstract_tag.get("content", "Abstract not found.")
            return "Abstract not found."
        except Exception as e:
            self.log(f"Error fetching abstract: {str(e)}")
            return "Abstract fetch error."

    def extract_claim_fragments_from_excel(self, df):
        """
        Feb10: Extract claim fragments starting from any recognized Expert/Reviewer Comments row.
        Supports multiple header labels instead of only 'Expert Comments'.
        """
        claim_parts = []
        try:
            expert_comments_variations = ['Expert Comments', 'Expert/Reviewer Comments', 'Reviewer Comments']
            claim_start_idx = None

            for variation in expert_comments_variations:
                try:
                    claim_start_idx = df[df[0] == variation].index[0] + 1
                    break
                except (IndexError, KeyError):
                    continue

            if claim_start_idx is None:
                self.log("Error: No Expert/Reviewer Comments row found in Excel file.")
                return []

            while claim_start_idx < len(df):
                cell_value = df.iloc[claim_start_idx, 0]
                if pd.isna(cell_value) or str(cell_value).strip() == "":
                    break
                claim_parts.append(str(cell_value).strip())
                claim_start_idx += 1
        except Exception as e:
            self.log(f"Error extracting claim fragments from Excel: {str(e)}")
        return claim_parts

    def format_number_with_commas(self, number_str):
        match = re.match(r"\d+", number_str)
        if match:
            digits = match.group()
            return "{:,}".format(int(digits))
        else:
            return number_str

    def get_claim_from_google_patents(self, patent_number, claim_num):
        try:
            url = f"https://patents.google.com/patent/{patent_number}/en"
            response = requests.get(url, timeout=10)
            if response.status_code != 200:
                self.log(f"Failed to fetch claim {claim_num}: HTTP {response.status_code}")
                return None
            soup = BeautifulSoup(response.content, "html.parser")
            claims_div = soup.find("section", itemprop="claims")
            if not claims_div:
                self.log(f"No claims section found for claim {claim_num}")
                return None
            claims_text = claims_div.get_text(separator="\n")
            claim_lines = claims_text.strip().split("\n")
            claim_full_text = ""
            collect = False
            for line in claim_lines:
                if re.match(rf"^{claim_num}\.", line.strip()):
                    collect = True
                elif re.match(r"^\d+\.", line.strip()) and not line.strip().startswith(f"{claim_num}."):
                    if collect:
                        break
                if collect:
                    processed_line = re.sub(
                        r'\bUS(\d{7,})\b',
                        lambda m: self.format_number_with_commas(m.group(1)),
                        line.strip()
                    )
                    claim_full_text += processed_line + "\n"
            return claim_full_text.strip()
        except Exception as e:
            self.log(f"Error fetching claim {claim_num}: {str(e)}")
            return None

    def insert_paragraph_after(self, paragraph, text=""):
        new_p = OxmlElement('w:p')
        paragraph._p.addnext(new_p)
        para = Paragraph(new_p, paragraph._parent)
        if text:
            run = para.add_run(text)
            run.font.name = 'Inter'
        return para

    def insert_table_after_paragraph(self, doc, table, paragraph):
        tbl_el = table._tbl
        body = doc.element.body
        body.remove(tbl_el)
        paragraph._p.addnext(tbl_el)

    def delete_row(self, table, row):
        table._tbl.remove(row._tr)

    def clone_table_structure(self, source_table):
        """
        Clone the entire table structure for creating new mapping tables.
        
        Args:
            source_table: The table to clone
            
        Returns:
            A new table with the same structure as the source table
        """
        new_tbl = deepcopy(source_table._tbl)
        from docx.table import Table
        return Table(new_tbl, source_table._parent)

    def populate_table_with_claim(self, table, claim_number, color_cycle):
        """
        Fill table with claim data using the improved approach from the notebook.
        
        Args:
            table: The table to populate
            claim_number: The claim number to process
            color_cycle: List of colors for alternating claim elements
        """
        claim_fragments, fragment_rows = self.get_claim_fragments_for_claim(claim_number)
        
        # Filter empty fragments
        filtered_fragments = []
        filtered_rows = []
        for i, fragment in enumerate(claim_fragments):
            if fragment.strip():
                filtered_fragments.append(fragment)
                filtered_rows.append(fragment_rows[i] if i < len(fragment_rows) else -1)
        
        claim_fragments = filtered_fragments
        fragment_rows = filtered_rows
        
        # CRITICAL: Save the template data row BEFORE deleting anything
        template_row_copy = None
        if len(table.rows) > 1:
            template_row_copy = deepcopy(table.rows[1]._tr)
        
        # Remove all data rows (keep header)
        for row in list(table.rows[1:]):
            table._tbl.remove(row._tr)
        
        # Add rows for fragments using the saved template
        for frag_idx, fragment in enumerate(claim_fragments):
            # Clone from the saved template row
            if template_row_copy is not None:
                new_tr = deepcopy(template_row_copy)
                table._tbl.append(new_tr)
                from docx.table import _Row
                new_row = _Row(new_tr, table)
            else:
                new_row = table.add_row()
            
            self.clear_cell_keep_formatting(new_row.cells[0])
            self.clear_cell_keep_formatting(new_row.cells[1])
            
            # Left cell
            p_left = new_row.cells[0].paragraphs[0]
            run_left = p_left.add_run(fragment)
            run_left.font.name = 'Inter'
            run_left.font.size = Pt(9)
            run_left.bold = True
            run_left.font.color.rgb = color_cycle[self.global_color_index % 2]
            self.global_color_index += 1
            if frag_idx > 0:
                p_left.paragraph_format.left_indent = Inches(0.23)
            
            # Right cell
            if fragment.strip() and fragment_rows[frag_idx] != -1:
                for p in list(new_row.cells[1].paragraphs):
                    p._element.getparent().remove(p._element)
                
                main_para = new_row.cells[1].add_paragraph()
                
                # Set default font for the paragraph to prevent Calibri fallback
                self.set_paragraph_default_font(main_para, 'Inter', 9)
                # Ensure zero spacing before/after as per Colab changes
                main_para.paragraph_format.space_after = Pt(0)
                main_para.paragraph_format.space_before = Pt(0)
                
                for i, ref in enumerate(self.sorted_references):
                    if i > 0:
                        spacing_run = main_para.add_run("\n\n")
                        spacing_run.font.name = 'Inter'
                        spacing_run.font.size = Pt(9)
                    
                    if ref.isNPL:
                        heading_text = f"{ref.Rank}. {ref.Title}"
                    else:
                        heading_text = f"{ref.Rank}. {ref.RawPublicationNumber}"
                    
                    heading_run = main_para.add_run(heading_text)
                    heading_run.font.name = 'Inter'
                    heading_run.font.size = Pt(9)
                    heading_run.bold = True
                    
                    # Add placeholder text only if not the last reference
                    if i < len(self.sorted_references) - 1:
                        newline_run = main_para.add_run("\n")
                        newline_run.font.name = 'Inter'
                        newline_run.font.size = Pt(9)
                
                # Final font setting to ensure no Calibri fallback
                # Add a final empty run to ensure no unstyled trailing line
                pPr = main_para._p.get_or_add_pPr()
                rPr = OxmlElement('w:rPr')
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Inter')
                rFonts.set(qn('w:hAnsi'), 'Inter')
                rPr.append(rFonts)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '18')  # 9pt in half-points
                rPr.append(sz)
                pPr.append(rPr)
                final_run = main_para.add_run("\n")
                final_run.font.name = 'Inter'
                final_run.font.size = Pt(9)
                self.set_paragraph_default_font(main_para, 'Inter', 9)

    def clear_cell_keep_formatting(self, cell):
        """Clear text but keep cell formatting and shading."""
        for p in cell.paragraphs:
            for run in p.runs:
                run.text = ""
            # Ensure zero spacing before/after to prevent unintended gaps
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
        if len(cell.paragraphs) == 0:
            para = cell.add_paragraph()
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)

    def set_paragraph_default_font(self, paragraph, font_name='Inter', font_size=9):
        """
        Set the default font for a paragraph to prevent Calibri fallback.
        
        Args:
            paragraph: The paragraph to set default font for
            font_name: Font name to use (default: 'Inter')
            font_size: Font size to use (default: 9)
        """
        # Ensure the paragraph has at least one run
        if not paragraph.runs:
            paragraph.add_run("")
        
        # Set font for all existing runs
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
        
        # Set paragraph style to ensure consistent font
        paragraph.style = None  # Clear any inherited styles

    def extract_patent_at_issue_and_claims(self):
        """
        Extract patent-at-issue information and claim numbers from Excel data.
        Uses improved patent number extraction and supports various claim input formats.
        """
        self.log("Extracting patent-at-issue and claims...")
        try:
            # Extract patent number using improved extraction method
            self.PatentAtIssue_Number = self.extract_patent_number(self.df.iloc[1, 0])
            self.short_patent_name = self.get_short_patent_name_with_suffix(self.PatentAtIssue_Number)
            self.short_patent_name_v2 = self.get_short_patent_name_v2(self.PatentAtIssue_Number)
            self.short_patent_name_lower = self.short_patent_name.replace(" Patent", " patent")

            # Look for "Required Claims" row first (new method)
            required_claims_row_idx = None
            for idx, row in self.df.iterrows():
                for col_idx, cell in enumerate(row):
                    if pd.notna(cell) and 'required claim' in str(cell).strip().lower():
                        required_claims_row_idx = idx
                        break
                if required_claims_row_idx is not None:
                    break

            self.ClaimNumbers = []
            if required_claims_row_idx is not None:
                claim_input = self.df.iloc[required_claims_row_idx, 1] if len(self.df.columns) > 1 else None
                if pd.notna(claim_input):
                    self.ClaimNumbers = self.parse_claim_numbers(claim_input, self.PatentAtIssue_Number)

            # Feb10: FALLBACK – if no "Required Claims" row, derive claims from any Expert/Reviewer Comments header
            if not self.ClaimNumbers:
                expert_comments_variations = ['Expert Comments', 'Expert/Reviewer Comments', 'Reviewer Comments']
                expert_comments_row_idx = None

                for variation in expert_comments_variations:
                    try:
                        expert_comments_row_idx = self.df[self.df[0] == variation].index[0]
                        self.log(f"Found '{variation}' row for claim numbers")
                        break
                    except (IndexError, KeyError):
                        continue

                if expert_comments_row_idx is None:
                    raise RuntimeError("Could not find Expert Comments, Expert/Reviewer Comments, or Reviewer Comments row in Excel")

                claim_row_idx = expert_comments_row_idx + 1
                for cell in self.df.iloc[claim_row_idx:, 0]:
                    if pd.isna(cell) or str(cell).strip() == "":
                        break
                    cell_str = str(cell).strip()
                    if re.match(r"^\d", cell_str):
                        claim_num = cell_str.replace(".", "")[:2].strip()
                        self.ClaimNumbers.append(claim_num)

            self.claim_word = "claim" if len(self.ClaimNumbers) == 1 else "claims"
            self.log("Patent-at-issue and claims extracted.")
        except Exception as e:
            self.log(f"Error extracting patent-at-issue and claims: {str(e)}")
            raise

    def extract_search_results(self):
        self.log("Extracting search results...")
        search_results = []
        database_rows = []
        for idx, row in self.df.iterrows():
            for col_idx, cell in enumerate(row):
                if pd.notna(cell) and str(cell).strip().lower() == 'database':
                    database_rows.append((idx, col_idx))
                    break
        if not database_rows:
            self.log("No 'Database' header found in Excel.")
            return pd.DataFrame(), 0
        header_row, header_col = database_rows[0]
        if header_col + 3 >= len(self.df.columns):
            self.log("Excel columns insufficient for search results.")
            return pd.DataFrame(), 0
        database_col = header_col
        scope_col = header_col + 1
        hits_col = header_col + 2
        query_col = header_col + 3
        data_rows = []
        current_row = header_row + 1
        s_no = 1
        while current_row < len(self.df):
            database_val = self.df.iloc[current_row, database_col] if database_col < len(self.df.columns) else None
            scope_val = self.df.iloc[current_row, scope_col] if scope_col < len(self.df.columns) else None
            query_val = self.df.iloc[current_row, query_col] if query_col < len(self.df.columns) else None
            hits_val = self.df.iloc[current_row, hits_col] if hits_col < len(self.df.columns) else None
            if pd.isna(database_val) or str(database_val).strip() == '':
                break
            data_rows.append({
                'S/No': s_no,
                'Database': str(database_val).strip(),
                'Scope': str(scope_val).strip() if pd.notna(scope_val) else '',
                'Query': str(query_val).strip() if pd.notna(query_val) else '',
                'Hits': str(hits_val).strip() if pd.notna(hits_val) else ''
            })
            s_no += 1
            current_row += 1
        search_df = pd.DataFrame(data_rows)
        total_hits = 0
        for hits in search_df['Hits']:
            try:
                hits_clean = str(hits).replace(',', '').strip()
                if hits_clean.isdigit():
                    total_hits += int(hits_clean)
            except:
                continue
        self.log("Search results extracted.")
        return search_df, total_hits

    def process_references(self):
        self.log("Processing references...")
        self.top_references = []
        self.related_references = []

        def process_reference(rank_cell, is_related=False):
            ref = self.Reference()
            ref.Rank = rank_cell
            current_row, current_col = rank_cell_coords
            try:
                if not is_related:
                    ref.RawPublicationNumber = str(self.df.iloc[current_row-9, current_col])
                    ref.PublicationNumber = self.clean_publication_number(self.df.iloc[current_row-9, current_col])
                    # Feb10: use Excel-displayed dates via row/col indices
                    ref.PriorityDate = self.format_date(current_row-8, current_col)
                    ref.FilingDate = self.format_date(current_row-7, current_col)
                    ref.PublicationDate = self.format_date(current_row-6, current_col)
                    # Align assignee offsets with notebook: Original at -4, Current at -5 for top refs
                    ref.OriginalAssignee = str(self.df.iloc[current_row-4, current_col])
                    ref.CurrentAssignee = str(self.df.iloc[current_row-5, current_col])
                    ref.Title = str(self.df.iloc[current_row-3, current_col]).strip()
                    ref.URL = str(self.df.iloc[current_row-2, current_col])
                    ref.isNPL = False if "patents.google" in ref.URL else True
                    self.top_references.append(ref)
                else:
                    ref.URL = str(self.df.iloc[current_row-2, current_col])
                    ref.isNPL = False if "patents.google" in ref.URL else True
                    ref.Title = str(self.df.iloc[current_row-3, current_col]).strip()
                    ref.OriginalAssignee = str(self.df.iloc[current_row-4, current_col])
                    ref.CurrentAssignee = str(self.df.iloc[current_row-5, current_col])
                    ref.PublicationNumber = self.clean_publication_number(self.df.iloc[current_row-9, current_col])
                    self.related_references.append(ref)
            except IndexError as e:
                self.log(f"Error processing reference at row {current_row}, col {current_col}: {str(e)}")

        try:
            rank_header_row = self.df[self.df[0] == 'Rank'].index[0]
            current_col = 1
            while current_col < self.df.shape[1]:
                current_row = rank_header_row
                while current_row < self.df.shape[0]:
                    cell_value = self.df.iloc[current_row, current_col]
                    if pd.isna(cell_value):
                        break
                    if str(cell_value).strip() in list("ABCDEFGHIJKLMNOPQRSTUVWXYZ"):
                        rank_cell_coords = (current_row, current_col)
                        process_reference(str(cell_value), is_related=False)
                    elif str(cell_value).strip() in ['RR', 'RR NPL']:
                        rank_cell_coords = (current_row, current_col)
                        process_reference(str(cell_value), is_related=True)
                    current_row += 3
                current_col += 1
            self.include_other_related_references = len(self.related_references) > 0
            self.log("References processed.")
        except Exception as e:
            self.log(f"Error processing references: {str(e)}")
            raise

    def replace_in_paragraphs_and_tables(self, doc, replacements):
        try:
            for p in doc.paragraphs:
                for key, val in replacements.items():
                    if key in p.text:
                        inline = p.runs
                        text = "".join(run.text for run in inline)
                        text = text.replace(key, val)
                        for i in range(len(inline) - 1, -1, -1):
                            p.runs[i].text = ""
                        if not inline:
                            p.add_run(text)
                        else:
                            p.runs[0].text = text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for key, val in replacements.items():
                                if key in p.text:
                                    inline = p.runs
                                    text = "".join(run.text for run in inline)
                                    text = text.replace(key, val)
                                    for i in range(len(inline) - 1, -1, -1):
                                        p.runs[i].text = ""
                                    if not inline:
                                        p.add_run(text)
                                    else:
                                        p.runs[0].text = text
        except Exception as e:
            self.log(f"Error replacing text in paragraphs/tables: {str(e)}")

    def replace_in_textboxes(self, doc, replacements):
        try:
            parts = [doc.part]
            for section in doc.sections:
                try:
                    if section.header:
                        parts.append(section.header.part)
                except Exception:
                    pass
                try:
                    if section.footer:
                        parts.append(section.footer.part)
                except Exception:
                    pass
            for part in parts:
                root = part.element
                tx_nodes = root.xpath(".//*[local-name()='txbxContent']//*[local-name()='t']")
                for t in tx_nodes:
                    if t.text:
                        for key, val in replacements.items():
                            if key in t.text:
                                t.text = t.text.replace(key, val)
        except Exception as e:
            self.log(f"Error replacing text in textboxes: {str(e)}")

    def find_paragraph_with_placeholder(self, doc, placeholder):
        try:
            for p in doc.paragraphs:
                if placeholder in p.text:
                    return p
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if placeholder in p.text:
                                return p
            self.log(f"Placeholder '{placeholder}' not found.")
            return None
        except Exception as e:
            self.log(f"Error finding paragraph with placeholder '{placeholder}': {str(e)}")
            return None

    def find_table_with_placeholder(self, doc, placeholder):
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if any(placeholder in p.text for p in cell.paragraphs):
                            return table
            self.log(f"Table with placeholder '{placeholder}' not found.")
            return None
        except Exception as e:
            self.log(f"Error finding table with placeholder '{placeholder}': {str(e)}")
            return None

    def find_row_with_placeholder(self, table, placeholder):
        try:
            for row in table.rows:
                for cell in row.cells:
                    if any(placeholder in p.text for p in cell.paragraphs):
                        return row
            self.log(f"Row with placeholder '{placeholder}' not found in table.")
            return None
        except Exception as e:
            self.log(f"Error finding row with placeholder '{placeholder}': {str(e)}")
            return None

    def clone_row_after(self, table, template_row):
        try:
            tbl = table._tbl
            tr = template_row._tr
            new_tr = deepcopy(tr)
            tbl.append(new_tr)
            return table.rows[-1]
        except Exception as e:
            self.log(f"Error cloning table row: {str(e)}")
            raise

    def clear_cell(self, cell):
        try:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.text = ""
                if len(p.runs) == 0:
                    p.add_run("")
        except Exception as e:
            self.log(f"Error clearing cell: {str(e)}")

    def set_cell_text(self, cell, text, bold=False, size=10, color_rgb=None):
        try:
            self.clear_cell(cell)
            p = cell.paragraphs[0]
            run = p.runs[0]
            run.text = text
            run.font.name = 'Inter'
            run.font.size = Pt(size)
            run.bold = bold
            if color_rgb:
                run.font.color.rgb = color_rgb
        except Exception as e:
            self.log(f"Error setting cell text: {str(e)}")

    def add_hyperlink_to_paragraph(self, doc, paragraph, url, text, size=10):
        try:
            part = doc.part
            r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            run_element = OxmlElement('w:r')
            run_props = OxmlElement('w:rPr')
            font_element = OxmlElement('w:rFonts')
            font_element.set(qn('w:ascii'), 'Inter')
            font_element.set(qn('w:hAnsi'), 'Inter')
            run_props.append(font_element)
            size_element = OxmlElement('w:sz')
            size_element.set(qn('w:val'), str(size * 2))
            run_props.append(size_element)
            size_cs_element = OxmlElement('w:szCs')
            size_cs_element.set(qn('w:val'), str(size * 2))
            run_props.append(size_cs_element)
            color_element = OxmlElement('w:color')
            color_element.set(qn('w:val'), '0000FF')
            run_props.append(color_element)
            underline_element = OxmlElement('w:u')
            underline_element.set(qn('w:val'), 'single')
            run_props.append(underline_element)
            run_element.append(run_props)
            text_element = OxmlElement('w:t')
            text_element.text = text
            run_element.append(text_element)
            hyperlink.append(run_element)
            paragraph._p.append(hyperlink)
        except Exception as e:
            self.log(f"Error adding hyperlink to paragraph: {str(e)}")

    def format_claims_list(self, claims_text_joined: str) -> str:
        claims = [c.strip() for c in claims_text_joined.split(",") if c.strip()]
        if len(claims) > 1:
            return ", ".join(claims[:-1]) + ", and " + claims[-1]
        elif claims:
            return claims[0]
        else:
            return ""

    def process_title_page(self):
      self.log("Processing title page...")
      try:
          pub_number_raw = str(self.df.iloc[1, 0]) if pd.notna(self.df.iloc[1, 0]) else ""
          assignee = str(self.df.iloc[1, 3]) if pd.notna(self.df.iloc[1, 3]) else ""
          title = str(self.df.iloc[1, 4]) if pd.notna(self.df.iloc[1, 4]) else ""

          pub_number_display = self.format_patent_display(pub_number_raw, include_prefix=True)

          current_date_str = datetime.now().strftime("%B %d, %Y")
          if self.report_type == "FTO":
              current_date_str = current_date_str.upper()

          client_name = "Unknown Client"
          if hasattr(self, 'excel_filename') and self.excel_filename:
              match = re.match(r".*?([A-Za-z0-9]+-\d+)\s*([A-Za-z\s][A-Za-z\s\.\-&]*?)\s*([A-Z]{2}\d+[A-Z]?\d*|US\d+|\d+)(?:.*)?\.xlsx", self.excel_filename, re.IGNORECASE)
              if match:
                  client_name = match.group(2).strip().upper()
              else:
                  match = re.match(r".*?([A-Za-z0-9]+-\d+)\s*(.+?)\.xlsx$", self.excel_filename, re.IGNORECASE)
                  if match:
                      client_name = match.group(2).strip().upper()
          else:
              self.log("Warning: excel_filename not set, using default client name 'Unknown Client'.")

          title_replacements = {
              "[DATE]": current_date_str.upper(),
              "[CLIENT]": client_name,
              "[PUBLICATION_NUMBER]": pub_number_display or "",
              "[ASSIGNEE]": assignee or "",
              "[PATENT_TITLE]": title or "",
              "[SHORT_PATENT_NAME]": self.short_patent_name,
              "[SHORT_PATENT_NAME_V2]": self.short_patent_name_v2,
              "[SHORT_PATENT_NAME_LOWER]": self.short_patent_name_lower,
          }

          # Use target_doc: gen_doc for update mode, doc otherwise
          target_doc = self.get_target_doc("title")
          self.replace_in_textboxes(target_doc, title_replacements)
          self.replace_in_paragraphs_and_tables(target_doc, title_replacements)
          self.log("Title page processed.")
      except Exception as e:
          self.log(f"Error processing title page: {str(e)}")
          raise

    def process_objectives(self):
        """
        Process the objectives section of the report.
        Updates the objective text with patent information and claim ranges.
        """
        self.log("Processing objectives section...")
        try:
            # Get target document: gen_doc for update mode, doc otherwise
            target_doc = self.get_target_doc("objectives")
            
            # Handle US vs non-US patent formatting
            if self.PatentAtIssue_Number.upper().startswith("US"):
                formatted_name = self.format_number_with_commas(self.PatentAtIssue_Number[2:])
                patent_prefix = "U.S. Patent No. "
            else:
                formatted_name = self.PatentAtIssue_Number
                patent_prefix = ""  # NO prefix for non-US patents
            claims_text_joined = self.format_claims_as_ranges(self.ClaimNumbers)

            obj_para = self.find_paragraph_with_placeholder(target_doc, "[OBJECTIVE_TEXT]")
            if obj_para:
                obj_para.text = ""
                if self.report_type == "Invalidity":
                    run1 = obj_para.add_run(
                        f"This report presents the mappings of the various elements of {self.claim_word} {claims_text_joined} "
                        f"of {patent_prefix}{formatted_name} "
                    )
                    run1.font.name = "Inter"
                    run1.font.size = Pt(10)

                    run2 = obj_para.add_run(f"({self.short_patent_name})")
                    run2.font.name = "Inter"
                    run2.font.size = Pt(10)
                    run2.bold = True

                    run3 = obj_para.add_run(
                        " with the most relevant disclosures found from the following references in the publication search:"
                    )
                    run3.font.name = "Inter"
                    run3.font.size = Pt(10)
                else:  # FTO
                    run1 = obj_para.add_run(
                        f"This report presents the mappings of the various elements of the client’s invention with the most relevant disclosures found from the following references in the publication search:"
                    )
                    run1.font.name = "Inter"
                    run1.font.size = Pt(10)

            ref_anchor = self.find_paragraph_with_placeholder(target_doc, "[REFERENCE_LIST]")
            if ref_anchor:
                ref_anchor.text = ""

                numbering_part = target_doc.part.numbering_part
                if numbering_part is None:
                    from docx.parts.numbering import NumberingPart
                    numbering_part = NumberingPart.new()
                    target_doc.part.relate_to(numbering_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering')

                abstractNum = OxmlElement('w:abstractNum')
                abstractNum.set(qn('w:abstractNumId'), '1')
                lvl = OxmlElement('w:lvl')
                lvl.set(qn('w:ilvl'), '0')
                numFmt = OxmlElement('w:numFmt')
                numFmt.set(qn('w:val'), 'upperLetter')
                lvl.append(numFmt)
                start = OxmlElement('w:start')
                start.set(qn('w:val'), '1')
                lvl.append(start)
                lvlText = OxmlElement('w:lvlText')
                lvlText.set(qn('w:val'), '%1.')
                lvl.append(lvlText)
                lvlJc = OxmlElement('w:lvlJc')
                lvlJc.set(qn('w:val'), 'left')
                lvl.append(lvlJc)
                pPr = OxmlElement('w:pPr')
                ind = OxmlElement('w:ind')
                ind.set(qn('w:left'), '851')
                ind.set(qn('w:hanging'), '425')
                pPr.append(ind)
                lvl.append(pPr)
                rPr = OxmlElement('w:rPr')
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Inter SemiBold')
                rFonts.set(qn('w:hAnsi'), 'Inter SemiBold')
                rPr.append(rFonts)
                b = OxmlElement('w:b')
                rPr.append(b)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '20')
                rPr.append(sz)
                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), '20')
                rPr.append(szCs)
                lvl.append(rPr)
                abstractNum.append(lvl)
                numbering_part.element.append(abstractNum)
                num = OxmlElement('w:num')
                num.set(qn('w:numId'), '1')
                abstractNumId = OxmlElement('w:abstractNumId')
                abstractNumId.set(qn('w:val'), '1')
                num.append(abstractNumId)
                numbering_part.element.append(num)

                def rank_index(rank_value):
                    r = (rank_value or "").strip().upper()
                    return ord(r) - ord('A') if len(r) == 1 and 'A' <= r <= 'Z' else 999

                self.sorted_references = sorted(self.top_references, key=lambda r: (rank_index(r.Rank), (r.PublicationNumber or "")))

                for i, ref in enumerate(self.sorted_references):
                    self.isUSPatent(ref)
                    main_para = ref_anchor.insert_paragraph_before()
                    if main_para._p.pPr is not None:
                        main_para._p.remove(main_para._p.pPr)
                    main_para.paragraph_format.left_indent = Cm(1.5)
                    main_para.paragraph_format.hanging_indent = Cm(0.75)
                    main_para.paragraph_format.space_after = Pt(0)
                    main_para.paragraph_format.space_before = Pt(18) if i == 0 else Pt(0)

                    pPr = main_para._p.get_or_add_pPr()
                    numPr = OxmlElement('w:numPr')
                    ilvl = OxmlElement('w:ilvl')
                    ilvl.set(qn('w:val'), '0')
                    numId = OxmlElement('w:numId')
                    numId.set(qn('w:val'), '1')
                    numPr.append(ilvl)
                    numPr.append(numId)
                    pPr.append(numPr)

                    if ref.isNPL:
                        pub_text = f'"{ref.Title}"'
                    elif ref.PublicationNumber.startswith("US"):
                        if '/' in (ref.PublicationName or ""):
                            pub_text = f"U.S. Pat. App. Pub. No. {ref.PublicationName}"
                        else:
                            pub_text = f"U.S. Patent No. {ref.PublicationName or ''}"
                    else:
                        pub_text = ref.PublicationNumber

                    run_pub = main_para.add_run(pub_text)
                    run_pub.font.name = 'Inter SemiBold'
                    run_pub.font.size = Pt(10)
                    run_pub.bold = True

                    if ref.isNPL:
                        # Try both Current and Original assignee fields for author/publisher
                        pub = None
                        if ref.CurrentAssignee and str(ref.CurrentAssignee).lower() not in ("nan", ""):
                            pub = ref.CurrentAssignee
                        elif ref.OriginalAssignee and str(ref.OriginalAssignee).lower() not in ("nan", ""):
                            pub = ref.OriginalAssignee

                        if pub:
                            # Determine if this is a DOI reference to use correct label
                            is_doi_ref = False
                            pub_num = ref.PublicationNumber
                            if pub_num and pub_num.lower() != "nan":
                                if str(pub_num).startswith("10.") or ("doi.org" in str(pub_num).lower()):
                                    is_doi_ref = True

                            author_label = "Author: " if is_doi_ref else "Author/Publisher: "
                            detail_para = ref_anchor.insert_paragraph_before(f"{author_label}{pub}")
                            detail_para.paragraph_format.left_indent = Cm(1.5)
                            detail_para.paragraph_format.space_after = Pt(0)
                            detail_para.paragraph_format.space_before = Pt(0)
                            run_detail = detail_para.runs[0]
                            run_detail.font.name = 'Inter'
                            run_detail.font.size = Pt(10)
                            run_detail.bold = False
                        if ref.PublicationDate:
                            detail_para = ref_anchor.insert_paragraph_before(f"Publication Date: {ref.PublicationDate}")
                            detail_para.paragraph_format.left_indent = Cm(1.5)
                            detail_para.paragraph_format.space_after = Pt(0)
                            detail_para.paragraph_format.space_before = Pt(0)
                            run_detail = detail_para.runs[0]
                            run_detail.font.name = 'Inter'
                            run_detail.font.size = Pt(10)
                            run_detail.bold = False
                        
                        # Add DOI or URL link for NPL references
                        # Determine link URL and display format using same logic as Related References
                        link_url = None
                        display_label = "Link: "
                        display_text = None
                        is_doi = False
                        pub_num = ref.PublicationNumber

                        if pub_num and pub_num.lower() != "nan":
                            # Case 1: Raw DOI (starts with "10.")
                            if str(pub_num).startswith("10."):
                                link_url = f"https://doi.org/{pub_num}"
                                display_label = "DOI: "
                                display_text = pub_num
                                is_doi = True
                            # Case 2: Already a URL containing doi.org
                            elif (str(pub_num).startswith("http://") or str(pub_num).startswith("https://")) and "doi" in str(pub_num).lower():
                                link_url = pub_num
                                # Extract DOI part after doi.org/
                                if "/10." in pub_num:
                                    doi_part = pub_num.split("/10.", 1)[1]
                                    display_label = "DOI: "
                                    display_text = "10." + doi_part
                                    is_doi = True
                                else:
                                    display_text = pub_num
                            # Case 3: Other URL
                            elif str(pub_num).startswith("http://") or str(pub_num).startswith("https://"):
                                link_url = pub_num
                                display_text = pub_num

                        # Fallback to URL field if no link found from PublicationNumber
                        if not link_url and ref.URL and ref.URL.lower() != "nan":
                            link_url = ref.URL
                            display_text = ref.URL

                        # Add the link paragraph if we have a URL
                        if link_url:
                            link_para = ref_anchor.insert_paragraph_before("")
                            link_para.paragraph_format.left_indent = Cm(1.5)
                            link_para.paragraph_format.space_after = Pt(0)
                            link_para.paragraph_format.space_before = Pt(0)

                            # Add label ("DOI: " or "Link: ") - plain text, not clickable
                            link_label_run = link_para.add_run(display_label)
                            link_label_run.font.name = 'Inter'
                            link_label_run.font.size = Pt(10)

                            # Add the hyperlink with appropriate display text
                            self.add_hyperlink_to_paragraph(target_doc, link_para, link_url, display_text if display_text else link_url)

                    elif ref.PublicationNumber.startswith("US"):
                        if ref.Title and str(ref.Title).lower() != 'nan':
                            detail_para = ref_anchor.insert_paragraph_before(f'"{ref.Title}"')
                            detail_para.paragraph_format.left_indent = Cm(1.5)
                            detail_para.paragraph_format.space_after = Pt(0)
                            detail_para.paragraph_format.space_before = Pt(0)
                            run_detail = detail_para.runs[0]
                            run_detail.font.name = 'Inter'
                            run_detail.font.size = Pt(10)
                            run_detail.bold = False
                        if ref.CurrentAssignee and str(ref.CurrentAssignee).lower() != 'nan':
                            if ref.CurrentAssignee == ref.OriginalAssignee:
                                detail_para = ref_anchor.insert_paragraph_before(f"Original & Current Assignee: {ref.CurrentAssignee}")
                            else:
                                detail_para = ref_anchor.insert_paragraph_before(f"Current Assignee: {ref.CurrentAssignee}")
                            detail_para.paragraph_format.left_indent = Cm(1.5)
                            detail_para.paragraph_format.space_after = Pt(0)
                            run_detail = detail_para.runs[0]
                            run_detail.font.name = 'Inter'
                            run_detail.font.size = Pt(10)
                            run_detail.bold = False
                        if ref.OriginalAssignee and str(ref.OriginalAssignee).lower() != 'nan':
                            if ref.CurrentAssignee != ref.OriginalAssignee:
                                detail_para = ref_anchor.insert_paragraph_before(f"Original Assignee: {ref.OriginalAssignee}")
                                detail_para.paragraph_format.left_indent = Cm(1.5)
                                detail_para.paragraph_format.space_after = Pt(0)
                                run_detail = detail_para.runs[0]
                                run_detail.font.name = 'Inter'
                                run_detail.font.size = Pt(10)
                                run_detail.bold = False
                        if ref.PriorityDate:
                            detail_para = ref_anchor.insert_paragraph_before(f"Priority Date: {ref.PriorityDate}")
                            detail_para.paragraph_format.left_indent = Cm(1.5)
                            detail_para.paragraph_format.space_after = Pt(0)
                            run_detail = detail_para.runs[0]
                            run_detail.font.name = 'Inter'
                            run_detail.font.size = Pt(10)
                            run_detail.bold = False
                        if ref.FilingDate:
                            detail_para = ref_anchor.insert_paragraph_before(f"Filing Date: {ref.FilingDate}")
                            detail_para.paragraph_format.left_indent = Cm(1.5)
                            detail_para.paragraph_format.space_after = Pt(0)
                            run_detail = detail_para.runs[0]
                            run_detail.font.name = 'Inter'
                            run_detail.font.size = Pt(10)
                            run_detail.bold = False
                        if ref.PublicationDate:
                            detail_para = ref_anchor.insert_paragraph_before(f"Publication Date: {ref.PublicationDate}")
                            detail_para.paragraph_format.left_indent = Cm(1.5)
                            detail_para.paragraph_format.space_after = Pt(0)
                            run_detail = detail_para.runs[0]
                            run_detail.font.name = 'Inter'
                            run_detail.font.size = Pt(10)
                            run_detail.bold = False

                    else:
                        if ref.Title and str(ref.Title).lower() != 'nan':
                            detail_para = ref_anchor.insert_paragraph_before(f'"{ref.Title}"')
                            detail_para.paragraph_format.left_indent = Cm(1.5)
                            detail_para.paragraph_format.space_after = Pt(0)
                            run_detail = detail_para.runs[0]
                            run_detail.font.name = 'Inter'
                            run_detail.font.size = Pt(10)
                            run_detail.bold = False
                        if ref.CurrentAssignee and str(ref.CurrentAssignee).lower() != 'nan':
                            if ref.CurrentAssignee == ref.OriginalAssignee:
                                detail_para = ref_anchor.insert_paragraph_before(f"Original & Current Assignee: {ref.CurrentAssignee}")
                            else:
                                detail_para = ref_anchor.insert_paragraph_before(f"Current Assignee: {ref.CurrentAssignee}")
                            detail_para.paragraph_format.left_indent = Cm(1.5)
                            detail_para.paragraph_format.space_after = Pt(0)
                            run_detail = detail_para.runs[0]
                            run_detail.font.name = 'Inter'
                            run_detail.font.size = Pt(10)
                            run_detail.bold = False
                        if ref.OriginalAssignee and str(ref.OriginalAssignee).lower() != 'nan':
                            if ref.CurrentAssignee != ref.OriginalAssignee:
                                detail_para = ref_anchor.insert_paragraph_before(f"Original Assignee: {ref.OriginalAssignee}")
                                detail_para.paragraph_format.left_indent = Cm(1.5)
                                detail_para.paragraph_format.space_after = Pt(0)
                                run_detail = detail_para.runs[0]
                                run_detail.font.name = 'Inter'
                                run_detail.font.size = Pt(10)
                                run_detail.bold = False
                        if ref.PriorityDate:
                            detail_para = ref_anchor.insert_paragraph_before(f"Priority Date: {ref.PriorityDate}")
                            detail_para.paragraph_format.left_indent = Cm(1.5)
                            detail_para.paragraph_format.space_after = Pt(0)
                            run_detail = detail_para.runs[0]
                            run_detail.font.name = 'Inter'
                            run_detail.font.size = Pt(10)
                            run_detail.bold = False
                        if ref.FilingDate:
                            detail_para = ref_anchor.insert_paragraph_before(f"Filing Date: {ref.FilingDate}")
                            detail_para.paragraph_format.left_indent = Cm(1.5)
                            detail_para.paragraph_format.space_after = Pt(0)
                            run_detail = detail_para.runs[0]
                            run_detail.font.name = 'Inter'
                            run_detail.font.size = Pt(10)
                            run_detail.bold = False
                        if ref.PublicationDate:
                            detail_para = ref_anchor.insert_paragraph_before(f"Publication Date: {ref.PublicationDate}")
                            detail_para.paragraph_format.left_indent = Cm(1.5)
                            detail_para.paragraph_format.space_after = Pt(0)
                            run_detail = detail_para.runs[0]
                            run_detail.font.name = 'Inter'
                            run_detail.font.size = Pt(10)
                            run_detail.bold = False



                    if i < len(self.sorted_references) - 1:
                        spacer = ref_anchor.insert_paragraph_before("")
                        spacer.paragraph_format.left_indent = Cm(1.5)
                        spacer.paragraph_format.space_after = Pt(0)
                        spacer.paragraph_format.space_before = Pt(0)

                if ref_anchor:
                    parent = ref_anchor._element.getparent()
                    anchor_index = list(parent).index(ref_anchor._element)
                    parent.remove(ref_anchor._element)
                    while anchor_index < len(parent):
                        elem = parent[anchor_index]
                        if elem.tag.endswith('p'):
                            text_content = ''.join(elem.itertext()).strip()
                            if not text_content:
                                parent.remove(elem)
                            else:
                                break
                        else:
                            break
            else:
                self.log("Warning: [REFERENCE_LIST] placeholder not found.")
            self.log("Objectives section processed.")
        except Exception as e:
            self.log(f"Error processing objectives: {str(e)}")
            raise

    def add_page_break_before_paragraph(self, doc, target_paragraph):
        try:
            page_break = OxmlElement('w:br')
            page_break.set(qn('w:type'), 'page')
            run_element = OxmlElement('w:r')
            run_element.append(page_break)
            target_paragraph._p.insert(0, run_element)
        except Exception as e:
            self.log(f"Error adding page break: {str(e)}")

    def process_other_related_references(self):
        self.log("Processing other related references...")
        try:
            # Get target document: gen_doc for update mode, doc otherwise
            target_doc = self.get_target_doc("references")
            
            if self.include_other_related_references:
                table_rr = self.find_table_with_placeholder(target_doc, "**[REF_INDEX]**") or \
                          self.find_table_with_placeholder(target_doc, "[REF_ENTRY]") or \
                          self.find_table_with_placeholder(target_doc, "[REF_OWNER]")
                if table_rr:
                    # Check if an ORR heading exists anywhere; if so, skip generating another heading
                    other_refs_heading = None
                    for paragraph in target_doc.paragraphs:
                        if "other related references" in paragraph.text.lower():
                            other_refs_heading = paragraph
                            break
                    if other_refs_heading:
                        # Page break before existing ORR heading
                        self.add_page_break_before_paragraph(target_doc, other_refs_heading)
                    else:
                        # Insert ORR header immediately before the table if missing
                        try:
                            from docx.oxml import OxmlElement
                            from docx.oxml.ns import qn
                            new_p = OxmlElement('w:p')
                            r = OxmlElement('w:r')
                            rPr = OxmlElement('w:rPr')
                            rFonts = OxmlElement('w:rFonts')
                            rFonts.set(qn('w:ascii'), 'Inter')
                            rFonts.set(qn('w:hAnsi'), 'Inter')
                            rPr.append(rFonts)
                            b = OxmlElement('w:b')
                            rPr.append(b)
                            sz = OxmlElement('w:sz')
                            sz.set(qn('w:val'), '20')  # 10pt
                            rPr.append(sz)
                            szCs = OxmlElement('w:szCs')
                            szCs.set(qn('w:val'), '20')
                            rPr.append(szCs)
                            color = OxmlElement('w:color')
                            color.set(qn('w:val'), '000000')
                            rPr.append(color)
                            r.append(rPr)
                            t = OxmlElement('w:t')
                            t.text = 'OTHER RELATED REFERENCES FOUND'
                            r.append(t)
                            pPr = OxmlElement('w:pPr')
                            spacing = OxmlElement('w:spacing')
                            spacing.set(qn('w:after'), '160')  # 8pt
                            spacing.set(qn('w:line'), '216')   # 1.08 * 200
                            spacing.set(qn('w:lineRule'), 'auto')
                            pPr.append(spacing)
                            new_p.append(pPr)
                            new_p.append(r)
                            # Insert before table
                            table_rr._tbl.addprevious(new_p)
                        except Exception as e:
                            self.log(f"Warning: Could not insert ORR header before table: {str(e)}")
                    row_template = self.find_row_with_placeholder(table_rr, "**[REF_INDEX]**") or table_rr.rows[-1]
                    granted_us_patents, us_applications, foreign_patents, npl_references = [], [], [], []

                    for ref in self.related_references:
                        # Normalize computed publication name fields
                        self.isUSPatent(ref)
                        if ref.isNPL:
                            npl_references.append(ref)
                        elif isinstance(ref.PublicationNumber, str) and ref.PublicationNumber.startswith("US"):
                            # US references: split into granted vs applications using '/'
                            if ref.PublicationName and '/' in str(ref.PublicationName):
                                us_applications.append(ref)
                            else:
                                granted_us_patents.append(ref)
                        else:
                            # Non-US patents
                            foreign_patents.append(ref)

                    # Sort within categories for stable deterministic ordering
                    def sort_key_pat(ref_obj):
                        return (ref_obj.PublicationNumber or "").upper()
                    def sort_key_npl(ref_obj):
                        return (ref_obj.Title or "").upper()

                    granted_us_patents.sort(key=sort_key_pat)
                    us_applications.sort(key=sort_key_pat)
                    foreign_patents.sort(key=sort_key_pat)
                    npl_references.sort(key=sort_key_npl)

                    sorted_related_refs = granted_us_patents + us_applications + foreign_patents + npl_references

                    def render_ref_into_row(row_cells, idx, ref_obj):
                        from docx.enum.table import WD_ALIGN_VERTICAL
                        for cell in row_cells:
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        self.set_cell_text(row_cells[0], str(idx), size=9, bold=True)
                        if ref_obj.isNPL:
                            self.clear_cell(row_cells[1])
                            p1 = row_cells[1].paragraphs[0]
                            r = p1.add_run(f'"{ref_obj.Title}"')
                            r.font.name = 'Inter'
                            r.font.size = Pt(9)
                            
                            # Determine link URL: prioritize DOI from PublicationNumber, fallback to URL field
                            link_url = None
                            pub_num = ref_obj.PublicationNumber

                            if pub_num and pub_num.lower() != "nan":
                                # Case 1: Raw DOI (starts with "10.")
                                if str(pub_num).startswith("10."):
                                    link_url = f"https://doi.org/{pub_num}"
                                # Case 2: Already a URL
                                elif str(pub_num).startswith("http://") or str(pub_num).startswith("https://"):
                                    link_url = pub_num

                            # Fallback to URL field if no DOI found
                            if not link_url and ref_obj.URL and ref_obj.URL.lower() != "nan":
                                link_url = ref_obj.URL

                            # Add hyperlink if we have a valid URL
                            if link_url:
                                p2 = row_cells[1].add_paragraph()
                                p2.add_run("[")
                                self.add_hyperlink_to_paragraph(target_doc, p2, link_url, "Link", size=9)
                                p2.add_run("]")
                                self.apply_font_style(p2, size=9)
                        else:
                            if ref_obj.PublicationNumber and ref_obj.PublicationNumber.startswith("US"):
                                if ref_obj.PublicationName and '/' in str(ref_obj.PublicationName):
                                    pub_text = f"U.S. Pat. App. Pub. No. {ref_obj.PublicationName}"
                                else:
                                    pub_text = f"U.S. Patent No. {ref_obj.PublicationName or ''}"
                            else:
                                pub_text = ref_obj.PublicationNumber or "Unknown Patent Number"
                            self.set_cell_text(row_cells[1], pub_text, size=9)
                        auth = ref_obj.CurrentAssignee if ref_obj.CurrentAssignee else ref_obj.OriginalAssignee
                        if auth == "nan" or auth == "":
                            auth = ""
                        self.set_cell_text(row_cells[2], auth or "", size=9)

                    if sorted_related_refs:
                        render_ref_into_row(row_template.cells, 1, sorted_related_refs[0])
                        for idx, ref in enumerate(sorted_related_refs[1:], start=2):
                            new_row = self.clone_row_after(table_rr, row_template)
                            render_ref_into_row(new_row.cells, idx, ref)
                    else:
                        for cell in row_template.cells:
                            self.clear_cell(cell)
                else:
                    self.log("Warning: Other related references table not found.")
            else:
                # No related references: keep section and show message (match notebook behavior)
                other_refs_heading = None
                for paragraph in target_doc.paragraphs:
                    if "other related references" in paragraph.text.lower():
                        other_refs_heading = paragraph
                        break

                if other_refs_heading:
                    # Add page break before the heading to start section on a new page
                    self.add_page_break_before_paragraph(target_doc, other_refs_heading)

                # Find and clear the table with placeholders, then insert message
                table_rr = self.find_table_with_placeholder(target_doc, "**[REF_INDEX]**") or \
                           self.find_table_with_placeholder(target_doc, "[REF_ENTRY]") or \
                           self.find_table_with_placeholder(target_doc, "[REF_OWNER]")
                if table_rr:
                    row_template = self.find_row_with_placeholder(table_rr, "**[REF_INDEX]**") or table_rr.rows[-1]

                    # Clear all cells in template row
                    for cell in row_template.cells:
                        self.clear_cell(cell)

                    # Add message in the second cell (where references would go)
                    # Keep first and third cells empty
                    self.set_cell_text(row_template.cells[0], "", size=9)
                    self.set_cell_text(row_template.cells[1], "No related references found in this search.", size=9)
                    self.set_cell_text(row_template.cells[2], "", size=9)
            self.log("Other related references processed.")
        except Exception as e:
            self.log(f"Error processing other related references: {str(e)}")
            raise

    def process_patent_at_issue(self):
        self.log("Processing patent-at-issue section...")
        try:
            # Get target document: gen_doc for update mode, doc otherwise
            target_doc = self.get_target_doc("patent")
            
            # Handle US vs non-US patent formatting
            if self.PatentAtIssue_Number.upper().startswith("US"):
                patent_number_display = self.format_number_with_commas(self.PatentAtIssue_Number[2:])
                patent_display_text = f"U.S. Patent No. {patent_number_display}"
            else:
                # Non-US patent - JUST the publication number, NO prefix
                patent_display_text = self.PatentAtIssue_Number
            
            assignee_display = str(self.df.iloc[1, 3]) if pd.notna(self.df.iloc[1, 3]) else ""
            # Feb10: use Excel-displayed date for priority using row/col indices
            priority_display = self.format_date(1, 1)

            abstract_text = self.fetch_abstract(self.PatentAtIssue_Number)
            abstract_text = abstract_text.lstrip()

            patent_replacements = {
                "[PATENT_AT_ISSUE_NUMBER]": patent_display_text,
                "[PATENT_AT_ISSUE_ASSIGNEE]": f"Current Assignee: {assignee_display}",
                "[PATENT_AT_ISSUE_PRIORITY_DATE]": f"Earliest Priority Date: {priority_display}",
                "[PATENT_AT_ISSUE_ABSTRACT]": abstract_text,
            }
            self.replace_in_paragraphs_and_tables(target_doc, patent_replacements)
            self.log("Patent-at-issue section processed.")
        except Exception as e:
            self.log(f"Error processing patent-at-issue section: {str(e)}")
            raise

    def process_criteria(self):
        """
        Process the criteria section of the report.
        Handles both Invalidity and FTO report types with appropriate claim formatting.
        Supports merging with edited reports in update mode.
        """
        self.log("Processing criteria section...")
        self.log(f"DEBUG: update_mode = {self.update_mode}")
        self.log(f"DEBUG: edited_doc is None = {self.edited_doc is None}")
        try:
            def merge_claim_fragments(fragments):
                """
                Merge fragments so that lone "claim X" lines are combined with the
                surrounding text, matching the notebook's behavior.
                """
                merged = []
                i = 0
                while i < len(fragments):
                    fragment = fragments[i].replace('\n', ' ').strip()
                    if i + 1 < len(fragments):
                        next_frag = fragments[i + 1].strip()
                        if re.match(r'^(of\s+)?claim\s+\d+[,.]?$', next_frag, re.IGNORECASE):
                            combined = fragment + " " + next_frag
                            if i + 2 < len(fragments):
                                following_frag = fragments[i + 2].strip()
                                if following_frag and following_frag[0] in ',.;:':
                                    combined += following_frag
                                else:
                                    combined += " " + following_frag
                                i += 3
                            else:
                                i += 2
                            merged.append(combined)
                            continue
                    merged.append(fragment)
                    i += 1
                return merged

            # Extract and preserve criteria section from edited document
            if self.update_mode:
                self.log("DEBUG: In update mode, extracting criteria section...")
                preserved_criteria_elements = extract_criteria_section(self.edited_doc)
            else:
                self.log("DEBUG: Not in update mode, skipping criteria extraction")
                preserved_criteria_elements = []

            # Filter out preserved criteria heading so we only merge content starting from the claims text
            def _filter_preserved_criteria_elements(elements):
                filtered = []
                for el in elements:
                    try:
                        if el.tag == qn('w:p'):
                            # Get text directly from element without creating Paragraph object
                            text_elem = el.xpath('.//w:t')
                            text = ''.join([t.text for t in text_elem if t.text is not None]) if text_elem else ""
                            if "criteria for the publication search" in text.lower().strip():
                                # Skip preserved section header
                                continue
                    except Exception:
                        pass
                    filtered.append(el)
                return filtered

            filtered_preserved_criteria_elements = _filter_preserved_criteria_elements(preserved_criteria_elements)

            criteria_num_elements = len(preserved_criteria_elements)
            if len(preserved_criteria_elements) > 0:
                self.log(f"✓ Extracted {criteria_num_elements} elements from Criteria section.")
            else:
                self.log("⚠ No Criteria section found in edited report. Generating fresh Criteria content.")

            # Always replace placeholders first, regardless of mode
            criteria_intro = f"{self.claim_word.capitalize()} {self.format_claims_as_ranges(self.ClaimNumbers)} of the {self.short_patent_name}"
            
            # Find criteria header and insert empty paragraph after it
            # criteria_header = None
            # for p in self.doc.paragraphs:
            #     if "criteria for the publication search" in p.text.lower():
            #         criteria_header = p
            #         # Insert an empty paragraph after the criteria header
            #         empty_para = self.insert_paragraph_after(p, "")
            #         empty_para.paragraph_format.space_after = Pt(0)
            #         empty_para.paragraph_format.space_before = Pt(0)
            #         break
            
            self.replace_in_paragraphs_and_tables(self.doc, {
                "[CRITERIA_TEXT]": criteria_intro,
                "[CRITERIA_CLAIM/S]": ""
            })
            criteria_anchor = self.find_paragraph_with_placeholder(self.doc, "[CRITERIA_TEXT]") or self.find_paragraph_with_placeholder(self.doc, criteria_intro)
            if criteria_anchor:
                for run in criteria_anchor.runs:
                    run.bold = True
                    run.font.name = 'Inter'
                    run.font.size = Pt(10)

            # Check if we have preserved criteria elements from edited document
            if self.update_mode and preserved_criteria_elements:
                self.log("✓ Using preserved criteria elements")
                self.log("DEBUG: Processing preserved criteria elements...")
                # Instead of removing sections, let's just replace the criteria content in place
                # This is safer and won't accidentally remove mapping placeholders

                # Clear the original criteria content by removing the criteria anchor and everything after it
                # until we hit the mappings section
                if criteria_anchor:
                    # First, find a new anchor point before the criteria section
                    # Look for the paragraph that contains "CRITERIA FOR THE PUBLICATION SEARCH" heading
                    new_anchor = None
                    for p in self.doc.paragraphs:
                        if "criteria for the publication search" in p.text.lower():
                            new_anchor = p
                            break

                    if new_anchor:
                        # Remove the criteria anchor paragraph (which contains the generated claims text)
                        criteria_anchor._element.getparent().remove(criteria_anchor._element)

                        # Remove everything after the criteria heading until we hit the mappings section
                        current_para = new_anchor._element.getnext()
                        while current_para is not None:
                            if current_para.tag == qn('w:p'):
                                p = Paragraph(current_para, self.doc)
                                text_lower = p.text.lower().strip()
                                # Stop when we hit the mappings section
                                if "mappings based on selected references" in text_lower:
                                    break
                                # Don't remove paragraphs that contain mapping placeholders
                                if any(placeholder in p.text for placeholder in ['[CLAIM_HEADER', '[MAPPINGS_PARAGRAPH]', '[REF_INDEX]', '[REF_ENTRY]', '[REF_OWNER]']):
                                    current_para = current_para.getnext()
                                    continue
                            # Remove this element
                            next_para = current_para.getnext()
                            current_para.getparent().remove(current_para)
                            current_para = next_para

                        # Insert the preserved criteria elements after the criteria heading (skip preserved header)
                        current_anchor = new_anchor
                        
                        # Insert a blank line after the criteria header in update mode
                        empty_para_after_header = self.insert_paragraph_after(new_anchor, "")
                        empty_para_after_header.paragraph_format.space_after = Pt(0)
                        empty_para_after_header.paragraph_format.space_before = Pt(0)
                        current_anchor = empty_para_after_header
                        
                        for el in filtered_preserved_criteria_elements:
                            new_el = insert_element_after(current_anchor, el)
                            if new_el:
                                current_anchor = new_el
                        # Update last_inserted_para to point to the last inserted element from criteria section
                        last_inserted_para = current_anchor
                        self.last_inserted_para = current_anchor
                    else:
                        # Fallback: use the original criteria_anchor
                        current_anchor = criteria_anchor
                        for el in filtered_preserved_criteria_elements:
                            new_el = insert_element_after(current_anchor, el)
                            if new_el:
                                current_anchor = new_el
                        last_inserted_para = current_anchor
                        self.last_inserted_para = current_anchor
                else:
                    # Fallback if no criteria_anchor found
                    current_anchor = None
                    for el in filtered_preserved_criteria_elements:
                        if current_anchor is None:
                            # Insert at the end of the document
                            current_anchor = self.doc.add_paragraph()
                        new_el = insert_element_after(current_anchor, el)
                        if new_el:
                            current_anchor = new_el
                    last_inserted_para = current_anchor
                    self.last_inserted_para = current_anchor
            if not self.update_mode or not preserved_criteria_elements:
                # Insert all claims in order, immediately after the intro
                if criteria_anchor:
                    last_inserted_para = criteria_anchor
                    self.last_inserted_para = criteria_anchor
                    for ClaimNumber in self.ClaimNumbers:
                        web_scraped_claim = self.get_claim_from_google_patents(self.PatentAtIssue_Number, ClaimNumber)
                        if web_scraped_claim:
                            lines = web_scraped_claim.split('\n')
                            claim_parts = [line.strip() for line in lines if line.strip()]
                            if len(claim_parts) == 1:
                                text = claim_parts[0]
                                patterns = [r'(?=\bwherein\b)', r'(?=\bcomprising\b)', r'(?=\bfurther comprising\b)',
                                            r'(?=\bcharacterized by\b)', r'(?=\band\b)', r'(?=\bor\b)']
                                for pattern in patterns:
                                    parts = re.split(pattern, text, flags=re.IGNORECASE)
                                    if len(parts) > 1:
                                        claim_parts = [part.strip() for part in parts if part.strip()]
                                        break
                            fragments_to_use = claim_parts if claim_parts else [web_scraped_claim]
                        else:
                            fragments_to_use = self.extract_claim_fragments_from_excel(self.df)

                        # Merge fragments around "claim X" references
                        merged_fragments = merge_claim_fragments(fragments_to_use)
                        for idx, fragment in enumerate(merged_fragments):
                            new_para = self.insert_paragraph_after(last_inserted_para, fragment)
                            if idx != 0:
                                new_para.paragraph_format.left_indent = Inches(0.23)
                            new_para.paragraph_format.keep_together = True
                            self.apply_font_style(new_para)
                            last_inserted_para = new_para
                            self.last_inserted_para = new_para
            else:  # FTO
                criteria_text = self.extract_claim_fragments_from_excel(self.df)
                criteria_anchor = self.find_paragraph_with_placeholder(self.doc, "[CRITERIA_CLAIM/S]")
                if criteria_anchor:
                    criteria_anchor.text = criteria_anchor.text.replace("[CRITERIA_CLAIM/S]", "").strip()
                    if criteria_text:
                        criteria_anchor.text = criteria_text[0]
                        self.apply_font_style(criteria_anchor, size=10, bold=False)
                        last_inserted_para = criteria_anchor
                        for text in criteria_text[1:]:
                            new_para = self.insert_paragraph_after(last_inserted_para, text)
                            # FTO requires 0.5 inch indent to match Colab
                            new_para.paragraph_format.left_indent = Inches(0.5)
                            self.apply_font_style(new_para, size=10, bold=False)
                            last_inserted_para = new_para
                            self.last_inserted_para = new_para
                    else:
                        criteria_anchor.text = ""
                        self.log("Warning: No criteria text found in Excel for FTO report.")
                else:
                    self.log("Warning: [CRITERIA_CLAIM/S] placeholder not found for FTO report.")
            self.log("Criteria section processed.")
        except Exception as e:
            self.log(f"Error processing criteria section: {str(e)}")
            raise

    def find_mapping_tables(self, doc):
        try:
            tables = []
            for t in doc.tables:
                header_no = None
                for cell in t.rows[0].cells:
                    for p in cell.paragraphs:
                        m = re.search(r"\[CLAIM_HEADER(\d+)\]", p.text)
                        if m:
                            header_no = int(m.group(1))
                            break
                    if header_no is not None:
                        break
                if header_no is not None:
                    tables.append((header_no, t))
            tables.sort(key=lambda x: x[0])
            return tables
        except Exception as e:
            self.log(f"Error finding mapping tables: {str(e)}")
            return []

    def update_headers(self, table, claim_number_header):
        """
        Update table headers with the correct claim number.
        Handles both placeholder text and already processed headers.
        """
        try:
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    # Check if this is a claim header (left column)
                    if "[CLAIM_HEADER" in p.text:
                        # Replace placeholder with actual claim number
                        p.text = re.sub(r"\[CLAIM_HEADER\d+\]",
                                        f"{self.short_patent_name}'s Claim {claim_number_header} Elements", p.text)
                        for r in p.runs:
                            r.bold = True
                            r.font.name = 'Inter'
                            r.font.size = Pt(10)
                    elif "Claim" in p.text and "Elements" in p.text:
                        # Update existing claim header with new claim number
                        p.text = f"{self.short_patent_name}'s Claim {claim_number_header} Elements"
                        for r in p.runs:
                            r.bold = True
                            r.font.name = 'Inter'
                            r.font.size = Pt(10)
                    
                    # Check if this is a reference header (right column)
                    if "[REFERENCE_HEADER" in p.text:
                        p.text = re.sub(r"\[REFERENCE_HEADER\d+\]",
                                        "Related Disclosures from the Selected References", p.text)
                        for r in p.runs:
                            r.bold = True
                            r.font.name = 'Inter'
                            r.font.size = Pt(10)
                    elif "Related Disclosures" in p.text:
                        # Update existing reference header
                        p.text = "Related Disclosures from the Selected References"
                        for r in p.runs:
                            r.bold = True
                            r.font.name = 'Inter'
                            r.font.size = Pt(10)
        except Exception as e:
            self.log(f"Error updating table headers: {str(e)}")

    def find_placeholder_row_obj(self, table):
        try:
            for row in table.rows:
                if len(row.cells) >= 2:
                    left_has = any("[CLAIM_ELEMENT]" in p.text for p in row.cells[0].paragraphs)
                    right_has = any("[REFERENCE_DISCLOSURE/S]" in p.text for p in row.cells[1].paragraphs)
                    if left_has or right_has:
                        return row
            self.log("Warning: Placeholder row with [CLAIM_ELEMENT] or [REFERENCE_DISCLOSURE/S] not found.")
            return None
        except Exception as e:
            self.log(f"Error finding placeholder row: {str(e)}")
            return None

    def clear_cell_strict(self, cell):
        try:
            for p in list(cell.paragraphs):
                p._element.getparent().remove(p._element)
            p = cell.add_paragraph()
            r = p.add_run("")
            r.font.name = 'Inter'
            r.font.size = Pt(10)
        except Exception as e:
            self.log(f"Error clearing cell strictly: {str(e)}")

    def get_claim_fragments_for_claim(self, claim_number):
        try:
            # Feb10: support multiple possible Expert/Reviewer Comments header labels
            expert_comments_variations = ['Expert Comments', 'Expert/Reviewer Comments', 'Reviewer Comments']
            claim_start_idx = None
            for variation in expert_comments_variations:
                try:
                    claim_start_idx = self.df[self.df[0] == variation].index[0] + 1
                    break
                except (IndexError, KeyError):
                    continue

            claim_fragments = []
            fragment_rows = []
            if claim_start_idx is not None:
                current_idx = claim_start_idx
                found_claim = False

                while current_idx < len(self.df):
                    cell_value = self.df.iloc[current_idx, 0]
                    if pd.isna(cell_value) or str(cell_value).strip() == "":
                        break
                    cell_str = str(cell_value).strip()
                    if re.match(rf"^{claim_number}\.", cell_str):
                        found_claim = True
                        claim_fragments.append(cell_str)
                        fragment_rows.append(current_idx)
                    elif found_claim and re.match(r"^\d+\.", cell_str):
                        break
                    elif found_claim:
                        claim_fragments.append(cell_str)
                        fragment_rows.append(current_idx)
                    current_idx += 1

            if not claim_fragments:
                web_scraped = self.get_claim_from_google_patents(self.PatentAtIssue_Number, claim_number)
                if web_scraped:
                    claim_fragments = [line.strip() for line in web_scraped.split('\n') if line.strip()]
                    fragment_rows = [-1] * len(claim_fragments)

            if not claim_fragments or not claim_fragments[0].startswith(f"{claim_number}."):
                claim_fragments.insert(0, f"{claim_number}. [Claim text not found]")
                if len(fragment_rows) == 0:
                    fragment_rows.insert(0, -1)

            return claim_fragments, fragment_rows
        except Exception as e:
            self.log(f"Error getting claim fragments for claim {claim_number}: {str(e)}")
            return [f"{claim_number}. [Error retrieving claim text]"], [-1]

    def get_mapped_references_for_fragment(self, claim_number, target_row_idx):
        try:
            mapped = []
            for ref in self.sorted_references:
                if ref.isNPL:
                    label = f'{ref.Rank}. "{ref.Title}"' if ref.Title else f"{ref.Rank}. [No Title]"
                else:
                    label = f"{ref.Rank}. {ref.RawPublicationNumber}"
                mapped.append((ref.Rank or "", label))
            return [label for _, label in mapped]
        except Exception as e:
            self.log(f"Error getting mapped references: {str(e)}")
            return []

    def process_mappings(self):
        """
        Process the mappings section of the report.
        Creates mapping tables for each claim with color-coded elements.
        Supports merging with edited reports in update mode.
        """
        self.log("Processing mappings section...")
        self.log(f"DEBUG: update_mode = {self.update_mode}")
        self.log(f"DEBUG: edited_doc is None = {self.edited_doc is None}")
        try:
            # Diagnostics: capture section indices before any changes
            try:
                def _idx_of(text):
                    from docx.text.paragraph import Paragraph as _Paragraph
                    body = self.doc.element.body
                    for i, el in enumerate(list(body)):
                        if el.tag.endswith('p'):
                            p = _Paragraph(el, self.doc)
                            if text.lower() in (p.text or '').lower():
                                return i
                    return None
                criteria_idx_pre = _idx_of('criteria for the publication search')
                mappings_idx_pre = _idx_of('mappings based on selected references')
                about_idx_pre = _idx_of('about us')
                disclaimer_idx_pre = _idx_of('disclaimer')
                self.log(f"DEBUG: [pre-mappings] indices → criteria={criteria_idx_pre}, mappings={mappings_idx_pre}, about={about_idx_pre}, disclaimer={disclaimer_idx_pre}")
            except Exception:
                pass
            def rank_index(rank_value):
                r = (rank_value or "").strip().upper()
                return ord(r) - ord('A') if len(r) == 1 and 'A' <= r <= 'Z' else 999
            self.sorted_references = sorted(self.top_references, key=lambda r: (rank_index(r.Rank), (r.PublicationNumber or "")))
            color_cycle = [RGBColor(0x00, 0x70, 0xC0), RGBColor(0xC0, 0x00, 0x00)]
            
            # Initialize global color index for consistent coloring across claims
            self.global_color_index = 0

            # Extract and preserve mapping section from edited document
            if self.update_mode:
                self.log("DEBUG: In update mode, extracting mapping section...")
                preserved_mapping_elements = extract_mapping_section(self.edited_doc)
            else:
                self.log("DEBUG: Not in update mode, skipping mapping extraction")
                preserved_mapping_elements = []

            num_elements = len(preserved_mapping_elements)
            num_tables = sum(1 for el in preserved_mapping_elements if el.tag == qn('w:tbl'))
            if len(preserved_mapping_elements) > 0:
                self.log(f"✓ Extracted {num_elements} elements from Mapping section, including {num_tables} tables.")
            else:
                self.log("⚠ No Mapping section found in edited report. Generating fresh Mapping tables.")

            # Ensure last_inserted_para is defined - it should be the last paragraph from criteria section
            if not hasattr(self, 'last_inserted_para') or self.last_inserted_para is None:
                # This should not happen now since we set last_inserted_para in both criteria paths
                # But keep as fallback for safety
                self.log("⚠ Warning: last_inserted_para not defined, using fallback logic")
                # Find the last paragraph that was processed in the criteria section
                # Look for the criteria anchor or the last paragraph before mappings
                self.last_inserted_para = None
                if not self.update_mode:
                    # Look for criteria section by searching for common patterns
                    for p in self.doc.paragraphs:
                        text_lower = p.text.lower().strip()
                        if any(pattern in text_lower for pattern in ["criteria for the publication search", "criteria for publication search", "criteria"]):
                            self.last_inserted_para = p
                            break
                else:
                    # In update mode, look for the last paragraph in the document
                    for p in reversed(self.doc.paragraphs):
                        text_lower = p.text.lower().strip()
                        # Skip if it's a heading or placeholder
                        if not any(keyword in text_lower for keyword in ["mappings based", "criteria for the", "disclaimer", "appendix"]):
                            self.last_inserted_para = p
                            break

                # If not found, use the last paragraph in the document
                if self.last_inserted_para is None and self.doc.paragraphs:
                    self.last_inserted_para = self.doc.paragraphs[-1]

            # If we have preserved elements, replace the generated Mapping section with them
            if self.update_mode and preserved_mapping_elements:
                self.log("✓ Using preserved mapping elements")
                self.log("DEBUG: Removing existing mapping section...")
                # Remove from "Mappings Based" to just before "Disclaimer" (don't remove the disclaimer itself)
                # Find the disclaimer paragraph first to know where to stop
                disclaimer_para = None
                for p in self.doc.paragraphs:
                    if "disclaimer" in p.text.lower():
                        disclaimer_para = p
                        break
                
                if disclaimer_para:
                    # Remove from "Mappings Based" up to but not including the disclaimer paragraph
                    body = self.doc.element.body
                    elems_to_remove = []
                    remove_started = False
                    for child in list(body):
                        if child.tag == qn('w:p'):
                            p = Paragraph(child, self.doc)
                            text_lower = p.text.lower().strip()
                            if "mappings based on selected references" in text_lower:
                                remove_started = True
                                elems_to_remove.append(child)
                                continue
                        if remove_started:
                            if child == disclaimer_para._element:
                                break  # Stop before removing the disclaimer
                            elems_to_remove.append(child)
                    for el in elems_to_remove:
                        body.remove(el)
                    self.log("DEBUG: Mapping section removal completed")
                else:
                    # Fallback to original method
                    remove_section(self.doc, "mappings based on selected references", "disclaimer")
                    self.log("DEBUG: Mapping section removal completed (fallback)")
                current_anchor = self.last_inserted_para
                if isinstance(current_anchor, Table):
                    anchor_text = "Table"
                elif hasattr(current_anchor, 'text'):
                    anchor_text = current_anchor.text[:100]
                else:
                    anchor_text = str(type(current_anchor))
                self.log(f"DEBUG: Inserting mapping elements after: {anchor_text}")
                self.log(f"DEBUG: Number of mapping elements to insert: {len(preserved_mapping_elements)}")
                
                # Find mappings header and insert blank line after it in update mode
                mappings_header = None
                for p in self.doc.paragraphs:
                    if "mappings based on selected references" in p.text.lower():
                        mappings_header = p
                        # Insert an empty paragraph after the mappings header
                        empty_para_after_header = self.insert_paragraph_after(p, "")
                        empty_para_after_header.paragraph_format.space_after = Pt(0)
                        empty_para_after_header.paragraph_format.space_before = Pt(0)
                        break
                
                # Add page break before inserting preserved mapping elements
                # page_break_para = self.doc.add_paragraph()
                # run = page_break_para.add_run()
                # run.add_break(WD_BREAK.PAGE)
                # current_anchor._element.addnext(page_break_para._p)
                # current_anchor = page_break_para
                for idx, el in enumerate(preserved_mapping_elements):
                    # Check if this element is a table that follows another table (new claim, needs page break)
                    if el.tag == qn('w:tbl') and idx > 0 and idx < len(preserved_mapping_elements):
                        prev_element = preserved_mapping_elements[idx - 1]
                        if prev_element.tag == qn('w:tbl'):
                            self.log(f"DEBUG: Detected claim boundary at element {idx+1}, adding page break")
                            # Insert a page break paragraph before this table
                            page_break_para = self.doc.add_paragraph()
                            run = page_break_para.add_run()
                            run.add_break(WD_BREAK.PAGE)
                            current_anchor._element.addnext(page_break_para._p)
                            current_anchor = page_break_para
                    
                    # Insert element
                    new_el = insert_element_after(current_anchor, el)
                    if new_el:
                        self.log(f"DEBUG: Inserted mapping element {idx+1}/{len(preserved_mapping_elements)}")
                        
                        # Check if this is the mappings paragraph text and add blank line after it
                        if el.tag == qn('w:p'):
                            try:
                                p = Paragraph(el, self.doc)
                                text_lower = p.text.lower().strip()
                                if "these are the mappings of the elements" in text_lower:
                                    # # Add blank line after mappings paragraph text
                                    # empty_para_after_text = self.insert_paragraph_after(new_el, "")
                                    # empty_para_after_text.paragraph_format.space_after = Pt(0)
                                    # empty_para_after_text.paragraph_format.space_before = Pt(0)
                                    # current_anchor = empty_para_after_text
                                    # self.log("DEBUG: Added blank line after mappings paragraph text")
                                    continue
                            except Exception as e:
                                self.log(f"DEBUG: Error checking paragraph text: {e}")
                        
                        current_anchor = new_el
                    else:
                        self.log(f"DEBUG: Failed to insert mapping element {idx+1}/{len(preserved_mapping_elements)}")
                
                # Note: No page break before disclaimer in update mode
            elif not self.update_mode or not preserved_mapping_elements:
                self.log("⚠ No preserved mapping elements found, generating fresh mapping tables")
                # Generate fresh mapping tables
                if self.report_type == "Invalidity":
                    mapping_tables = self.find_mapping_tables(self.doc)
                    if not mapping_tables:
                        self.log("Warning: No mapping tables found for Invalidity report.")
                        return
                    
                    # Use the first table as the master template
                    if len(self.ClaimNumbers) > 0:
                        master_idx, master_table = mapping_tables[0]
                        
                        # Process FIRST claim in the existing template table
                        first_claim = self.ClaimNumbers[0]
                        self.update_headers(master_table, first_claim)
                        self.populate_table_with_claim(master_table, first_claim, color_cycle)
                        
                        # For remaining claims, clone the processed master table
                        last_table = master_table
                        
                        for claim_idx in range(1, len(self.ClaimNumbers)):
                            claim_number = self.ClaimNumbers[claim_idx]
                            
                            # Add page break
                            page_break_para = self.doc.add_paragraph()
                            run = page_break_para.add_run()
                            run.add_break(WD_BREAK.PAGE)
                            
                            # Insert page break after last table
                            last_table._tbl.addnext(page_break_para._p)
                            
                            # Clone the master table structure
                            new_table = self.clone_table_structure(master_table)
                            page_break_para._p.addnext(new_table._tbl)
                            
                            # Process this claim
                            self.update_headers(new_table, claim_number)
                            self.populate_table_with_claim(new_table, claim_number, color_cycle)
                            
                            last_table = new_table
                else:  # FTO
                    mapping_table = None
                    for t in self.doc.tables:
                        for row in t.rows:
                            if len(row.cells) >= 2:
                                left_has = any("[CLAIM_ELEMENT]" in p.text for p in row.cells[0].paragraphs)
                                right_has = any("[REFERENCE_DISCLOSURE/S]" in p.text for p in row.cells[1].paragraphs)
                                if left_has or right_has:
                                    mapping_table = t
                                    break
                        if mapping_table:
                            break
                    if not mapping_table:
                        self.log("Warning: No mapping table found for FTO report.")
                        return
                    criteria_fragments = self.extract_claim_fragments_from_excel(self.df)
                    filtered_fragments = [frag for frag in criteria_fragments if frag.strip()]
                    if filtered_fragments:
                        template_row = self.find_placeholder_row_obj(mapping_table)
                        if template_row is None:
                            template_row = self.clone_row_after(mapping_table, mapping_table.rows[-1])
                        for frag_idx, fragment in enumerate(filtered_fragments):
                            if frag_idx == 0:
                                row = template_row
                                self.clear_cell_strict(row.cells[0])
                                self.clear_cell_strict(row.cells[1])
                            else:
                                row = self.clone_row_after(mapping_table, template_row)
                                self.clear_cell_strict(row.cells[0])
                                self.clear_cell_strict(row.cells[1])
                            p_left = row.cells[0].paragraphs[0]
                            run_left = p_left.add_run(fragment)
                            run_left.font.name = 'Inter'
                            run_left.font.size = Pt(9)
                            run_left.bold = True
                            run_left.font.color.rgb = color_cycle[self.global_color_index % 2]
                            self.global_color_index += 1
                            if frag_idx > 0:
                                p_left.paragraph_format.left_indent = Inches(0.23)
                            if fragment.strip():
                                for p in list(row.cells[1].paragraphs):
                                    p._element.getparent().remove(p._element)
                                main_para = row.cells[1].add_paragraph()
                                # Ensure zero spacing and single line spacing as per Colab changes
                                main_para.paragraph_format.space_after = Pt(0)
                                main_para.paragraph_format.space_before = Pt(0)
                                main_para.paragraph_format.line_spacing = 1.0
                                for i, ref in enumerate(self.sorted_references):
                                    if i > 0:
                                        main_para.add_run("\n\n")
                                    if ref.isNPL:
                                        heading_text = f"{ref.Rank}. {ref.Title}"
                                    else:
                                        heading_text = f"{ref.Rank}. {ref.RawPublicationNumber}"
                                    heading_run = main_para.add_run(heading_text)
                                    heading_run.font.name = 'Inter'
                                    heading_run.font.size = Pt(9)
                                    heading_run.bold = True
                                    placeholder_run = main_para.add_run("\n")
                                    placeholder_run.font.name = 'Inter'
                                    placeholder_run.font.size = Pt(9)
                    else:
                        self.log("Warning: No criteria fragments found for FTO mapping table.")

            # Build mappings intro text
            claims_text_joined = self.format_claims_as_ranges(self.ClaimNumbers)
            para_text = (
                f"These are the mappings of the elements of {self.claim_word} {claims_text_joined} of the {self.short_patent_name_lower} "
                "against similar disclosures from the selected references. Matching with the claim elements "
                "may vary from somewhat relevant to strongly-matched."
            )

            if self.update_mode:
                # In Update mode: force placement immediately after the header,
                # with a blank paragraph before and after the intro paragraph.
                mappings_header = None
                for p in self.doc.paragraphs:
                    if "mappings based on selected references" in p.text.lower():
                        mappings_header = p
                        break
                if mappings_header is not None:
                    # Blank before intro
                    pre_blank = self.insert_paragraph_after(mappings_header, "")
                    pre_blank.paragraph_format.space_after = Pt(0)
                    pre_blank.paragraph_format.space_before = Pt(0)
                    # Intro paragraph (ensure 0pt spacing before/after)
                    intro_para = self.insert_paragraph_after(pre_blank, para_text)
                    self.apply_font_style(intro_para)
                    intro_para.paragraph_format.space_after = Pt(0)
                    intro_para.paragraph_format.space_before = Pt(0)
                    # Blank after intro
                    post_blank = self.insert_paragraph_after(intro_para, "")
                    post_blank.paragraph_format.space_after = Pt(0)
                    post_blank.paragraph_format.space_before = Pt(0)

                    # Remove any duplicate mappings intro paragraphs between the mappings header and disclaimer
                    try:
                        body = self.doc.element.body
                        keyphrase = "these are the mappings of the elements"
                        disclaimer_el = None
                        # Find disclaimer paragraph element
                        for child in list(body):
                            if child.tag == qn('w:p'):
                                p = Paragraph(child, self.doc)
                                if "disclaimer" in (p.text or "").lower():
                                    disclaimer_el = child
                                    break
                        remove_started = False
                        to_remove = []
                        for child in list(body):
                            if child.tag == qn('w:p'):
                                p = Paragraph(child, self.doc)
                                text_lower = (p.text or "").lower().strip()
                                if not remove_started and "mappings based on selected references" in text_lower:
                                    remove_started = True
                                    continue
                                if remove_started:
                                    if disclaimer_el is not None and child is disclaimer_el:
                                        break
                                    # Mark duplicates that are not the freshly inserted intro paragraph
                                    if keyphrase in text_lower and child is not intro_para._p:
                                        to_remove.append(child)
                        for el in to_remove:
                            body.remove(el)
                    except Exception:
                        pass
                else:
                    # Fallback to placeholder if header not found
                    mappings_paragraph = self.find_paragraph_with_placeholder(self.doc, "[MAPPINGS_PARAGRAPH]")
                    if mappings_paragraph:
                        mappings_paragraph.text = para_text
                        self.apply_font_style(mappings_paragraph)
                        # Add surrounding blanks
                        pre_blank = self.insert_paragraph_after(mappings_paragraph, "")
                        pre_blank.paragraph_format.space_after = Pt(0)
                        pre_blank.paragraph_format.space_before = Pt(0)
                        post_blank = self.insert_paragraph_after(pre_blank, "")
                        post_blank.paragraph_format.space_after = Pt(0)
                        post_blank.paragraph_format.space_before = Pt(0)
                    else:
                        self.log("Warning: Could not place mappings intro paragraph (no header or placeholder found).")
            else:
                # New mode: prefer placeholder; if missing, place after header (no-op if neither exists)
                mappings_paragraph = self.find_paragraph_with_placeholder(self.doc, "[MAPPINGS_PARAGRAPH]")
                if mappings_paragraph:
                    mappings_paragraph.text = para_text
                    self.apply_font_style(mappings_paragraph)
                else:
                    mappings_header = None
                    for p in self.doc.paragraphs:
                        if "mappings based on selected references" in p.text.lower():
                            mappings_header = p
                            break
                    if mappings_header is not None:
                        intro_para = self.insert_paragraph_after(mappings_header, para_text)
                        self.apply_font_style(intro_para)
                    else:
                        self.log("Warning: [MAPPINGS_PARAGRAPH] placeholder not found and header missing.")
            self.log("Mappings section processed.")
            # Diagnostics: capture section indices after changes
            try:
                def _idx_of2(text):
                    from docx.text.paragraph import Paragraph as _Paragraph
                    body = self.doc.element.body
                    for i, el in enumerate(list(body)):
                        if el.tag.endswith('p'):
                            p = _Paragraph(el, self.doc)
                            if text.lower() in (p.text or '').lower():
                                return i
                    return None
                criteria_idx_post = _idx_of2('criteria for the publication search')
                mappings_idx_post = _idx_of2('mappings based on selected references')
                about_idx_post = _idx_of2('about us')
                disclaimer_idx_post = _idx_of2('disclaimer')
                self.log(f"DEBUG: [post-mappings] indices → criteria={criteria_idx_post}, mappings={mappings_idx_post}, about={about_idx_post}, disclaimer={disclaimer_idx_post}")
            except Exception:
                pass
        except Exception as e:
            self.log(f"Error processing mappings section: {str(e)}")
            raise

    def process_search_strings(self):
      self.log("Processing search strings section...")
      try:
          # Get target document: gen_doc for update mode, doc otherwise
          target_doc = self.get_target_doc("search")
          
          self.search_results_df, self.total_search_hits = self.extract_search_results()
          if not self.search_results_df.empty:
              self.replace_in_paragraphs_and_tables(target_doc, {"[HITS_TOTAL]": f"{self.total_search_hits:,}"})
              self.search_results_df['Database_norm'] = self.search_results_df['Database'].astype(str).str.strip().str.lower()
              typo_map = {
                  "gogle patents": "google patents",
                  "google patents ": "google patents",
                  "espacent": "espacenet",
                  "espacents": "espacenet",
                  "espacenet ": "espacenet",
                  "google scholar": "google scholar",
                  "google search": "google search",
                  "science direct": "science direct",
                  "wipo patentscope": "wipo patentscope",
                  "j-platpat": "j-platpat",
                  "kipris": "kipris",
                  "manualslib": "manualslib",
                  "internet archive": "internet archive",
                  "amazon": "amazon",
                  "alibaba": "alibaba",
              }
              self.search_results_df['Database_norm'] = self.search_results_df['Database_norm'].replace(typo_map)
              canonical = [
                  'orbit', 'google patents', 'espacenet', 'pqai', 'kipris', 'j-platpat',
                  'wipo patentscope', 'google scholar', 'google search', 'science direct', 'ieee',
                  'manualslib', 'amazon', 'alibaba', 'internet archive'
              ]
              def canonicalize(name):
                  if name in canonical:
                      return name
                  match = difflib.get_close_matches(name, canonical, n=1, cutoff=0.75)
                  return match[0] if match else name
              self.search_results_df['Database_norm'] = self.search_results_df['Database_norm'].apply(canonicalize)
              sort_order = {
                  'orbit': 1, 'google patents': 2, 'espacenet': 3, 'pqai': 4, 'kipris': 5, 'j-platpat': 6,
                  'wipo patentscope': 7, 'google scholar': 8, 'google search': 9, 'science direct': 10, 'ieee': 11,
                  'manualslib': 12, 'amazon': 13, 'alibaba': 14, 'internet archive': 15
              }
              self.search_results_df['Database_rank'] = self.search_results_df['Database_norm'].map(sort_order).fillna(16)
              self.search_results_df = self.search_results_df.sort_values(by=['Database_rank', 'Database_norm'], ascending=[True, True], kind='mergesort').reset_index(drop=True)
              self.search_results_df['S/No'] = range(1, len(self.search_results_df) + 1)
              display_map = {
                  'orbit': 'Orbit', 'google patents': 'Google Patents', 'espacenet': 'Espacenet', 'pqai': 'PQAI',
                  'kipris': 'KIPRIS', 'j-platpat': 'J-PlatPat', 'wipo patentscope': 'WIPO Patentscope',
                  'google scholar': 'Google Scholar', 'google search': 'Google Search', 'science direct': 'Science Direct',
                  'ieee': 'IEEE', 'manualslib': 'ManualsLib', 'amazon': 'Amazon', 'alibaba': 'Alibaba',
                  'internet archive': 'Internet Archive'
              }
              self.search_results_df['Database'] = self.search_results_df['Database_norm'].map(display_map).fillna(self.search_results_df['Database'])
              self.search_results_df['Hits'] = self.search_results_df['Hits'].apply(lambda x: f"{int(str(x).replace(',', '')):,}" if str(x).replace(',', '').isdigit() else str(x))
              self.search_results_df = self.search_results_df.drop(columns=['Database_norm', 'Database_rank'])

              ss_table = self.find_table_with_placeholder(target_doc, "[ROW_INDEX]") or \
                        self.find_table_with_placeholder(target_doc, "[DB]") or \
                        self.find_table_with_placeholder(target_doc, "[SCOPE]") or \
                        self.find_table_with_placeholder(target_doc, "[QUERY]") or \
                        self.find_table_with_placeholder(target_doc, "[HITS]")
              if ss_table:
                  template_row = self.find_row_with_placeholder(ss_table, "[ROW_INDEX]") or ss_table.rows[-1]
                  def format_query_cell(cell, query_text, database_name):
                      self.clear_cell(cell)
                      para = cell.paragraphs[0]
                      special_databases = ['PQAI', 'Google Search', 'Amazon', 'Alibaba']
                      needs_quotes_and_italics = database_name in special_databases
                      # Extended operator pattern to match boolean and proximity tokens
                      op_pattern = r'(\bAND\b|\bOR\b|\bNOT\b|\bNEAR/\d+\b|\bNEAR\d*\b|\bADJ\d*\b|\b[FPS]\b|\b\d+[DW]\b)'
                      if needs_quotes_and_italics:
                          quote_run = para.add_run('"')
                          quote_run.font.name = 'Inter'
                          quote_run.font.size = Pt(9)
                          parts = re.split(op_pattern, query_text)
                          for part in parts:
                              if not part:
                                  continue
                              if re.match(r'^(AND|OR|NOT|NEAR/\d+|NEAR\d*|ADJ\d*|[FPS]|\d+[DW])$', part):
                                  run = para.add_run(part)
                                  run.italic = True
                                  run.font.name = 'Inter'
                                  run.font.size = Pt(9)
                              else:
                                  run = para.add_run(part)
                                  run.italic = True
                                  run.font.name = 'Inter'
                                  run.font.size = Pt(9)
                          quote_run = para.add_run('"')
                          quote_run.font.name = 'Inter'
                          quote_run.font.size = Pt(9)
                      else:
                          parts = re.split(op_pattern, query_text)
                          for part in parts:
                              if not part:
                                  continue
                              if re.match(r'^(AND|OR|NOT|NEAR/\d+|NEAR\d*|ADJ\d*|[FPS]|\d+[DW])$', part):
                                  run = para.add_run(part)
                                  run.bold = True
                                  run.font.name = 'Inter'
                                  run.font.size = Pt(9)
                              else:
                                  run = para.add_run(part)
                                  run.font.name = 'Inter'
                                  run.font.size = Pt(9)
                  if len(self.search_results_df) > 0:
                      first = self.search_results_df.iloc[0]
                      self.set_cell_text(template_row.cells[0], str(first['S/No']), size=9, bold=True)
                      self.set_cell_text(template_row.cells[1], first['Database'], size=9)
                      self.set_cell_text(template_row.cells[2], first['Scope'], size=9)
                      format_query_cell(template_row.cells[3], first['Query'], first['Database'])
                      self.set_cell_text(template_row.cells[4], first['Hits'], size=9)
                      template_row.cells[4].paragraphs[0].alignment = 1
                      for _, row in self.search_results_df.iloc[1:].iterrows():
                          new_row = self.clone_row_after(ss_table, template_row)
                          self.set_cell_text(new_row.cells[0], str(row['S/No']), size=9, bold=True)
                          self.set_cell_text(new_row.cells[1], row['Database'], size=9)
                          self.set_cell_text(new_row.cells[2], row['Scope'], size=9)
                          format_query_cell(new_row.cells[3], row['Query'], row['Database'])
                          self.set_cell_text(new_row.cells[4], row['Hits'], size=9)
                          new_row.cells[4].paragraphs[0].alignment = 1
                      total_row = self.clone_row_after(ss_table, template_row)
                      # Remove borders for all cells in the TOTAL row
                      for cell in total_row.cells:
                          tcPr = cell._element.get_or_add_tcPr()
                          tcBorders = tcPr.xpath('./w:tcBorders')
                          if tcBorders:
                              tcBorders[0].clear()
                          else:
                              tcBorders = OxmlElement('w:tcBorders')
                              tcPr.append(tcBorders)
                          for border_name in ['left', 'right', 'bottom']:
                              border = OxmlElement(f'w:{border_name}')
                              border.set(qn('w:val'), 'nil')
                              tcBorders.append(border)
                      self.set_cell_text(total_row.cells[0], "", size=9)
                      self.set_cell_text(total_row.cells[1], "", size=9)
                      self.set_cell_text(total_row.cells[2], "", size=9)
                      self.set_cell_text(total_row.cells[3], "TOTAL", bold=True, size=10)
                      self.set_cell_text(total_row.cells[4], f"{self.total_search_hits:,}", bold=True, size=10)
                      total_row.cells[4].paragraphs[0].alignment = 1
              else:
                  self.log("Warning: Search strings table not found.")
          else:
              self.log("Warning: No search results found in Excel.")
          self.log("Search strings section processed.")
      except Exception as e:
          self.log(f"Error processing search strings section: {str(e)}")
          raise

    def find_paragraph_contains(self, doc, text):
        """Find paragraph containing specific text (case-insensitive)."""
        search_text = text.lower()
        for p in doc.paragraphs:
            if search_text in (p.text or '').lower():
                return p
        return None

    def simple_replace_section(self, src_doc, dst_doc, start_heading_text, end_heading_text):
        """
        Simple section replacement that just copies content without complex boundary detection.
        Matches colab implementation (lines 2399-2485).
        """
        try:
            self.log(f"    🔍 Looking for '{start_heading_text}' in source document...")
            start_p = self.find_paragraph_contains(src_doc, start_heading_text)
            if not start_p:
                self.log(f"    ❌ Could not find '{start_heading_text}' in source document")
                return False

            self.log(f"    ✅ Found '{start_heading_text}' in source document")

            # Find end boundary
            end_p = None
            if end_heading_text:
                elems = list(src_doc.element.body)
                start_idx = elems.index(start_p._p)

                # Look for end_heading_text after the start position
                for i in range(start_idx + 1, len(elems)):
                    if elems[i].tag == qn('w:p'):
                        p = Paragraph(elems[i], src_doc)
                        if end_heading_text.lower() in (p.text or '').lower():
                            end_p = p
                            break

                if not end_p:
                    self.log(f"    ⚠️  Could not find '{end_heading_text}' after '{start_heading_text}' in source document, using end of document")

            # Get source content (excluding the heading to avoid duplicates)
            elems = list(src_doc.element.body)
            start_idx = elems.index(start_p._p)
            end_idx = elems.index(end_p._p) if end_p is not None else len(elems)
            # Skip the heading paragraph to avoid duplicates
            src_slice = elems[start_idx + 1:end_idx]
            self.log(f"    📋 Found {len(src_slice)} elements to copy from source (excluding heading)")

            # Find destination section
            self.log(f"    🔍 Looking for '{start_heading_text}' in destination document...")
            dst_start_p = self.find_paragraph_contains(dst_doc, start_heading_text)
            if not dst_start_p:
                # If destination start heading doesn't exist, insert before end_heading if possible
                self.log(f"    ❌ Could not find '{start_heading_text}' in destination document")
                if end_heading_text:
                    insert_before = self.find_paragraph_contains(dst_doc, end_heading_text)
                    if insert_before:
                        self.log(f"    ➕ Inserting new section '{start_heading_text}' before '{end_heading_text}' in destination (including source heading)")
                        # Insert the source heading and its content before insert_before
                        elems_dst = list(dst_doc.element.body)
                        insert_idx = elems_dst.index(insert_before._p)
                        # Build slice including heading in source
                        full_src_slice = elems[start_idx:end_idx]
                        for el in reversed(full_src_slice):
                            new_el = deepcopy(el)
                            insert_before._p.addprevious(new_el)
                        self.log(f"    ✅ Inserted section '{start_heading_text}' into destination")
                        return True
                return False

            self.log(f"    ✅ Found '{start_heading_text}' in destination document")

            # Find end boundary in destination
            dst_end_p = None
            if end_heading_text:
                dst_elems = list(dst_doc.element.body)
                dst_start_idx = dst_elems.index(dst_start_p._p)

                # Look for end_heading_text after the start position
                for i in range(dst_start_idx + 1, len(dst_elems)):
                    if dst_elems[i].tag == qn('w:p'):
                        p = Paragraph(dst_elems[i], dst_doc)
                        if end_heading_text.lower() in (p.text or '').lower():
                            dst_end_p = p
                            break

                if not dst_end_p:
                    # Fallback: find any next major section as boundary to avoid wiping whole doc
                    self.log(f"    ⚠️  Could not find '{end_heading_text}' after '{start_heading_text}' in destination document, searching for next section boundary")
                    major_sections = [
                        'other related references found',
                        'patent-at-issue',
                        'criteria for the publication search',
                        'mappings based on selected references',
                        'disclaimer',
                        'appendix'
                    ]
                    for i in range(dst_start_idx + 1, len(dst_elems)):
                        if dst_elems[i].tag == qn('w:p'):
                            p = Paragraph(dst_elems[i], dst_doc)
                            text_lower = (p.text or '').lower()
                            if any(ms in text_lower for ms in major_sections):
                                dst_end_p = p
                                self.log(f"    ✅ Found boundary at: '{p.text.strip()}'")
                                break
                    if not dst_end_p:
                        self.log(f"    ⚠️  No section boundary found, using end of document")

            # Remove existing content (keep the heading)
            dst_elems = list(dst_doc.element.body)
            dst_start_idx = dst_elems.index(dst_start_p._p)
            dst_end_idx = dst_elems.index(dst_end_p._p) if dst_end_p is not None else len(dst_elems)

            # Remove elements after the heading
            elements_to_remove = dst_end_idx - dst_start_idx - 1
            self.log(f"    🗑️  Removing {elements_to_remove} existing elements from destination")

            # Remove elements in reverse order
            for i in range(dst_end_idx - 1, dst_start_idx, -1):
                if i < len(dst_elems):
                    el = dst_elems[i]
                    dst_doc.element.body.remove(el)

            # Insert new content after the heading
            self.log(f"    ➕ Inserting {len(src_slice)} elements into destination")

            # Insert elements in reverse order to maintain correct order
            for i, el in enumerate(reversed(src_slice)):
                new_el = deepcopy(el)
                dst_start_p._p.addnext(new_el)
                if i % 10 == 0:  # Log every 10th element to avoid spam
                    self.log(f"      Inserted element {i+1}/{len(src_slice)}")

            self.log(f"    ✅ Successfully replaced section '{start_heading_text}'")
            return True
        except Exception as e:
            self.log(f"    ❌ Error replacing section '{start_heading_text}': {str(e)}")
            return False

    def merge_generated_sections(self):
        """
        Merge sections from gen_doc into doc in update mode.
        This copies regenerated sections (Title, Objectives, References, Patent-at-Issue, Criteria, Search Strings)
        from gen_doc into doc, which already contains preserved Mappings.
        Matches the colab implementation (lines 2488-2578).
        """
        if not self.update_mode or self.gen_doc is None:
            self.log("Not in update mode or gen_doc not available - skipping merge")
            return
        
        self.log("🔄 Starting merge process in update mode...")

        # Debug: List all headings in both documents to see what we're working with
        def list_headings(doc_obj, doc_name):
            self.log(f"\n📋 Headings in {doc_name}:")
            count = 0
            for i, p in enumerate(doc_obj.paragraphs):
                text = p.text.strip()
                if text and len(text) < 100 and (text.isupper() or any(word in text.lower() for word in ['title', 'contents', 'objective', 'references', 'patent', 'criteria', 'mappings', 'search', 'appendix', 'disclaimer', 'about'])):
                    self.log(f"  {i}: '{text}'")
                    count += 1
                    if count > 20:  # Limit output
                        break

        list_headings(self.gen_doc, "Generated Document")
        list_headings(self.doc, "Edited Document")

        # Additional debug: Check if gen_doc has the expected content
        self.log(f"\n🔍 Checking gen_doc content:")
        self.log(f"  - gen_doc has {len(self.gen_doc.paragraphs)} paragraphs")
        self.log(f"  - gen_doc has {len(self.gen_doc.tables)} tables")

        # Check if the search strategies content exists in gen_doc
        search_strategies_found = False
        for p in self.gen_doc.paragraphs:
            if "search strategy below resulted in" in p.text.lower():
                search_strategies_found = True
                self.log(f"  ✅ Found search strategies content: '{p.text[:100]}...'")
                break

        if not search_strategies_found:
            self.log("  ❌ No search strategies content found in gen_doc")

        # Replace full Title Page (first-page content) from gen_doc into doc (up to OBJECTIVE)
        self.log("\n📄 Replacing Title Page (full first-page content)...")
        try:
            gen_obj_p = self.find_paragraph_contains(self.gen_doc, "objective")
            doc_obj_p = self.find_paragraph_contains(self.doc, "objective")
            if gen_obj_p is None or doc_obj_p is None:
                self.log("  ⚠️  Could not find 'OBJECTIVE' heading in one of the documents; skipping title page replacement")
            else:
                gen_body = list(self.gen_doc.element.body)
                doc_body = self.doc.element.body
                gen_cut_idx = gen_body.index(gen_obj_p._p)

                # Remove existing first-page elements in destination (before its OBJECTIVE)
                doc_elems = list(doc_body)
                doc_obj_idx = doc_elems.index(doc_obj_p._p)
                for i in range(doc_obj_idx - 1, -1, -1):
                    doc_body.remove(doc_elems[i])

                # Insert source first-page elements before destination OBJECTIVE, maintaining order
                insert_before = doc_obj_p._p
                for el in gen_body[:gen_cut_idx]:
                    new_el = deepcopy(el)
                    insert_before.addprevious(new_el)

                self.log("  ✅ Title page replaced from generated document")
        except Exception as e:
            self.log(f"  ⚠️  Warning: Could not replace title page: {str(e)}")

        # Check if sections exist in edited document before copying
        self.log("\n🔍 Checking if sections exist in edited document:")
        obj_exists = self.find_paragraph_contains(self.doc, "objective") is not None
        other_refs_exists = self.find_paragraph_contains(self.doc, "other related references found") is not None
        patent_exists = self.find_paragraph_contains(self.doc, "patent-at-issue") is not None
        criteria_exists = self.find_paragraph_contains(self.doc, "criteria for the publication search") is not None
        appendix_b_exists = self.find_paragraph_contains(self.doc, "appendix b") is not None

        self.log(f"  - Objective exists: {'✅' if obj_exists else '❌'}")
        self.log(f"  - Other Related References exists: {'✅' if other_refs_exists else '❌'}")
        self.log(f"  - Patent-at-Issue exists: {'✅' if patent_exists else '❌'}")
        self.log(f"  - Criteria exists: {'✅' if criteria_exists else '❌'}")
        self.log(f"  - Appendix B exists: {'✅' if appendix_b_exists else '❌'}")
        
        if self.update_mode:
            self.log("  ℹ️  In update mode, Criteria section is preserved from edited document")

        # Simple approach: Replace the existing sections with generated content
        self.log("\n📄 Replacing existing sections with generated content...")

        # Replace Objective section
        self.log("📄 Replacing Objective section...")
        success1 = self.simple_replace_section(self.gen_doc, self.doc, "objective", "other related references found")
        self.log(f"  Result: {'✅ Success' if success1 else '❌ Failed'}")

        # Replace Other Related References section
        self.log("📄 Replacing Other Related References section...")
        success2 = self.simple_replace_section(self.gen_doc, self.doc, "other related references found", "patent-at-issue")
        self.log(f"  Result: {'✅ Success' if success2 else '❌ Failed'}")

        # Replace Patent-at-Issue section
        self.log("📄 Replacing Patent-at-Issue section...")
        success3 = self.simple_replace_section(self.gen_doc, self.doc, "patent-at-issue", "criteria for the publication search")
        self.log(f"  Result: {'✅ Success' if success3 else '❌ Failed'}")

        # Skip Criteria section in update mode - it's already preserved in doc
        if not self.update_mode:
            # Replace Criteria section (only in New mode)
            self.log("📄 Replacing Criteria section...")
            success4 = self.simple_replace_section(self.gen_doc, self.doc, "criteria for the publication search", "mappings based on selected references")
            self.log(f"  Result: {'✅ Success' if success4 else '❌ Failed'}")
        else:
            self.log("📄 Skipping Criteria section (already preserved from edited document)")
            success4 = True  # Mark as success since we're intentionally skipping

        # Replace Search Strings section
        self.log("📄 Replacing Search Strings section...")
        # Pre-merge diagnostics for Appendix B boundaries
        try:
            def _find_idx(doc_obj, text):
                from docx.text.paragraph import Paragraph as _Paragraph
                body = doc_obj.element.body
                for i, el in enumerate(list(body)):
                    if el.tag.endswith('p'):
                        p = _Paragraph(el, doc_obj)
                        if text.lower() in (p.text or '').lower():
                            return i
                return None
            dst_idx_appb = _find_idx(self.doc, 'appendix b') or _find_idx(self.doc, 'appendix b: search strategies')
            dst_idx_map = _find_idx(self.doc, 'mappings based on selected references')
            dst_idx_about = _find_idx(self.doc, 'about us')
            dst_idx_disc = _find_idx(self.doc, 'disclaimer')
            self.log(f"DEBUG: [pre-AppB-merge] dest indices → appB={dst_idx_appb}, mappings={dst_idx_map}, about={dst_idx_about}, disclaimer={dst_idx_disc}")
        except Exception:
            pass
        success5 = (self.simple_replace_section(self.gen_doc, self.doc, "appendix b: search strategies", "disclaimer") or
                   self.simple_replace_section(self.gen_doc, self.doc, "appendix b", "disclaimer") or
                   self.simple_replace_section(self.gen_doc, self.doc, "search strategies", "disclaimer"))
        self.log(f"  Result: {'✅ Success' if success5 else '❌ Failed'}")

        # Post-merge diagnostics: check if ABOUT US precedes MAPPINGS
        try:
            dst_idx_map2 = _find_idx(self.doc, 'mappings based on selected references')
            dst_idx_about2 = _find_idx(self.doc, 'about us')
            self.log(f"DEBUG: [post-AppB-merge] dest indices → mappings={dst_idx_map2}, about={dst_idx_about2}")
            if dst_idx_about2 is not None and dst_idx_map2 is not None and dst_idx_about2 < dst_idx_map2:
                self.log("WARN: ABOUT US appears before the MAPPINGS section after merge. This may cause Mappings to appear after About.")
        except Exception:
            pass

        # Summary - criteria section is skipped in update mode
        if self.update_mode:
            total_expected = 4  # Objective, References, Patent, Search Strings
            total_success = sum([success1, success2, success3, success5]) + (1 if success4 else 0)  # Add criteria as success
            self.log(f"✅ Merge process completed! {total_success}/{total_expected} sections copied successfully.")
        else:
            total_success = sum([success1, success2, success3, success4, success5])
            self.log(f"✅ Merge process completed! {total_success}/5 sections copied successfully.")

        if (self.update_mode and total_success < 4) or (not self.update_mode and total_success < 5):
            self.log("⚠️  Some sections could not be copied. Check the debug output above for details.")

        # After merge, ensure spacing and formatting around ORR and Patent-at-Issue
        try:
            self.ensure_patent_at_issue_spacing_and_format(self.doc)
            # Remove stray ORR heading if it precedes Patent-at-Issue without a table
            self.remove_stray_orr_heading(self.doc)
            # Ensure ORR header is present above the ORR table and add page break before Patent-at-Issue
            self.ensure_orr_header_and_spacing(self.doc)
            # Ensure a page break exists between Criteria and Mappings sections
            self.ensure_page_break_before_mappings(self.doc)
            # Deep diagnostics for mappings placement
            self.debug_mappings_placement(self.doc)
        except Exception as e:
            self.log(f"Warning: Could not normalize Patent-at-Issue heading formatting: {str(e)}")

        # Final safeguard: if ABOUT US precedes MAPPINGS, relocate MAPPINGS to immediately follow CRITERIA
        try:
            self.relocate_mappings_after_criteria_if_needed(self.doc)
        except Exception as e:
            self.log(f"WARN: relocate_mappings_after_criteria_if_needed failed: {e}")

    def relocate_mappings_after_criteria_if_needed(self, doc):
        """
        If 'ABOUT US' appears before 'Mappings Based on Selected References', move the entire
        Mappings section (header + content up to next major section) to immediately follow the
        Criteria section, inserting a page break before the Mappings header.
        """
        try:
            from docx.text.paragraph import Paragraph
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            body = doc.element.body
            elems = list(body)

            # Helper to find first index of paragraph containing text
            def find_idx(substrs):
                for i, el in enumerate(elems):
                    if el.tag.endswith('p'):
                        p = Paragraph(el, doc)
                        t = (p.text or '').strip().lower()
                        for s in (substrs if isinstance(substrs, (list, tuple)) else [substrs]):
                            if s in t:
                                return i
                return None

            criteria_idx = find_idx('criteria for the publication search')
            mappings_idx = find_idx('mappings based on selected references')
            about_idx = find_idx('about us')
            disclaimer_idx = find_idx('disclaimer')

            self.log(f"DEBUG: [relocate] indices before → criteria={criteria_idx}, mappings={mappings_idx}, about={about_idx}, disclaimer={disclaimer_idx}")

            if mappings_idx is None or criteria_idx is None or about_idx is None:
                return
            if about_idx < mappings_idx:
                # Determine end of mappings block: stop at the next major section after mappings
                major_keys = [
                    'disclaimer', 'appendix', 'parola analytics', 'about us',
                    'objective', 'patent-at-issue', 'criteria for the publication search'
                ]
                end_idx = len(elems)
                for i in range(mappings_idx + 1, len(elems)):
                    el = elems[i]
                    if el.tag.endswith('p'):
                        p = Paragraph(el, doc)
                        t = (p.text or '').strip().lower()
                        if any(k in t for k in major_keys):
                            end_idx = i
                            break

                # If there is a page-break paragraph immediately BEFORE the original mappings header, remove it
                try:
                    if mappings_idx - 1 >= 0 and elems[mappings_idx - 1].tag.endswith('p'):
                        prev_el = elems[mappings_idx - 1]
                        has_prev_page_br = any(br.get(qn('w:type')) == 'page' for br in prev_el.xpath('.//w:br'))
                        if has_prev_page_br:
                            body.remove(prev_el)
                            elems = list(body)
                            # Recompute mappings_idx after removal
                            mappings_idx = None
                            for i, el in enumerate(elems):
                                if el.tag.endswith('p'):
                                    p = Paragraph(el, doc)
                                    if 'mappings based on selected references' in (p.text or '').strip().lower():
                                        mappings_idx = i
                                        break
                except Exception:
                    pass

                # Collect elements to move [mappings_idx, end_idx)
                to_move = elems[mappings_idx:end_idx]
                if not to_move:
                    return

                # Find insertion point: end of criteria section.
                # Walk forward from criteria_idx+1 until hitting a major section, insert after the last content before it.
                insert_after_idx = criteria_idx
                for i in range(criteria_idx + 1, len(elems)):
                    el = elems[i]
                    if el.tag.endswith('p'):
                        p = Paragraph(el, doc)
                        t = (p.text or '').strip().lower()
                        if any(k in t for k in ['mappings based on selected references', 'disclaimer', 'appendix', 'about us', 'parola analytics']):
                            break
                    insert_after_idx = i

                insert_ref_el = elems[insert_after_idx]

                # Do NOT insert any additional page break before the moved block.
                # Also strip pageBreakBefore from the MAPPINGS header to avoid implicit breaks.
                header_el = to_move[0]
                try:
                    # Remove any w:pageBreakBefore on the header paragraph
                    pPr_candidates = header_el.xpath('.//w:pPr')
                    if pPr_candidates:
                        pPr_h = pPr_candidates[0]
                        for child in list(pPr_h):
                            if child.tag == qn('w:pageBreakBefore'):
                                pPr_h.remove(child)
                except Exception:
                    pass
                after_break_ref = insert_ref_el

                # Insert copies in order after the break, then remove originals
                for el in to_move:
                    new_el = deepcopy(el)
                    after_break_ref.addnext(new_el)
                    after_break_ref = new_el

                # Remove originals
                for el in to_move:
                    try:
                        body.remove(el)
                    except Exception:
                        pass

                # Refresh indices for diagnostics
                elems2 = list(body)
                def find_idx2(substrs):
                    for i, el in enumerate(elems2):
                        if el.tag.endswith('p'):
                            p = Paragraph(el, doc)
                            t = (p.text or '').strip().lower()
                            for s in (substrs if isinstance(substrs, (list, tuple)) else [substrs]):
                                if s in t:
                                    return i
                    return None
                mappings_idx2 = find_idx2('mappings based on selected references')
                about_idx2 = find_idx2('about us')
                criteria_idx2 = find_idx2('criteria for the publication search')
                self.log(f"DEBUG: [relocate] indices after → criteria={criteria_idx2}, mappings={mappings_idx2}, about={about_idx2}")
                if about_idx2 is not None and mappings_idx2 is not None and about_idx2 < mappings_idx2:
                    self.log("WARN: Relocation attempted but ABOUT US still precedes MAPPINGS.")
                else:
                    self.log("✓ Relocated MAPPINGS to immediately follow CRITERIA with a page break.")
        except Exception as e:
            self.log(f"WARN: relocate_mappings_after_criteria_if_needed encountered an error: {e}")

    def remove_stray_orr_heading(self, doc):
        """
        Remove an ORR heading that appears immediately before Patent-at-Issue
        without an intervening ORR table.
        """
        try:
            body = doc.element.body
            elems = list(body)
            from docx.text.paragraph import Paragraph
            for i, el in enumerate(elems):
                if el.tag.endswith('p'):
                    p = Paragraph(el, doc)
                    txt = (p.text or '').strip().lower()
                    if txt == 'other related references found':
                        # Scan forward until next major section or a table
                        j = i + 1
                        found_table = False
                        hit_boundary = False
                        while j < len(elems):
                            nxt = elems[j]
                            if nxt.tag.endswith('tbl'):
                                found_table = True
                                break
                            if nxt.tag.endswith('p'):
                                pn = Paragraph(nxt, doc)
                                t = (pn.text or '').strip().lower()
                                if any(k in t for k in ['patent-at-issue', 'criteria for', 'mappings based', 'disclaimer', 'appendix']):
                                    hit_boundary = True
                                    break
                            j += 1
                        if hit_boundary and not found_table:
                            body.remove(el)
                            break
        except Exception:
            pass

    def ensure_orr_header_and_spacing(self, doc):
        """
        Ensure an ORR header exists immediately before the ORR table, and add
        a page break before the Patent-at-Issue heading when it follows the ORR table.
        Also normalizes Patent-at-Issue heading formatting.
        """
        try:
            from docx.text.paragraph import Paragraph
            from docx.table import Table
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            body = doc.element.body
            elems = list(body)

            # Locate ORR heading if present
            orr_heading_idx = None
            for i, el in enumerate(elems):
                if el.tag.endswith('p'):
                    p = Paragraph(el, doc)
                    text = (p.text or '').strip().lower()
                    if 'other related references' in text:
                        orr_heading_idx = i
                        break

            # Locate ORR table after heading
            orr_table_idx = None
            if orr_heading_idx is not None:
                for j in range(orr_heading_idx + 1, len(elems)):
                    if elems[j].tag.endswith('tbl'):
                        orr_table_idx = j
                        break

            # Fallback: detect ORR table by header texts
            if orr_table_idx is None:
                for i, el in enumerate(elems):
                    if el.tag.endswith('tbl'):
                        tbl = Table(el, doc)
                        try:
                            if len(tbl.rows) > 0 and len(tbl.rows[0].cells) >= 2:
                                row_texts = [
                                    " ".join((p.text or '') for p in c.paragraphs).strip().lower()
                                    for c in tbl.rows[0].cells
                                ]
                                joined = " | ".join(row_texts)
                                if ('references found' in joined) and (
                                    'assignee' in joined or 'author/publisher' in joined or 'inventor' in joined
                                ):
                                    orr_table_idx = i
                                    break
                        except Exception:
                            continue

            if orr_table_idx is not None:
                # If no ORR heading exists anywhere, insert one immediately before the table
                if orr_heading_idx is None:
                    has_header = False
                    elems = list(body)
                    if orr_table_idx - 1 >= 0 and elems[orr_table_idx - 1].tag.endswith('p'):
                        prev_p = Paragraph(elems[orr_table_idx - 1], doc)
                        if (prev_p.text or '').strip().lower() == 'other related references found':
                            has_header = True
                    if not has_header:
                        new_p = OxmlElement('w:p')
                        r = OxmlElement('w:r')
                        rPr = OxmlElement('w:rPr')
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:ascii'), 'Inter')
                        rFonts.set(qn('w:hAnsi'), 'Inter')
                        rPr.append(rFonts)
                        b = OxmlElement('w:b')
                        rPr.append(b)
                        sz = OxmlElement('w:sz')
                        sz.set(qn('w:val'), '20')  # 10pt
                        rPr.append(sz)
                        szCs = OxmlElement('w:szCs')
                        szCs.set(qn('w:val'), '20')
                        rPr.append(szCs)
                        color = OxmlElement('w:color')
                        color.set(qn('w:val'), '000000')
                        rPr.append(color)
                        r.append(rPr)
                        t = OxmlElement('w:t')
                        t.text = 'OTHER RELATED REFERENCES FOUND'
                        r.append(t)
                        pPr = OxmlElement('w:pPr')
                        spacing = OxmlElement('w:spacing')
                        spacing.set(qn('w:after'), '160')  # 8pt
                        spacing.set(qn('w:line'), '216')   # ~1.08
                        spacing.set(qn('w:lineRule'), 'auto')
                        pPr.append(spacing)
                        new_p.append(pPr)
                        new_p.append(r)
                        body.insert(orr_table_idx, new_p)

                # Add page break before Patent-at-Issue if it immediately follows a table
                pat_p = None
                pat_idx = None
                for j, el in enumerate(list(body)):
                    if el.tag.endswith('p'):
                        p = Paragraph(el, doc)
                        if (p.text or '').strip().lower() == 'patent-at-issue':
                            pat_p = p
                            pat_idx = j
                            break
                if pat_p is not None:
                    elems2 = list(body)
                    prev_is_tbl = (pat_idx - 1 >= 0 and elems2[pat_idx - 1].tag.endswith('tbl'))
                    if prev_is_tbl:
                        # Insert a separate page-break paragraph BEFORE the heading paragraph
                        br = OxmlElement('w:br')
                        br.set(qn('w:type'), 'page')
                        run_element = OxmlElement('w:r')
                        rPr = OxmlElement('w:rPr')
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:ascii'), 'Inter')
                        rFonts.set(qn('w:hAnsi'), 'Inter')
                        rPr.append(rFonts)
                        run_element.append(rPr)
                        run_element.append(br)
                        new_break_p = OxmlElement('w:p')
                        new_break_p.append(run_element)
                        pat_p._p.addprevious(new_break_p)

                    # Normalize PATENT-AT-ISSUE heading formatting with robust style and spacing
                    try:
                        # Apply Heading 1 (fallbacks included)
                        try:
                            pat_p.style = 'Heading 1'
                        except Exception:
                            try:
                                pat_p.style = 'Heading1'
                            except Exception:
                                pass

                        # Enforce Heading 1 at XML level (pStyle) and outline level 0
                        pPr = pat_p._p.get_or_add_pPr()
                        # Remove existing pStyle to avoid duplicates
                        for el in list(pPr):
                            if el.tag == qn('w:pStyle'):
                                pPr.remove(el)
                        pStyle = OxmlElement('w:pStyle')
                        pStyle.set(qn('w:val'), 'Heading1')
                        pPr.append(pStyle)
                        # Ensure outline level 0
                        for el in list(pPr):
                            if el.tag == qn('w:outlineLvl'):
                                pPr.remove(el)
                        outline = OxmlElement('w:outlineLvl')
                        outline.set(qn('w:val'), '0')
                        pPr.append(outline)

                        # Paragraph spacing
                        pf = pat_p.paragraph_format
                        pf.line_spacing = 1.0
                        pf.space_before = Pt(0)
                        pf.space_after = Pt(0)

                        # Hard enforce 0pt spacing at XML level and disable auto-spacing
                        spacing_el = None
                        for el in list(pPr):
                            if el.tag == qn('w:spacing'):
                                spacing_el = el
                                break
                        if spacing_el is None:
                            spacing_el = OxmlElement('w:spacing')
                            pPr.append(spacing_el)
                        spacing_el.set(qn('w:before'), '0')
                        spacing_el.set(qn('w:after'), '0')
                        spacing_el.set(qn('w:beforeAutospacing'), '0')
                        spacing_el.set(qn('w:afterAutospacing'), '0')
                        spacing_el.set(qn('w:lineRule'), 'auto')

                        # Make sure the heading has visible text and correct run formatting
                        if not pat_p.runs:
                            pat_p.add_run("")
                        for r in pat_p.runs:
                            if r.text:
                                r.text = r.text.upper()
                            r.font.bold = True
                            r.font.size = Pt(12)
                            try:
                                r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
                            except Exception:
                                pass

                        # Normalize the next paragraph's spacing-before to 0 to eliminate visual gap
                        nxt_el = pat_p._p.getnext()
                        if nxt_el is not None and nxt_el.tag == qn('w:p'):
                            from docx.text.paragraph import Paragraph as _Paragraph
                            nxt_p = _Paragraph(nxt_el, doc)
                            npf = nxt_p.paragraph_format
                            npf.space_before = Pt(0)
                            pPr_n = nxt_p._p.get_or_add_pPr()
                            spacing_n = None
                            for el in list(pPr_n):
                                if el.tag == qn('w:spacing'):
                                    spacing_n = el
                                    break
                            if spacing_n is None:
                                spacing_n = OxmlElement('w:spacing')
                                pPr_n.append(spacing_n)
                            spacing_n.set(qn('w:before'), '0')
                            spacing_n.set(qn('w:beforeAutospacing'), '0')
                    except Exception:
                        pass
        except Exception:
            pass

    def ensure_page_break_before_mappings(self, doc):
        """
        Ensure that the "Mappings Based on Selected References" section starts on a new page
        immediately after the "CRITERIA FOR THE PUBLICATION SEARCH" section.
        """
        try:
            from docx.text.paragraph import Paragraph
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            body = doc.element.body
            elems = list(body)

            criteria_idx = None
            mappings_idx = None

            # Locate criteria heading and mappings heading
            self.log("DEBUG: ensure_page_break_before_mappings - scanning for section boundaries...")
            for i, el in enumerate(elems):
                if el.tag.endswith('p'):
                    p = Paragraph(el, doc)
                    text_lower = (p.text or '').strip().lower()
                    if criteria_idx is None and 'criteria for the publication search' in text_lower:
                        criteria_idx = i
                        self.log(f"  - Found CRITERIA heading at index {criteria_idx} -> '{(p.text or '').strip()}'")
                    if mappings_idx is None and 'mappings based on selected references' in text_lower:
                        mappings_idx = i
                        self.log(f"  - Found MAPPINGS heading at index {mappings_idx} -> '{(p.text or '').strip()}'")
                    if criteria_idx is not None and mappings_idx is not None:
                        break

            if criteria_idx is None or mappings_idx is None:
                self.log(f"DEBUG: Section indices not found (criteria_idx={criteria_idx}, mappings_idx={mappings_idx}) - skipping break insert")
                return

            # If mappings header is not on a fresh page after criteria, insert a break
            # Check if the element right before mappings is a page break paragraph
            prev_el = elems[mappings_idx - 1] if mappings_idx - 1 >= 0 else None
            is_prev_break_p = False
            if prev_el is not None and prev_el.tag.endswith('p'):
                # Detect w:br inside the previous paragraph
                try:
                    for br in prev_el.xpath('.//w:br'):
                        if br.get(qn('w:type')) == 'page':
                            is_prev_break_p = True
                            break
                except Exception:
                    pass

            self.log(f"DEBUG: Before insert check -> is_prev_break_p={is_prev_break_p}, criteria_idx={criteria_idx}, mappings_idx={mappings_idx}")
            if not is_prev_break_p:
                # Insert a page break paragraph immediately before the mappings header
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                run_element = OxmlElement('w:r')
                run_element.append(br)
                break_para = OxmlElement('w:p')
                break_para.append(run_element)
                # Insert into body before mappings header element
                elems_m = list(body)
                body.insert(mappings_idx, break_para)
                self.log(f"DEBUG: Inserted page-break paragraph before mappings at index {mappings_idx}")
            else:
                self.log("DEBUG: Page break already present before mappings - no insertion needed")

            # Additionally enforce page break via pageBreakBefore on the mappings header itself
            try:
                elems_now = list(body)
                if mappings_idx < len(elems_now) and elems_now[mappings_idx].tag.endswith('p'):
                    p_m = Paragraph(elems_now[mappings_idx], doc)
                    pPr_m = p_m._p.get_or_add_pPr()
                    # Remove existing pageBreakBefore if any, then set to true
                    for el in list(pPr_m):
                        if el.tag == qn('w:pageBreakBefore'):
                            pPr_m.remove(el)
                    pb = OxmlElement('w:pageBreakBefore')
                    pb.set(qn('w:val'), 'true')
                    pPr_m.append(pb)
                    self.log("DEBUG: Set pageBreakBefore on mappings heading paragraph")
            except Exception:
                pass

            # Diagnostics: section summary for the document
            try:
                self.log("DEBUG: Document sections summary (count and start types):")
                try:
                    from docx.enum.section import WD_SECTION_START
                except Exception:
                    WD_SECTION_START = None
                for si, sec in enumerate(doc.sections):
                    stype = None
                    try:
                        stype = getattr(sec, 'start_type', None)
                        if WD_SECTION_START and stype is not None:
                            stype = WD_SECTION_START(stype).name
                    except Exception:
                        pass
                    self.log(f"  Section {si}: start_type={stype}")
            except Exception:
                pass
        except Exception:
            # Swallow exceptions but log minimal info to aid diagnosis
            try:
                import traceback
                self.log(f"WARN: ensure_page_break_before_mappings encountered an exception:\n{traceback.format_exc()}")
            except Exception:
                pass

    def debug_mappings_placement(self, doc):
        """
        Diagnostic logging to understand why the Mappings section placement may be incorrect
        in certain source documents. Logs surrounding elements, breaks, and paragraph props.
        """
        try:
            from docx.text.paragraph import Paragraph
            from docx.table import Table
            from docx.oxml.ns import qn

            body = doc.element.body
            elems = list(body)

            def find_idx(substr_list):
                for i, el in enumerate(elems):
                    if el.tag.endswith('p'):
                        p = Paragraph(el, doc)
                        t = (p.text or '').strip().lower()
                        for s in substr_list:
                            if s in t:
                                return i
                return None

            criteria_idx = find_idx(['criteria for the publication search'])
            mappings_idx = find_idx(['mappings based on selected references'])
            disclaimer_idx = find_idx(['disclaimer'])
            about_idx = find_idx(['about us'])

            self.log("DEBUG: debug_mappings_placement → indices:")
            self.log(f"  criteria_idx={criteria_idx}, mappings_idx={mappings_idx}, disclaimer_idx={disclaimer_idx}, about_idx={about_idx}")

            def dump_para_info(idx, label):
                try:
                    if idx is None or idx < 0 or idx >= len(elems):
                        self.log(f"  {label}: not found")
                        return
                    el = elems[idx]
                    if not el.tag.endswith('p'):
                        self.log(f"  {label}: element at {idx} is not a paragraph ({el.tag})")
                        return
                    p = Paragraph(el, doc)
                    t = (p.text or '').strip()
                    pPr = p._p.get_or_add_pPr()
                    # pageBreakBefore
                    pbb = None
                    for child in list(pPr):
                        if child.tag == qn('w:pageBreakBefore'):
                            pbb = child
                            break
                    spacing_el = None
                    for child in list(pPr):
                        if child.tag == qn('w:spacing'):
                            spacing_el = child
                            break
                    # keepNext / keepLines
                    keep_next = any(child.tag == qn('w:keepNext') for child in list(pPr))
                    keep_lines = any(child.tag == qn('w:keepLines') for child in list(pPr))
                    # style id
                    pstyle = None
                    for child in list(pPr):
                        if child.tag == qn('w:pStyle'):
                            pstyle = child.get(qn('w:val'))
                            break
                    has_br = any(True for _ in el.xpath('.//w:br'))
                    br_types = [br.get(qn('w:type')) for br in el.xpath('.//w:br')]
                    self.log(f"  {label} @ {idx}: text='{t[:80]}', style={pstyle}, keepNext={keep_next}, keepLines={keep_lines}, pageBreakBefore={'true' if pbb is not None else 'false'}, has_br={has_br}, br_types={br_types}")
                    if spacing_el is not None:
                        self.log(f"    spacing before={spacing_el.get(qn('w:before'))}, after={spacing_el.get(qn('w:after'))}, beforeAuto={spacing_el.get(qn('w:beforeAutospacing'))}, afterAuto={spacing_el.get(qn('w:afterAutospacing'))}")
                except Exception:
                    pass

            dump_para_info(criteria_idx, 'CRITERIA')
            dump_para_info(mappings_idx, 'MAPPINGS')
            dump_para_info(mappings_idx - 1 if mappings_idx else None, 'PREV_OF_MAPPINGS')
            dump_para_info(mappings_idx + 1 if mappings_idx else None, 'NEXT_OF_MAPPINGS')
            dump_para_info(about_idx, 'ABOUT_US')

            # Scan a small window around mappings to log types
            if mappings_idx is not None:
                start = max(0, mappings_idx - 5)
                end = min(len(elems), mappings_idx + 6)
                self.log("DEBUG: Elements around MAPPINGS (type summary):")
                for i in range(start, end):
                    el = elems[i]
                    kind = 'tbl' if el.tag.endswith('tbl') else ('p' if el.tag.endswith('p') else el.tag)
                    text = ''
                    if kind == 'p':
                        try:
                            p = Paragraph(el, doc)
                            text = (p.text or '').strip()
                        except Exception:
                            text = ''
                    if text:
                        self.log(f"  {i}: {kind} '{text[:80]}'")
                    else:
                        self.log(f"  {i}: {kind}")

            # Check if a section break exists right before or after mappings
            if mappings_idx is not None:
                try:
                    # Look for sectPr within previous or current paragraph
                    prev_el = elems[mappings_idx - 1] if mappings_idx - 1 >= 0 else None
                    cur_el = elems[mappings_idx]
                    prev_sect = prev_el.xpath('.//w:sectPr') if prev_el is not None else []
                    cur_sect = cur_el.xpath('.//w:sectPr') if cur_el is not None else []
                    self.log(f"DEBUG: Section props near MAPPINGS → prev_has_sectPr={len(prev_sect)>0}, cur_has_sectPr={len(cur_sect)>0}")
                    # Also list any standalone sectPr elements in the surrounding window and their types
                    window_start = max(0, mappings_idx - 10)
                    window_end = min(len(elems), mappings_idx + 11)
                    for i in range(window_start, window_end):
                        el = elems[i]
                        if el.tag.endswith('}sectPr') or el.tag.split('}')[-1] == 'sectPr':
                            # Determine section break type if present
                            sectType = None
                            try:
                                type_nodes = el.xpath('.//w:type')
                                if type_nodes:
                                    sectType = type_nodes[0].get(qn('w:val'))
                            except Exception:
                                sectType = None
                            self.log(f"  DEBUG: Found standalone sectPr at {i} (type={sectType})")
                except Exception:
                    pass
        except Exception:
            try:
                import traceback
                self.log(f"WARN: debug_mappings_placement encountered an exception:\n{traceback.format_exc()}")
            except Exception:
                pass

    def ensure_patent_at_issue_spacing_and_format(self, doc):
        """
        Ensure there is a page break before the Patent-at-Issue heading if preceded by the ORR table
        and normalize the heading formatting: Heading1, uppercase, bold, 12pt, color #404040,
        single line spacing, space after 0pt.
        """
        try:
            from docx.text.paragraph import Paragraph
            from docx.table import Table
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            body = doc.element.body
            elems = list(body)

            # Find Patent-at-Issue heading paragraph
            pat_p = None
            pat_idx = None
            for i, el in enumerate(elems):
                if el.tag.endswith('p'):
                    p = Paragraph(el, doc)
                    if (p.text or '').strip().lower() == 'patent-at-issue':
                        pat_p = p
                        pat_idx = i
                        break

            if pat_p is not None and pat_idx is not None:
                # Insert a page break at start if previous element is a table
                if pat_idx - 1 >= 0 and elems[pat_idx - 1].tag.endswith('tbl'):
                    br = OxmlElement('w:br')
                    br.set(qn('w:type'), 'page')
                    run_element = OxmlElement('w:r')
                    run_element.append(br)
                    pat_p._p.insert(0, run_element)

                # Normalize heading formatting
                try:
                    pat_p.style = 'Heading 1'
                except Exception:
                    try:
                        pat_p.style = 'Heading1'
                    except Exception:
                        pass
                pf = pat_p.paragraph_format
                pf.line_spacing = 1.0
                pf.space_after = Pt(0)
                if not pat_p.runs:
                    pat_p.add_run("")
                for r in pat_p.runs:
                    if r.text:
                        r.text = r.text.upper()
                    r.font.bold = True
                    r.font.size = Pt(12)
                    try:
                        r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
                    except Exception:
                        pass
        except Exception as e:
            self.log(f"Error ensuring Patent-at-Issue heading format: {str(e)}")

    def generate_report(self):
        # Document processing is now in main thread; this method is for data extraction only
        self.log("Data extraction complete in worker thread")
        return True

    def save_report(self, output_path):
        self.log("Saving report...")
        try:
            excel_name = os.path.splitext(self.excel_filename)[0] if self.excel_filename else "Report"
            date_str = datetime.now().strftime("%d%b%Y")
            if not output_path:
                output_path = f"GeneratedReport_{excel_name}_{date_str}.docx"
            self.log(f"Saving document to {output_path}...")
            self.doc.save(output_path)
            self.log(f"Document saved successfully to {output_path}")
            return output_path
        except Exception as e:
            import traceback
            self.log(f"Error in save_report: {str(e)}\n{traceback.format_exc()}")
            raise

class GenerationThread(QThread):
    """
    Background thread for processing patent report generation.
    
    This thread handles the data extraction phase of report generation,
    keeping the GUI responsive while processing large Excel files and
    performing web scraping operations.
    """
    log_signal = Signal(str)
    progress_signal = Signal(int, str)
    finished_signal = Signal(str, bool)
    request_save_dialog_signal = Signal()
    save_report_signal = Signal(str)
    document_process_signal = Signal()  # New signal for document processing in main thread

    def __init__(self, generator, excel_path, template_path, report_type):
        super().__init__()
        self.generator = generator
        self.excel_path = excel_path
        self.template_path = template_path
        self.report_type = report_type
        self.output_path = None

    def run(self):
        try:
            self.log_signal.emit("Thread started")
            self.log_signal.emit(f"Loading Excel file: {self.excel_path}")
            self.generator.load_excel(self.excel_path)
            self.log_signal.emit("Excel loaded, starting data extraction...")
            
            # Load edited report if in update mode
            if self.generator.update_mode:
                self.log_signal.emit("Loading edited report for update mode...")
                self.generator.load_edited_report()
                self.progress_signal.emit(5, "Loaded edited report")
            
            self.PatentAtIssue_Number = str(self.generator.df.iloc[1, 0])
            self.generator.extract_patent_at_issue_and_claims()
            self.progress_signal.emit(10, "Extracted patent and claims")
            self.generator.process_references()
            self.progress_signal.emit(20, "Processed references")
            self.generator.extract_search_results()
            self.progress_signal.emit(30, "Extracted search results")
            self.log_signal.emit("Data extraction complete, emitting document_process_signal")
            QThread.msleep(100)
            self.document_process_signal.emit()  # Trigger document processing in main thread
        except Exception as e:
            import traceback
            self.log_signal.emit(f"Error in thread data extraction: {str(e)}\n{traceback.format_exc()}")
            self.finished_signal.emit("", False)

    def set_output_path(self, output_path):
        self.output_path = output_path
        if output_path:
            self.log_signal.emit(f"Emitting save_report_signal for path: {output_path}")
            QThread.msleep(100)
            self.save_report_signal.emit(output_path)
        else:
            self.log_signal.emit("Save dialog cancelled")
            self.finished_signal.emit("", False)

class MainWindow(QMainWindow):
    """
    Main GUI window for the Patent Report Generator application.
    
    This class provides the user interface for:
    - Selecting Excel data files and Word templates
    - Choosing report types (Invalidity, FTO, etc.)
    - Monitoring report generation progress
    - Managing the report generation workflow
    """
    
    def __init__(self):
        """Initialize the main window and set up the user interface."""
        super().__init__()
        self.setWindowTitle("Report Generation Tool v.1.2")
        self.resize(800, 600)

        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(10)

        title_label = QLabel("Report Generation Tool")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title_label.setStyleSheet("font-size: 24px; font-weight: bold;")
        layout.addWidget(title_label)

        form_layout = QHBoxLayout()
        form_layout.setSpacing(10)

        self.excel_button = QPushButton("Select Excel File")
        self.excel_button.clicked.connect(self.select_excel)
        form_layout.addWidget(self.excel_button)

        self.template_button = QPushButton("Select Word Template")
        self.template_button.clicked.connect(self.select_template)
        form_layout.addWidget(self.template_button)

        self.report_type_combo = QComboBox()
        report_types = [
            "Invalidity",
            "FTO",
            "Patentability",
            "FTO Patentability",
            "Evidence of Use",
            "Design Patentability"
        ]
        self.report_type_combo.addItems(report_types)
        self.report_type_combo.setCurrentText("Invalidity")
        # Disable unsupported report types (keep only Invalidity enabled)
        for index, report_type in enumerate(report_types):
            if report_type not in ["Invalidity"]:
                self.report_type_combo.model().item(index).setEnabled(False)
        form_layout.addWidget(QLabel("Report Type:"))
        form_layout.addWidget(self.report_type_combo)

        # Add report mode selection
        self.report_mode_combo = QComboBox()
        self.report_mode_combo.addItems(["New Report", "Update Report"])
        self.report_mode_combo.setCurrentText("New Report")
        form_layout.addWidget(QLabel("Report Mode:"))
        form_layout.addWidget(self.report_mode_combo)
        self.report_mode_combo.currentTextChanged.connect(self.update_report_mode)

        # Add edited report file selection (initially hidden)
        self.edited_report_button = QPushButton("Select Existing Report")
        self.edited_report_button.clicked.connect(self.select_edited_report)
        self.edited_report_button.setVisible(False)
        form_layout.addWidget(self.edited_report_button)

        layout.addLayout(form_layout)

        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        layout.addWidget(self.log_text)

        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(True)
        self.progress_bar.setFormat("%p%")
        layout.addWidget(self.progress_bar)

        self.generate_button = QPushButton("Generate Report")
        self.generate_button.clicked.connect(self.generate_report)
        self.generate_button.setEnabled(False)
        layout.addWidget(self.generate_button)

        self.excel_path = None
        self.template_path = None
        self.edited_report_path = None
        self.report_type = "Invalidity"
        self.report_mode = "New Report"
        self.update_mode = False
        self.report_type_combo.currentTextChanged.connect(self.update_report_type)

        self.log_queue = Queue()
        self.log_timer = QTimer(self)
        self.log_timer.setSingleShot(False)
        self.log_timer.timeout.connect(self.process_log_queue)
        self.log_timer.start(100)

    def update_report_type(self, text):
        self.report_type = text
        self.check_enable_generate()

    def update_report_mode(self, text):
        self.report_mode = text
        self.update_mode = (text == "Update Report")
        self.log_text.append(f"DEBUG: Report mode changed to: {text}")
        self.log_text.append(f"DEBUG: update_mode set to: {self.update_mode}")
        self.edited_report_button.setVisible(self.update_mode)
        self.check_enable_generate()

    def select_edited_report(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select Edited Report", "", "Word Files (*.docx)")
        if path:
            self.edited_report_path = path
            self.log_text.append(f"Selected Edited Report: {path}")
            self.check_enable_generate()

    def select_excel(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select Excel File", "", "Excel Files (*.xlsx *.xls)")
        if path:
            self.excel_path = path
            self.log_text.append(f"Selected Excel: {path}")
            self.check_enable_generate()

    def select_template(self):
        template_path, _ = QFileDialog.getOpenFileName(self, "Select Word Template", "", "Word Files (*.docx)")
        if template_path:
            expected_templates = {
                "Invalidity": "Invalidity_Template.docx",
                "FTO": "FTO_Template.docx",
                "Patentability": "Patentability_Template.docx",
                "FTO Patentability": "FTO_Patentability_Template.docx",
                "Evidence of Use": "Evidence_of_Use_Template.docx",
                "Design Patentability": "Design_Patentability.docx"
            }
            selected_report_type = self.report_type_combo.currentText()
            expected_template = expected_templates.get(selected_report_type)
            template_filename = os.path.basename(template_path)

            if template_filename != expected_template:
                msg = QMessageBox()
                msg.setWindowTitle("Template Mismatch")
                msg.setText(
                    f"The selected template '{template_filename}' does not match the expected template "
                    f"'{expected_template}' for the '{selected_report_type}' report type.\n\n"
                    "Please select the correct template file."
                )
                msg.setStandardButtons(QMessageBox.StandardButton.Ok)
                msg.setIcon(QMessageBox.Icon.Warning)
                msg.exec()
                self.log_text.append(f"Template mismatch: Expected '{expected_template}' but got '{template_filename}'. Please select the correct template.")
                return

            self.template_path = template_path
            self.template_button.setText(f"Template: {os.path.basename(template_path)}")
            self.log_text.append(f"Selected template: {self.template_path}")
            self.check_enable_generate()

    def check_enable_generate(self):
        if self.excel_path and self.template_path and self.report_type:
            if self.update_mode:
                # For update mode, also need edited report
                if self.edited_report_path:
                    self.generate_button.setEnabled(True)
                else:
                    self.generate_button.setEnabled(False)
            else:
                # For New mode, just need excel and template
                self.generate_button.setEnabled(True)
        else:
            self.generate_button.setEnabled(False)

    def log_callback(self, message):
        self.log_queue.put(message)

    def process_log_queue(self):
        while not self.log_queue.empty():
            message = self.log_queue.get()
            self.log_text.append(message)

    def progress_callback(self, value, message):
        self.progress_bar.setValue(value)
        if message:
            self.progress_bar.setFormat(f"{message} (%p%)")
        else:
            self.progress_bar.setFormat("%p%")

    def generation_finished(self, output_path, success):
        self.log_queue.put(f"Entering generation_finished with path: {output_path}, success: {success}")
        try:
            self.progress_bar.setEnabled(True)
            self.log_text.setEnabled(True)
            self.log_queue.put("Re-enabled progress_bar and log_text")
            QThread.msleep(100)
            if success:
                self.log_queue.put("Showing success QMessageBox")
                QMessageBox.information(self, "Success", f"Report generated and saved to {output_path}")
            else:
                self.log_queue.put("Showing cancelled QMessageBox")
                QMessageBox.warning(self, "Cancelled", "Report generation cancelled")
            self.log_queue.put("Re-enabling generate_button")
            self.generate_button.setEnabled(True)
            self.progress_bar.setValue(0)
            self.log_queue.put("generation_finished completed")
        except Exception as e:
            import traceback
            self.log_queue.put(f"Error in generation_finished: {str(e)}\n{traceback.format_exc()}")

    def open_save_dialog(self):
        try:
            self.log_queue.put("Opening save dialog...")
            self.log_timer.stop()
            self.progress_bar.setEnabled(False)
            self.log_text.setEnabled(False)
            QApplication.processEvents()

            excel_name = os.path.splitext(os.path.basename(self.excel_path))[0] if self.excel_path else "Report"
            date_str = datetime.now().strftime("%d%b%Y")
            suggested_filename = f"GeneratedReport_{excel_name}_{date_str}.docx"
            output_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save Report",
                os.path.join(os.path.expanduser("~"), "Desktop", suggested_filename),
                "Word Documents (*.docx)"
            )

            self.progress_bar.setEnabled(True)
            self.log_text.setEnabled(True)
            self.log_timer.start(100)
            QApplication.processEvents()

            if output_path:
                self.log_queue.put(f"Selected output path: {output_path}")
                self.thread.set_output_path(output_path)
            else:
                self.log_queue.put("Save dialog cancelled by user")
                self.thread.set_output_path("")
        except Exception as e:
            import traceback
            self.log_queue.put(f"Error in open_save_dialog: {str(e)}\n{traceback.format_exc()}")
            self.progress_bar.setEnabled(True)
            self.log_text.setEnabled(True)
            self.log_timer.start(100)
            self.thread.set_output_path("")

    def save_report(self, output_path):
        try:
            self.log_queue.put(f"Processing save_report with path: {output_path}")
            saved_path = self.thread.generator.save_report(output_path)
            self.log_queue.put(f"save_report completed, emitting finished_signal")
            QThread.msleep(100)
            self.thread.finished_signal.emit(saved_path, True)
        except Exception as e:
            import traceback
            self.log_queue.put(f"Error in save_report: {str(e)}\n{traceback.format_exc()}")
            self.thread.finished_signal.emit("", False)

    def process_document(self):
        try:
            self.log_queue.put("Starting document processing in main thread...")
            self.progress_bar.setValue(40)
            self.thread.generator.load_template(self.template_path)
            # Set up update mode document structure (like colab lines 145-158)
            self.thread.generator.setup_update_mode_documents()
            self.thread.generator.process_title_page()
            self.progress_bar.setValue(50)
            self.thread.generator.process_objectives()
            self.progress_bar.setValue(60)
            self.thread.generator.process_other_related_references()
            self.thread.generator.process_patent_at_issue()
            self.progress_bar.setValue(70)
            self.thread.generator.process_criteria()
            self.progress_bar.setValue(80)
            self.thread.generator.process_mappings()
            self.progress_bar.setValue(90)
            self.thread.generator.process_search_strings()
            # Merge sections in update mode
            self.thread.generator.merge_generated_sections()
            self.progress_bar.setValue(100)
            self.log_queue.put("Document processing complete in main thread")
            self.thread.request_save_dialog_signal.emit()  # Trigger save dialog after document processing
        except Exception as e:
            import traceback
            self.log_queue.put(f"Error in process_document: {str(e)}\n{traceback.format_exc()}")
            self.thread.finished_signal.emit("", False)

    def generate_report(self):
        if not self.excel_path or not self.template_path:
            QMessageBox.warning(self, "Error", "Please select both an Excel file and a Word template.")
            self.log_text.append("Error: Excel file or Word template not selected.")
            return

        report_type = self.report_type_combo.currentText()
        self.generate_button.setEnabled(False)
        self.progress_bar.setEnabled(False)
        self.log_text.setEnabled(False)
        self.log_text.append("Starting report generation...")

        if report_type in ["Invalidity", "FTO"]:
            self.log_text.append(f"DEBUG: Creating PatentReportGenerator with:")
            self.log_text.append(f"  - report_type: {report_type}")
            self.log_text.append(f"  - update_mode: {self.update_mode}")
            self.log_text.append(f"  - edited_report_path: {self.edited_report_path}")
            
            self.thread = GenerationThread(
                PatentReportGenerator(
                    self.log_callback, 
                    self.progress_callback, 
                    report_type,
                    self.update_mode,
                    self.edited_report_path
                ),
                self.excel_path,
                self.template_path,
                report_type
            )
            self.thread.log_signal.connect(self.log_callback, Qt.ConnectionType.QueuedConnection)
            self.thread.progress_signal.connect(self.progress_callback, Qt.ConnectionType.QueuedConnection)
            self.thread.finished_signal.connect(self.generation_finished, Qt.ConnectionType.QueuedConnection)
            self.thread.request_save_dialog_signal.connect(self.open_save_dialog, Qt.ConnectionType.QueuedConnection)
            self.thread.save_report_signal.connect(self.save_report, Qt.ConnectionType.QueuedConnection)
            self.thread.document_process_signal.connect(self.process_document, Qt.ConnectionType.QueuedConnection)  # Connect new signal
            self.thread.start()
        else:
            self.log_text.append(f"Report type '{report_type}' is not yet supported.")
            self.generate_button.setEnabled(True)
            self.progress_bar.setEnabled(True)
            self.log_text.setEnabled(True)
            self.progress_bar.setValue(0)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())