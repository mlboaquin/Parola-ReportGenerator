# This script patches main.py to synchronize it with the latest Colab Notebook logic.
import re

with open('main.py', 'r', encoding='utf-8') as f:
    code = f.read()

# Define the global helpers from the notebook that are missing in main.py
global_helpers = """
# =========================================================
# Colab-imported Helper Functions and Ranks Parsing Logic
# =========================================================
def clean_text(value):
    if pd.isna(value):
        return ""
    s = str(value).strip()
    return "" if s.lower() == "nan" else s

def normalize_rank(rank_value):
    return clean_text(rank_value).strip()

def is_normal_letter_rank(rank_value):
    r = normalize_rank(rank_value).upper()
    return bool(re.fullmatch(r"[A-Z]", r))

def get_system_parent_info(rank_value):
    r = normalize_rank(rank_value)
    match = re.match(r'^\[?\\s*([A-Z])\\.\\s*["“](.+?)\\s*\\]?$', r, re.IGNORECASE)
    if not match:
        return None
    letter = match.group(1).upper()
    name = match.group(2).strip().strip('"').strip("”").strip("]").strip()
    return letter, name

def is_system_parent_rank(rank_value):
    return get_system_parent_info(rank_value) is not None

def is_system_child_rank(rank_value):
    r = normalize_rank(rank_value).upper()
    return bool(re.fullmatch(r"[A-Z]\\.\\d+", r))

def get_rank_parent_letter(rank_value):
    r = normalize_rank(rank_value).upper()
    if is_normal_letter_rank(r):
        return r
    if is_system_child_rank(r):
        return r.split(".")[0]
    parent_info = get_system_parent_info(rank_value)
    if parent_info:
        return parent_info[0]
    return ""

def get_child_number(rank_value):
    r = normalize_rank(rank_value).upper()
    if is_system_child_rank(r):
        return int(r.split(".")[1])
    return 0

def is_archive_link(url):
    return "web.archive.org" in clean_text(url).lower()

def is_video_link(url):
    url_lower = clean_text(url).lower()
    video_domains = ["youtube.com", "youtu.be", "vimeo.com", "dailymotion.com", "facebook.com/watch", "tiktok.com"]
    return any(domain in url_lower for domain in video_domains)

def find_paragraph_contains(doc, text_lower):
    search_text = text_lower.lower()
    for p in doc.paragraphs:
        if search_text in (p.text or '').lower():
            return p
    return None

def find_next_table_after(doc, paragraph):
    if paragraph is None:
        return None
    body = doc.element.body
    elems = list(body)
    try:
        idx = elems.index(paragraph._p)
    except ValueError:
        return None
    for i in range(idx + 1, len(elems)):
        el = elems[i]
        if el.tag.endswith('tbl'):
            from docx.table import Table
            return Table(el, paragraph._parent)
    return None

def rank_sort_key(ref):
    letter = get_rank_parent_letter(ref.Rank)
    if letter and len(letter) == 1 and "A" <= letter <= "Z":
        letter_index = ord(letter) - ord("A")
    else:
        letter_index = 999
    if is_system_child_rank(ref.Rank):
        return (letter_index, 1, get_child_number(ref.Rank))
    return (letter_index, 0, 0)

def number_to_letter(n):
    result = ""
    while n > 0:
        n -= 1
        result = chr(ord("A") + (n % 26)) + result
        n //= 26
    return result

def build_display_rank_map(sorted_refs):
    display_rank_map = {}
    parent_letter_to_display = {}
    main_counter = 0
    for ref in sorted_refs:
        raw_rank = clean_text(ref.Rank)
        parent_letter = get_rank_parent_letter(raw_rank)
        if is_system_child_rank(raw_rank):
            display_parent = parent_letter_to_display.get(parent_letter)
            if not display_parent:
                main_counter += 1
                display_parent = number_to_letter(main_counter)
                parent_letter_to_display[parent_letter] = display_parent
            display_rank_map[raw_rank] = f"{display_parent}.{get_child_number(raw_rank)}"
        else:
            main_counter += 1
            display_letter = number_to_letter(main_counter)
            display_rank_map[raw_rank] = display_letter
            if parent_letter:
                parent_letter_to_display[parent_letter] = display_letter
    return display_rank_map

def style_run(run, font_name="Inter", size=10, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold

def add_detail_line(ref_anchor, text, indent_cm=1.5):
    detail_para = ref_anchor.insert_paragraph_before(text)
    detail_para.paragraph_format.left_indent = Cm(indent_cm)
    detail_para.paragraph_format.space_after = Pt(0)
    detail_para.paragraph_format.space_before = Pt(0)
    if detail_para.runs:
        style_run(detail_para.runs[0], "Inter", 10, False)
    return detail_para

def get_ref_publisher(ref):
    if clean_text(ref.CurrentAssignee):
        return clean_text(ref.CurrentAssignee)
    if clean_text(ref.OriginalAssignee):
        return clean_text(ref.OriginalAssignee)
    return ""

def get_npl_link(ref):
    pub_num = clean_text(ref.PublicationNumber)
    url = clean_text(ref.URL)
    if pub_num:
        if pub_num.startswith("10."):
            return f"https://doi.org/{pub_num}", "DOI: ", pub_num
        if pub_num.startswith("http://") or pub_num.startswith("https://"):
            if "doi" in pub_num.lower() and "/10." in pub_num:
                doi_part = pub_num.split("/10.", 1)[1]
                return pub_num, "DOI: ", "10." + doi_part
            return pub_num, "Link: ", pub_num
    if url:
        return url, "Link: ", url
    return "", "Link: ", ""

def render_system_child(ref_anchor, target_doc, ref, next_ref_exists=True):
    child_no = get_child_number(ref.Rank)
    child_para = ref_anchor.insert_paragraph_before("")
    child_para.paragraph_format.left_indent = Cm(2.25)
    child_para.paragraph_format.hanging_indent = Cm(0.5)
    child_para.paragraph_format.space_after = Pt(0)
    child_para.paragraph_format.space_before = Pt(0)
    run_num = child_para.add_run(f"{child_no}. ")
    style_run(run_num, "Inter SemiBold", 10, True)
    title_text = clean_text(ref.Title)
    run_title = child_para.add_run(f'"{title_text}"')
    style_run(run_title, "Inter SemiBold", 10, True)
    publisher = get_ref_publisher(ref)
    label = "Author/Publisher: " if is_video_link(ref.URL) else "Publisher: "
    add_detail_line(ref_anchor, f"{label}{publisher}", indent_cm=2.75)
    date_label = "Archive Date" if is_archive_link(ref.URL) else "Publication Date"
    add_detail_line(ref_anchor, f"{date_label}: {clean_text(ref.PublicationDate)}", indent_cm=2.75)
    if not is_archive_link(ref.URL):
        add_detail_line(ref_anchor, "Retrieval Date:", indent_cm=2.75)
    link_url = clean_text(ref.URL)
    link_para = ref_anchor.insert_paragraph_before("")
    link_para.paragraph_format.left_indent = Cm(2.75)
    link_para.paragraph_format.space_after = Pt(0)
    link_para.paragraph_format.space_before = Pt(0)
    link_label_run = link_para.add_run("Link: ")
    style_run(link_label_run, "Inter", 10, False)
    if link_url:
        # We need a helper to reference target_doc.relate_to but we can pass it
        part = target_doc.part
        r_id = part.relate_to(link_url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        hyperlink.set(qn('w:anchor'), '')
        run_element = OxmlElement('w:r')
        run_props = OxmlElement('w:rPr')
        font_element = OxmlElement('w:rFonts')
        font_element.set(qn('w:ascii'), 'Inter')
        font_element.set(qn('w:hAnsi'), 'Inter')
        font_element.set(qn('w:cs'), 'Inter')
        run_props.append(font_element)
        size_element = OxmlElement('w:sz')
        size_element.set(qn('w:val'), '20')
        run_props.append(size_element)
        size_cs_element = OxmlElement('w:szCs')
        size_cs_element.set(qn('w:val'), '20')
        run_props.append(size_cs_element)
        color_element = OxmlElement('w:color')
        color_element.set(qn('w:val'), '0070C0')
        run_props.append(color_element)
        underline_element = OxmlElement('w:u')
        underline_element.set(qn('w:val'), 'single')
        run_props.append(underline_element)
        run_element.append(run_props)
        text_element = OxmlElement('w:t')
        text_element.text = link_url
        run_element.append(text_element)
        hyperlink.append(run_element)
        link_para._p.append(hyperlink)

def render_regular_reference_details(ref_anchor, target_doc, ref):
    if ref.isNPL:
        pub = get_ref_publisher(ref)
        is_doi_ref = False
        pub_num = clean_text(ref.PublicationNumber)
        if pub_num:
            if pub_num.startswith("10.") or ("doi.org" in pub_num.lower()):
                is_doi_ref = True
        if pub:
            author_label = "Author: " if is_doi_ref else "Author/Publisher: "
            add_detail_line(ref_anchor, f"{author_label}{pub}", indent_cm=1.5)
        if clean_text(ref.PublicationDate):
            add_detail_line(ref_anchor, f"Publication Date: {clean_text(ref.PublicationDate)}", indent_cm=1.5)
        if not is_doi_ref:
            add_detail_line(ref_anchor, "Retrieval Date:", indent_cm=1.5)
    else:
        if clean_text(ref.Title):
            add_detail_line(ref_anchor, f'"{clean_text(ref.Title)}"', indent_cm=1.5)
        if clean_text(ref.CurrentAssignee):
            if clean_text(ref.CurrentAssignee) == clean_text(ref.OriginalAssignee):
                add_detail_line(ref_anchor, f"Original & Current Assignee: {clean_text(ref.CurrentAssignee)}", indent_cm=1.5)
            else:
                add_detail_line(ref_anchor, f"Current Assignee: {clean_text(ref.CurrentAssignee)}", indent_cm=1.5)
        if clean_text(ref.OriginalAssignee):
            if clean_text(ref.CurrentAssignee) != clean_text(ref.OriginalAssignee):
                add_detail_line(ref_anchor, f"Original Assignee: {clean_text(ref.OriginalAssignee)}", indent_cm=1.5)
        if clean_text(ref.PriorityDate):
            add_detail_line(ref_anchor, f"Priority Date: {clean_text(ref.PriorityDate)}", indent_cm=1.5)
        if clean_text(ref.FilingDate):
            add_detail_line(ref_anchor, f"Filing Date: {clean_text(ref.FilingDate)}", indent_cm=1.5)
        if clean_text(ref.PublicationDate):
            add_detail_line(ref_anchor, f"Publication Date: {clean_text(ref.PublicationDate)}", indent_cm=1.5)
    
    if ref.isNPL:
        link_url, display_label, display_text = get_npl_link(ref)
        if link_url:
            link_para = ref_anchor.insert_paragraph_before("")
            link_para.paragraph_format.left_indent = Cm(1.5)
            link_para.paragraph_format.space_after = Pt(0)
            link_para.paragraph_format.space_before = Pt(0)
            link_label_run = link_para.add_run(display_label)
            style_run(link_label_run, "Inter", 10, False)
            
            part = target_doc.part
            r_id = part.relate_to(link_url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            hyperlink.set(qn('w:anchor'), '')
            run_element = OxmlElement('w:r')
            run_props = OxmlElement('w:rPr')
            font_element = OxmlElement('w:rFonts')
            font_element.set(qn('w:ascii'), 'Inter')
            font_element.set(qn('w:hAnsi'), 'Inter')
            font_element.set(qn('w:cs'), 'Inter')
            run_props.append(font_element)
            size_element = OxmlElement('w:sz')
            size_element.set(qn('w:val'), '20')
            run_props.append(size_element)
            size_cs_element = OxmlElement('w:szCs')
            size_cs_element.set(qn('w:val'), '20')
            run_props.append(size_cs_element)
            color_element = OxmlElement('w:color')
            color_element.set(qn('w:val'), '0070C0')
            run_props.append(color_element)
            underline_element = OxmlElement('w:u')
            underline_element.set(qn('w:val'), 'single')
            run_props.append(underline_element)
            run_element.append(run_props)
            text_element = OxmlElement('w:t')
            text_element.text = display_text if display_text else link_url
            run_element.append(text_element)
            hyperlink.append(run_element)
            link_para._p.append(hyperlink)

def get_mapping_display_rank(ref, sorted_refs):
    display_rank_map = build_display_rank_map(sorted_refs)
    return display_rank_map.get(clean_text(ref.Rank), clean_text(ref.Rank))

def should_include_ref_in_mapping(ref, sorted_refs):
    raw_rank = clean_text(ref.Rank)
    parent_info = get_system_parent_info(raw_rank)
    current_letter = get_rank_parent_letter(raw_rank)
    has_child_refs = any(
        is_system_child_rank(child_ref.Rank)
        and get_rank_parent_letter(child_ref.Rank) == current_letter
        for child_ref in sorted_refs
    )
    if is_system_child_rank(raw_rank):
        return True
    if parent_info or has_child_refs:
        return False
    return True
"""

# Insert global helpers at the top of main.py, e.g., right before unlock_password_protected_docx
unlock_idx = code.find("def unlock_password_protected_docx")
if unlock_idx != -1:
    code = code[:unlock_idx] + global_helpers + "\\n" + code[unlock_idx:]
    print("✓ Inserted global helpers.")
else:
    print("✗ Failed to insert global helpers.")

# Save modified code
with open('main.py', 'w', encoding='utf-8') as f:
    f.write(code)
