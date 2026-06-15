# Patch process_objectives in main.py
import re

with open('main.py', 'r', encoding='utf-8') as f:
    code = f.read()

new_process_objectives = """    def process_objectives(self):
        self.log("Processing objectives section...")
        try:
            target_doc = self.get_target_doc("objectives")
            
            if self.PatentAtIssue_Number.upper().startswith("US"):
                formatted_name = self.format_number_with_commas(self.PatentAtIssue_Number[2:])
                patent_prefix = "U.S. Patent No. "
            else:
                formatted_name = self.PatentAtIssue_Number
                patent_prefix = ""  # NO prefix for non-US patents
            claims_text_joined = self.format_claims_as_ranges(self.ClaimNumbers)

            obj_para = self.find_paragraph_with_placeholder(target_doc, "[OBJECTIVE_TEXT]")
            if obj_para:
                obj_para.text = ""  # clear the placeholder
                if self.report_type == "Invalidity":
                    run1 = obj_para.add_run(
                        f"This report presents the mappings of the various elements of {self.claim_word} {claims_text_joined} "
                        f"of {patent_prefix}{formatted_name} "
                    )
                    run1.font.name = "Inter"
                    run1.font.size = Pt(10)

                    run2 = obj_para.add_run(f"({self.short_patent_name})")
                    run2.font.name = "Inter"
                    run2.font.size = Pt(10)
                    run2.bold = True

                    run3 = obj_para.add_run(
                        " with the most relevant disclosures found from the following references in the publication search:"
                    )
                    run3.font.name = "Inter"
                    run3.font.size = Pt(10)
                else:  # FTO
                    run1 = obj_para.add_run(
                        f"This report presents the mappings of the various elements of the client’s invention with the most relevant disclosures found from the following references in the publication search:"
                    )
                    run1.font.name = "Inter"
                    run1.font.size = Pt(10)

            self.sorted_references = sorted(
                self.top_references,
                key=lambda r: (rank_sort_key(r), clean_text(r.PublicationNumber), clean_text(r.Title))
            )

            ref_anchor = self.find_paragraph_with_placeholder(target_doc, "[REFERENCE_LIST]")
            if ref_anchor:
                ref_anchor.text = ""

                numbering_part = target_doc.part.numbering_part
                if numbering_part is None:
                    from docx.parts.numbering import NumberingPart
                    numbering_part = NumberingPart.new()
                    target_doc.part.relate_to(
                        numbering_part,
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering'
                    )

                # Clear any existing numbering definitions to avoid conflicts
                abstract_nums = numbering_part.element.findall(f'.//{qn("w:abstractNum")}')
                for abstract_num in abstract_nums:
                    numbering_part.element.remove(abstract_num)

                nums = numbering_part.element.findall(f'.//{qn("w:num")}')
                for num in nums:
                    numbering_part.element.remove(num)

                # Create abstract numbering definition for letters A, B, C...
                abstractNum = OxmlElement('w:abstractNum')
                abstractNum.set(qn('w:abstractNumId'), '1')
                abstractNum.set(qn('w:restartNumberingAfterBreak'), '0')
                abstractNum.set(qn('w:multiLevelType'), 'hybridMultilevel')

                lvl = OxmlElement('w:lvl')
                lvl.set(qn('w:ilvl'), '0')

                numFmt = OxmlElement('w:numFmt')
                numFmt.set(qn('w:val'), 'upperLetter')
                lvl.append(numFmt)

                start = OxmlElement('w:start')
                start.set(qn('w:val'), '1')
                lvl.append(start)

                lvlText = OxmlElement('w:lvlText')
                lvlText.set(qn('w:val'), '%1.')
                lvl.append(lvlText)

                lvlJc = OxmlElement('w:lvlJc')
                lvlJc.set(qn('w:val'), 'left')
                lvl.append(lvlJc)

                pPr = OxmlElement('w:pPr')
                ind = OxmlElement('w:ind')
                ind.set(qn('w:left'), '851')
                ind.set(qn('w:hanging'), '425')
                pPr.append(ind)
                lvl.append(pPr)

                rPr = OxmlElement('w:rPr')
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Inter SemiBold')
                rFonts.set(qn('w:hAnsi'), 'Inter SemiBold')
                rFonts.set(qn('w:cs'), 'Inter SemiBold')
                rPr.append(rFonts)

                bold = OxmlElement('w:b')
                rPr.append(bold)

                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '20')
                rPr.append(sz)

                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), '20')
                rPr.append(szCs)

                lvl.append(rPr)
                abstractNum.append(lvl)
                numbering_part.element.append(abstractNum)

                # Default numbering definition
                default_abstract = OxmlElement('w:abstractNum')
                default_abstract.set(qn('w:abstractNumId'), '0')

                default_lvl = OxmlElement('w:lvl')
                default_lvl.set(qn('w:ilvl'), '0')

                default_numFmt = OxmlElement('w:numFmt')
                default_numFmt.set(qn('w:val'), 'decimal')
                default_lvl.append(default_numFmt)

                default_lvlText = OxmlElement('w:lvlText')
                default_lvlText.set(qn('w:val'), '%1.')
                default_lvl.append(default_lvlText)

                default_abstract.append(default_lvl)
                numbering_part.element.append(default_abstract)

                # Concrete numbering instance
                num = OxmlElement('w:num')
                num.set(qn('w:numId'), '1')

                abstractNumId = OxmlElement('w:abstractNumId')
                abstractNumId.set(qn('w:val'), '1')
                num.append(abstractNumId)

                numbering_part.element.append(num)

                if target_doc.part.numbering_part is None:
                    target_doc.part.relate_to(
                        numbering_part,
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering'
                    )

                for i, ref in enumerate(self.sorted_references):
                    self.isUSPatent(ref)

                    # System child ranks like F.1, F.2 are rendered as nested items, not main letters
                    if is_system_child_rank(ref.Rank):
                        render_system_child(
                            ref_anchor,
                            target_doc,
                            ref,
                            next_ref_exists=(i < len(self.sorted_references) - 1)
                        )

                        # Add normal spacing after the last child before the next main reference
                        if i < len(self.sorted_references) - 1:
                            next_ref = self.sorted_references[i + 1]

                            current_letter = get_rank_parent_letter(ref.Rank)
                            next_letter = get_rank_parent_letter(next_ref.Rank)

                            # No spacer between D.1 and D.2, but spacer after D.2 before E
                            if not (is_system_child_rank(next_ref.Rank) and current_letter == next_letter):
                                spacer = ref_anchor.insert_paragraph_before("")
                                spacer.paragraph_format.left_indent = Cm(1.5)
                                spacer.paragraph_format.space_after = Pt(0)
                                spacer.paragraph_format.space_before = Pt(0)

                        continue

                    # Main lettered paragraph
                    main_para = ref_anchor.insert_paragraph_before()

                    if main_para._p.pPr is not None:
                        main_para._p.remove(main_para._p.pPr)

                    main_para.paragraph_format.left_indent = Cm(1.5)
                    main_para.paragraph_format.hanging_indent = Cm(0.75)
                    main_para.paragraph_format.space_after = Pt(0)
                    main_para.paragraph_format.space_before = Pt(18) if i == 0 else Pt(0)

                    # Apply custom lettered numbering
                    pPr = main_para._p.get_or_add_pPr()
                    numPr = OxmlElement('w:numPr')

                    ilvl = OxmlElement('w:ilvl')
                    ilvl.set(qn('w:val'), '0')

                    numId = OxmlElement('w:numId')
                    numId.set(qn('w:val'), '1')

                    numPr.append(ilvl)
                    numPr.append(numId)
                    pPr.append(numPr)

                    jc = OxmlElement('w:jc')
                    jc.set(qn('w:val'), 'left')
                    pPr.append(jc)

                    # Determine main display text
                    parent_info = get_system_parent_info(ref.Rank)

                    if parent_info:
                        system_name = parent_info[1]
                        pub_text = f'"{system_name}"'

                    elif ref.isNPL:
                        pub_text = f'"{clean_text(ref.Title)}"'

                    elif clean_text(ref.PublicationNumber).startswith("US"):
                        if '/' in (ref.PublicationName or ""):
                            pub_text = f"U.S. Pat. App. Pub. No. {ref.PublicationName}"
                        else:
                            pub_text = f"U.S. Patent No. {ref.PublicationName or ''}"

                    else:
                        pub_text = clean_text(ref.PublicationNumber)

                    run_pub = main_para.add_run(pub_text)
                    style_run(run_pub, "Inter SemiBold", 10, True)

                    current_letter = get_rank_parent_letter(ref.Rank)
                    has_child_refs = any(
                        is_system_child_rank(child_ref.Rank)
                        and get_rank_parent_letter(child_ref.Rank) == current_letter
                        for child_ref in self.sorted_references
                    )

                    if parent_info or has_child_refs:
                        publisher = get_ref_publisher(ref)
                        add_detail_line(ref_anchor, f"Publisher: {publisher}", indent_cm=1.5)

                    else:
                        render_regular_reference_details(ref_anchor, target_doc, ref)

                    # Spacer after main references only if next item is another main reference.
                    if i < len(self.sorted_references) - 1:
                        next_ref = self.sorted_references[i + 1]

                        current_letter = get_rank_parent_letter(ref.Rank)
                        next_letter = get_rank_parent_letter(next_ref.Rank)

                        if not (is_system_child_rank(next_ref.Rank) and current_letter == next_letter):
                            spacer = ref_anchor.insert_paragraph_before("")
                            spacer.paragraph_format.left_indent = Cm(1.5)
                            spacer.paragraph_format.space_after = Pt(0)
                            spacer.paragraph_format.space_before = Pt(0)

                # Remove the anchor and any empty paragraphs after it
                if ref_anchor:
                    parent = ref_anchor._element.getparent()
                    anchor_index = list(parent).index(ref_anchor._element)
                    parent.remove(ref_anchor._element)
                    while anchor_index < len(parent):
                        elem = parent[anchor_index]
                        if elem.tag.endswith('p'):
                            text_content = ''.join(elem.itertext()).strip()
                            if not text_content:
                                parent.remove(elem)
                            else:
                                break
                        else:
                            break
            else:
                self.log("Warning: [REFERENCE_LIST] placeholder not found.")
            self.log("Objectives section processed.")
        except Exception as e:
            self.log(f"Error processing objectives: {str(e)}")
            raise"""

start_p = code.find("    def process_objectives(self):")
end_p = code.find("    def add_page_break_before_paragraph(self,")

if start_p != -1 and end_p != -1:
    code = code[:start_p] + new_process_objectives + "\\n\\n" + code[end_p:]
    print("✓ Patched process_objectives.")
else:
    print("✗ Failed to patch process_objectives.")

with open('main.py', 'w', encoding='utf-8') as f:
    f.write(code)
