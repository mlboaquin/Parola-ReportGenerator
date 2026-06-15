# Patch process_other_related_references in main.py
import re

with open('main.py', 'r', encoding='utf-8') as f:
    code = f.read()

new_orr_method = """    def process_other_related_references(self):
        self.log("Processing other related references...")
        try:
            target_doc = self.get_target_doc("references")
            
            if self.include_other_related_references:
                table_rr = self.find_table_with_placeholder(target_doc, "**[REF_INDEX]**") or \\
                          self.find_table_with_placeholder(target_doc, "[REF_ENTRY]") or \\
                          self.find_table_with_placeholder(target_doc, "[REF_OWNER]")
                if table_rr:
                    # Check if an ORR heading exists anywhere
                    other_refs_heading = None
                    for paragraph in target_doc.paragraphs:
                        if "other related references" in paragraph.text.lower():
                            other_refs_heading = paragraph
                            break
                    if other_refs_heading:
                        self.add_page_break_before_paragraph(target_doc, other_refs_heading)
                    else:
                        try:
                            from docx.oxml import OxmlElement
                            from docx.oxml.ns import qn
                            new_p = OxmlElement('w:p')
                            r = OxmlElement('w:r')
                            rPr = OxmlElement('w:rPr')
                            rFonts = OxmlElement('w:rFonts')
                            rFonts.set(qn('w:ascii'), 'Inter')
                            rFonts.set(qn('w:hAnsi'), 'Inter')
                            rPr.append(rFonts)
                            b = OxmlElement('w:b')
                            rPr.append(b)
                            sz = OxmlElement('w:sz')
                            sz.set(qn('w:val'), '20')
                            rPr.append(sz)
                            szCs = OxmlElement('w:szCs')
                            szCs.set(qn('w:val'), '20')
                            rPr.append(szCs)
                            color = OxmlElement('w:color')
                            color.set(qn('w:val'), '000000')
                            rPr.append(color)
                            r.append(rPr)
                            t = OxmlElement('w:t')
                            t.text = 'OTHER RELATED REFERENCES FOUND'
                            r.append(t)
                            pPr = OxmlElement('w:pPr')
                            spacing = OxmlElement('w:spacing')
                            spacing.set(qn('w:after'), '160')
                            spacing.set(qn('w:line'), '216')
                            spacing.set(qn('w:lineRule'), 'auto')
                            pPr.append(spacing)
                            new_p.append(pPr)
                            new_p.append(r)
                            table_rr._tbl.addprevious(new_p)
                        except Exception as e:
                            self.log(f"Warning: Could not insert ORR header: {str(e)}")
                            
                    row_template = self.find_row_with_placeholder(table_rr, "**[REF_INDEX]**") or table_rr.rows[-1]
                    granted_us_patents, us_applications, foreign_patents, npl_references = [], [], [], []
                    
                    for ref in self.related_references:
                        self.isUSPatent(ref)
                        if ref.isNPL:
                            npl_references.append(ref)
                        elif ref.PublicationNumber.startswith("US"):
                            if ref.PublicationName and '/' not in str(ref.PublicationName):
                                granted_us_patents.append(ref)
                            else:
                                us_applications.append(ref)
                        else:
                            foreign_patents.append(ref)
                            
                    sorted_related_refs = granted_us_patents + us_applications + foreign_patents + npl_references
                    
                    def render_ref_into_row(row_cells, idx, ref_obj):
                        from docx.enum.table import WD_ALIGN_VERTICAL
                        for cell in row_cells:
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        self.set_cell_text(row_cells[0], str(idx), size=9, bold=True)
                        if ref_obj.isNPL:
                            self.clear_cell(row_cells[1])
                            p1 = row_cells[1].paragraphs[0]
                            r = p1.add_run(f'"{ref_obj.Title}"')
                            r.font.name = 'Inter'
                            r.font.size = Pt(9)
                            
                            link_url = None
                            pub_num = ref_obj.PublicationNumber
                            if pub_num and pub_num.lower() != "nan":
                                if str(pub_num).startswith("10."):
                                    link_url = f"https://doi.org/{pub_num}"
                                elif str(pub_num).startswith("http://") or str(pub_num).startswith("https://"):
                                    link_url = pub_num
                                    
                            if not link_url and ref_obj.URL and ref_obj.URL.lower() != "nan":
                                link_url = ref_obj.URL
                                
                            if link_url:
                                p2 = row_cells[1].add_paragraph()
                                p2.add_run("[")
                                self.add_hyperlink_to_paragraph(target_doc, p2, link_url, "Link", size=9, color_hex="0070C0")
                                p2.add_run("]")
                                self.apply_font_style(p2, size=9)
                        else:
                            if ref_obj.PublicationNumber and ref_obj.PublicationNumber.startswith("US"):
                                if ref_obj.PublicationName and '/' in str(ref_obj.PublicationName):
                                    pub_text = f"U.S. Pat. App. Pub. No. {ref_obj.PublicationName}"
                                else:
                                    pub_text = f"U.S. Patent No. {ref_obj.PublicationName or ''}"
                            else:
                                pub_text = ref_obj.PublicationNumber or "Unknown Patent Number"
                            self.set_cell_text(row_cells[1], pub_text, size=9)
                            
                        auth = ref_obj.CurrentAssignee if ref_obj.CurrentAssignee else ref_obj.OriginalAssignee
                        if auth == "nan" or auth == "":
                            auth = ""
                        auth = auth.replace(" ,", ",")
                        self.set_cell_text(row_cells[2], auth or "", size=9)
                        
                    if sorted_related_refs:
                        render_ref_into_row(row_template.cells, 1, sorted_related_refs[0])
                        for idx, ref in enumerate(sorted_related_refs[1:], start=2):
                            new_row = self.clone_row_after(table_rr, row_template)
                            render_ref_into_row(new_row.cells, idx, ref)
                    else:
                        for cell in row_template.cells:
                            self.clear_cell(cell)
                else:
                    self.log("Warning: Table_rr not found.")
            else:
                other_refs_heading = None
                for paragraph in target_doc.paragraphs:
                    if "other related references" in paragraph.text.lower():
                        other_refs_heading = paragraph
                        break
                if other_refs_heading:
                    self.add_page_break_before_paragraph(target_doc, other_refs_heading)
                    
                table_rr = self.find_table_with_placeholder(target_doc, "**[REF_INDEX]**") or \\
                           self.find_table_with_placeholder(target_doc, "[REF_ENTRY]") or \\
                           self.find_table_with_placeholder(target_doc, "[REF_OWNER]")
                if table_rr:
                    row_template = self.find_row_with_placeholder(table_rr, "**[REF_INDEX]**") or table_rr.rows[-1]
                    for cell in row_template.cells:
                        self.clear_cell(cell)
                    self.set_cell_text(row_template.cells[0], "", size=9)
                    self.set_cell_text(row_template.cells[1], "No related references found in this search.", size=9)
                    self.set_cell_text(row_template.cells[2], "", size=9)
            self.log("Other related references processed.")
        except Exception as e:
            self.log(f"Error processing other related references: {str(e)}")
            raise"""

start_p = code.find("    def process_other_related_references(self):")
end_p = code.find("    def process_patent_at_issue(self):")

if start_p != -1 and end_p != -1:
    code = code[:start_p] + new_orr_method + "\\n\\n" + code[end_p:]
    print("✓ Patched process_other_related_references.")
else:
    print("✗ Failed to patch process_other_related_references.")

with open('main.py', 'w', encoding='utf-8') as f:
    f.write(code)
